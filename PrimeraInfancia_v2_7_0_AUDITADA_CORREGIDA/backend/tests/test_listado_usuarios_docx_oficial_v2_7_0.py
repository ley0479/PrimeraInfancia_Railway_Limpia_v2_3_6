from pathlib import Path
import shutil

from docx import Document

from services.listado_usuarios_docx_service import generate_list, template_info


def _template(path: Path):
    doc = Document()
    doc.add_paragraph("UCA: {{unidad}}")
    table = doc.add_table(rows=2, cols=7)
    headers = ["No.", "NOMBRE DEL BENEFICIARIO", "No. DOCUMENTO", "TELEFONO", "ACUDIENTE", "No. Documento", "FIRMA"]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    doc.save(path)


def test_genera_docx_oficial_y_mantiene_firma_vacia(tmp_path):
    info = template_info(tmp_path, tenant_id=7)
    template = Path(info["ruta"])
    template.parent.mkdir(parents=True, exist_ok=True)
    _template(template)
    output = tmp_path / "salida.docx"
    generate_list(tmp_path, output, [{
        "PrimerNombre": "ANA", "PrimerApellido": "PEREZ", "Documento": "123",
        "Telefono": "3001112233", "Acudiente": "MARIA", "DocumentoAcudiente": "456",
    }], {"unidad": "UCA SOL"}, tenant_id=7)
    result = Document(output)
    assert "UCA SOL" in " ".join(p.text for p in result.paragraphs)
    values = [cell.text for cell in result.tables[0].rows[1].cells]
    assert values == ["1", "ANA PEREZ", "123", "3001112233", "MARIA", "456", ""]


def test_plantilla_esta_aislada_por_fundacion(tmp_path):
    first = Path(template_info(tmp_path, tenant_id=1)["ruta"])
    first.parent.mkdir(parents=True, exist_ok=True)
    _template(first)
    assert template_info(tmp_path, tenant_id=1)["existe"] is True
    assert template_info(tmp_path, tenant_id=2)["existe"] is False
