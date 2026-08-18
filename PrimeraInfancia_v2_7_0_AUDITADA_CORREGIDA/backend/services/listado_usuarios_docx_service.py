"""Plantilla oficial DOCX de listado de usuarios, aislada por fundación."""
from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path
import shutil
import zipfile

from docx import Document

from modules.seguridad.tenant_context import current_tenant_id, tenant_storage_root

TEMPLATE_NAME = "plantilla_listado_usuarios_oficial.docx"


def _tenant_id(value=None) -> int:
    try:
        parsed = int(value or current_tenant_id() or 1)
    except (TypeError, ValueError):
        parsed = 1
    return max(1, parsed)


def template_dir(data_dir, tenant_id=None) -> Path:
    path = tenant_storage_root(data_dir, _tenant_id(tenant_id)) / "official_templates" / "listado_usuarios"
    path.mkdir(parents=True, exist_ok=True)
    (path / "backups").mkdir(exist_ok=True)
    return path


def template_info(data_dir, tenant_id=None) -> dict:
    tid = _tenant_id(tenant_id)
    path = template_dir(data_dir, tid) / TEMPLATE_NAME
    return {
        "nombre": "Listado oficial de usuarios",
        "archivo": TEMPLATE_NAME,
        "hoja": "Tabla editable Word",
        "version": "corporacion",
        "tipo_formato": "listado_usuarios",
        "codigo_tipo": "listado_usuarios",
        "preservar_estilos": True,
        "existe": path.exists(),
        "tamano_bytes": path.stat().st_size if path.exists() else 0,
        "fecha_actualizacion": datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds") if path.exists() else None,
        "ruta": str(path),
        "fundacion_id": tid,
    }


def _find_attendee_table(document):
    for table in document.tables:
        text = " ".join(cell.text.lower() for row in table.rows[:2] for cell in row.cells)
        if len(table.columns) >= 6 and "beneficiario" in text and "documento" in text:
            return table
    raise ValueError("El DOCX debe contener una tabla editable con Beneficiario y Documento.")


def validate_template(path: Path) -> None:
    if path.suffix.lower() != ".docx" or not zipfile.is_zipfile(path):
        raise ValueError("La plantilla de listado debe ser un archivo Word .docx válido.")
    document = Document(str(path))
    table = _find_attendee_table(document)
    if len(table.rows) < 2:
        raise ValueError("La tabla oficial debe incluir encabezado y al menos una fila editable.")


def replace_template(data_dir, uploaded_file, tenant_id=None) -> dict:
    filename = str(getattr(uploaded_file, "filename", "") or "")
    if Path(filename).suffix.lower() != ".docx":
        raise ValueError("La plantilla de listado oficial debe ser .docx.")
    base = template_dir(data_dir, tenant_id)
    destination = base / TEMPLATE_NAME
    temporary = base / f".upload_{datetime.now().strftime('%Y%m%d%H%M%S%f')}.docx"
    uploaded_file.save(temporary)
    try:
        validate_template(temporary)
        if destination.exists():
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            shutil.copy2(destination, base / "backups" / f"{stamp}_{TEMPLATE_NAME}")
        temporary.replace(destination)
    finally:
        temporary.unlink(missing_ok=True)
    return template_info(data_dir, tenant_id)


def restore_template(data_dir, tenant_id=None) -> dict:
    base = template_dir(data_dir, tenant_id)
    candidates = sorted((base / "backups").glob(f"*_{TEMPLATE_NAME}"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not candidates:
        raise FileNotFoundError("No hay copia anterior del listado oficial para esta corporación.")
    shutil.copy2(candidates[0], base / TEMPLATE_NAME)
    return template_info(data_dir, tenant_id)


def _value(data, *keys):
    for key in keys:
        value = (data or {}).get(key)
        if value not in (None, ""):
            return str(value).strip()
    return ""


def _name(user):
    direct = _value(user, "Nombre", "nombre", "nombres", "NombreCompleto", "nombre_completo")
    if direct:
        return direct
    return " ".join(filter(None, (_value(user, "PrimerNombre", "primer_nombre"), _value(user, "SegundoNombre", "segundo_nombre"), _value(user, "PrimerApellido", "primer_apellido"), _value(user, "SegundoApellido", "segundo_apellido"))))


def _replace_markers(document, metadata):
    mapping = {
        "{{tema}}": _value(metadata, "tema"), "{{fecha}}": _value(metadata, "fecha"),
        "{{hora_inicio}}": _value(metadata, "hora_inicio"), "{{hora_final}}": _value(metadata, "hora_final"),
        "{{unidad}}": _value(metadata, "unidad", "uca", "uds"),
        "{{profesional}}": _value(metadata, "profesional", "docente"), "{{cargo}}": _value(metadata, "cargo"),
        "{{modalidad}}": _value(metadata, "modalidad"), "{{servicio}}": _value(metadata, "servicio"),
    }
    containers = list(document.paragraphs)
    for table in document.tables:
        containers.extend(p for row in table.rows for cell in row.cells for p in cell.paragraphs)
    for paragraph in containers:
        original = paragraph.text
        updated = original
        for marker, value in mapping.items():
            updated = updated.replace(marker, value)
        if updated != original:
            if paragraph.runs:
                paragraph.runs[0].text = updated
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = updated


def generate_list(data_dir, output_path, users, metadata=None, tenant_id=None) -> str:
    info = template_info(data_dir, tenant_id)
    if not info["existe"]:
        raise FileNotFoundError("Esta corporación aún no ha cargado su plantilla DOCX de listado oficial.")
    document = Document(info["ruta"])
    _replace_markers(document, metadata or {})
    table = _find_attendee_table(document)
    data_rows = list(table.rows[1:])
    required = max(len(users or []), 1)
    while len(data_rows) < required:
        table._tbl.append(deepcopy(data_rows[-1]._tr))
        data_rows = list(table.rows[1:])
    for index, row in enumerate(data_rows):
        values = [str(index + 1), "", "", "", "", "", ""]
        if index < len(users or []):
            user = users[index]
            values = [
                str(index + 1), _name(user), _value(user, "Documento", "documento", "NUI", "nui"),
                _value(user, "Telefono", "telefono", "Celular", "celular"),
                _value(user, "Acudiente", "acudiente", "NombreAcudiente", "nombre_acudiente"),
                _value(user, "DocumentoAcudiente", "documento_acudiente"), "",
            ]
        for col, cell in enumerate(row.cells):
            cell.text = values[col] if col < len(values) else ""
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return str(output)
