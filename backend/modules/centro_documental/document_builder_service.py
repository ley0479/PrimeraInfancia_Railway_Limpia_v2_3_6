from __future__ import annotations

from copy import deepcopy
import hashlib
import json
from pathlib import Path
import shutil
import subprocess
import zipfile


def sha256_file(path: Path) -> str:
    digest=hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda:stream.read(1024*1024),b""): digest.update(chunk)
    return digest.hexdigest()


def _replace_in_paragraph(paragraph, values: dict[str,str]) -> bool:
    original=paragraph.text; replaced=original
    for key,value in values.items(): replaced=replaced.replace("{{ "+key+" }}",str(value or "")).replace("{{"+key+"}}",str(value or ""))
    if replaced==original: return False
    if paragraph.runs:
        paragraph.runs[0].text=replaced
        for run in paragraph.runs[1:]: run.text=""
    else: paragraph.text=replaced
    return True


def _walk_table(table, values: dict[str,str]) -> int:
    changes=0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs: changes+=int(_replace_in_paragraph(paragraph,values))
            for nested in cell.tables: changes+=_walk_table(nested,values)
    return changes


def build_docx(template_path: Path, output_path: Path, context: dict, narrative: str = "") -> dict:
    from docx import Document
    if template_path.suffix.lower() != ".docx": raise ValueError("La generación Word requiere una plantilla DOCX aprobada.")
    if not template_path.exists(): raise FileNotFoundError("La plantilla aprobada no está disponible.")
    before=sha256_file(template_path); document=Document(str(template_path))
    values={
        "fundacion.nombre":context.get("fundacion_nombre", ""), "uds.nombre":context.get("uds", ""),
        "actividad.tema":context.get("tema", ""), "actividad.fecha":context.get("fecha", ""),
        "actividad.responsable":context.get("responsable", ""), "documento.narrativa":narrative,
    }
    changes=0
    for paragraph in document.paragraphs: changes+=int(_replace_in_paragraph(paragraph,values))
    for table in document.tables: changes+=_walk_table(table,values)
    for section in document.sections:
        for paragraph in section.header.paragraphs: changes+=int(_replace_in_paragraph(paragraph,values))
        for table in section.header.tables: changes+=_walk_table(table,values)
        for paragraph in section.footer.paragraphs: changes+=int(_replace_in_paragraph(paragraph,values))
    if narrative and not any(token in context.get("mapped_fields",[]) for token in ("documento.narrativa","actividad.descripcion")) and changes==0:
        document.add_page_break(); document.add_heading("Continuación",level=1); document.add_paragraph(narrative)
    output_path.parent.mkdir(parents=True,exist_ok=True); document.save(str(output_path))
    if sha256_file(template_path)!=before: raise RuntimeError("La plantilla original cambió durante la generación.")
    Document(str(output_path))
    return {"path":str(output_path),"sha256":sha256_file(output_path),"campos_reemplazados":changes,"original_intacto":True}


def convert_pdf(word_path: Path, output_dir: Path, timeout=90) -> Path:
    executable=shutil.which("soffice") or shutil.which("libreoffice")
    if not executable: raise RuntimeError("No existe un convertidor PDF instalado; el Word permanece disponible.")
    output_dir.mkdir(parents=True,exist_ok=True)
    process=subprocess.run([executable,"--headless","--convert-to","pdf","--outdir",str(output_dir),str(word_path)],capture_output=True,text=True,timeout=timeout,check=False)
    result=output_dir/(word_path.stem+".pdf")
    if process.returncode or not result.exists() or result.stat().st_size<5: raise RuntimeError("No se pudo generar PDF; el Word permanece disponible.")
    if result.read_bytes()[:5]!=b"%PDF-": raise RuntimeError("El archivo PDF generado no es válido.")
    return result


def build_package(output_path: Path, files: list[tuple[str,Path]], manifest: dict) -> dict:
    output_path.parent.mkdir(parents=True,exist_ok=True)
    manifest=dict(manifest); manifest["archivos"]=[]
    with zipfile.ZipFile(output_path,"w",zipfile.ZIP_DEFLATED) as archive:
        for name,path in files:
            if path and path.exists():
                archive.write(path,name); manifest["archivos"].append({"nombre":name,"sha256":sha256_file(path),"tamano":path.stat().st_size})
        archive.writestr("08_MANIFIESTO.json",json.dumps(manifest,ensure_ascii=False,indent=2))
    return {"path":str(output_path),"sha256":sha256_file(output_path),"archivos":manifest["archivos"]}
