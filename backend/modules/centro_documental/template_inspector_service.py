from __future__ import annotations

from pathlib import Path
import zipfile


def inspect_template(path: Path) -> dict:
    extension = path.suffix.lower()
    if extension == ".docx":
        return _inspect_docx(path)
    if extension in {".xlsx", ".xlsm"}:
        return _inspect_xlsx(path)
    if extension == ".pdf":
        return _inspect_pdf(path)
    raise ValueError("Formato de plantilla no soportado.")


def _inspect_docx(path: Path) -> dict:
    if not zipfile.is_zipfile(path):
        raise ValueError("El archivo DOCX no tiene una estructura válida.")
    from docx import Document
    document = Document(str(path))
    sections = []
    for index, section in enumerate(document.sections, 1):
        sections.append({
            "numero": index,
            "margen_superior": section.top_margin.pt if section.top_margin else None,
            "margen_inferior": section.bottom_margin.pt if section.bottom_margin else None,
            "orientacion": str(section.orientation),
            "encabezados": len(section.header.paragraphs),
            "pies": len(section.footer.paragraphs),
        })
    tables = [{"indice": i, "filas": len(table.rows), "columnas": max((len(row.cells) for row in table.rows), default=0)} for i, table in enumerate(document.tables)]
    editable = bool(document.tables or any("{{" in p.text or "[" in p.text for p in document.paragraphs))
    return {
        "tipo": "DOCX", "parrafos": len(document.paragraphs), "tablas": tables,
        "secciones": sections, "imagenes": len(document.inline_shapes),
        "solo_imagen": bool(document.inline_shapes and not document.tables and not any(p.text.strip() for p in document.paragraphs)),
        "estructura_editable_detectada": editable,
    }


def _inspect_xlsx(path: Path) -> dict:
    if not zipfile.is_zipfile(path):
        raise ValueError("El archivo Excel no tiene una estructura válida.")
    from openpyxl import load_workbook
    workbook = load_workbook(path, read_only=False, data_only=False, keep_vba=path.suffix.lower() == ".xlsm")
    sheets = []
    for sheet in workbook.worksheets:
        sheets.append({
            "nombre": sheet.title, "estado": sheet.sheet_state,
            "filas": sheet.max_row, "columnas": sheet.max_column,
            "celdas_combinadas": [str(value) for value in sheet.merged_cells.ranges],
            "area_impresion": str(sheet.print_area) if sheet.print_area else None,
            "orientacion": sheet.page_setup.orientation,
            "formulas": sum(1 for row in sheet.iter_rows() for cell in row if isinstance(cell.value, str) and cell.value.startswith("=")),
        })
    workbook.close()
    return {"tipo": "XLSX", "hojas": sheets, "con_macros": path.suffix.lower() == ".xlsm"}


def _inspect_pdf(path: Path) -> dict:
    try:
        import fitz
        document = fitz.open(path)
        pages = [{"numero": i + 1, "ancho": page.rect.width, "alto": page.rect.height, "caracteres_texto": len(page.get_text("text"))} for i, page in enumerate(document)]
        document.close()
    except Exception as exc:
        raise ValueError(f"No se pudo inspeccionar el PDF: {type(exc).__name__}") from exc
    return {"tipo": "PDF", "paginas": pages, "requiere_revision_coordenadas": True}


def propose_mapping(inspection: dict) -> dict:
    """Propone zonas; nunca las aprueba ni inventa coordenadas."""
    if inspection.get("tipo") == "DOCX":
        candidates = [{"target_type": "table", "target_location": f"tabla:{t['indice']}", "status": "REQUIERE_REVISION"} for t in inspection.get("tablas", [])]
    elif inspection.get("tipo") == "XLSX":
        candidates = [{"target_type": "sheet", "target_location": f"hoja:{s['nombre']}", "status": "REQUIERE_REVISION"} for s in inspection.get("hojas", [])]
    else:
        candidates = [{"target_type": "page", "target_location": f"pagina:{p['numero']}", "status": "REQUIERE_REVISION"} for p in inspection.get("paginas", [])]
    return {"estado": "PROPUESTO", "campos": [], "zonas_detectadas": candidates, "requiere_aprobacion": True}
