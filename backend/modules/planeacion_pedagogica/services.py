from __future__ import annotations

import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
from werkzeug.utils import secure_filename

from .repository import PlaneacionRepository, now_iso
from .schema import TIPOS_DOCUMENTO_GENERABLES

MESES = {
    1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril', 5: 'mayo', 6: 'junio',
    7: 'julio', 8: 'agosto', 9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
}
MESES_INV = {v: k for k, v in MESES.items()}

TIPOS_ACTIVIDAD = [
    'Encuentro comunitario', 'Encuentro en el hogar', 'Encuentro en hogar comunitario',
    'Visita al hogar', 'Actividad pedagógica', 'Seguimiento familiar', 'Reunión de equipo',
    'Entrega de evidencia'
]


def normalizar(texto: Any) -> str:
    import unicodedata
    raw = str(texto or '').strip().lower()
    raw = unicodedata.normalize('NFKD', raw)
    raw = ''.join(ch for ch in raw if not unicodedata.combining(ch))
    raw = re.sub(r'[^a-z0-9]+', ' ', raw)
    return ' '.join(raw.split())


def periodo_actual() -> str:
    return datetime.now().strftime('%Y-%m')


def split_periodo(periodo: str | None) -> tuple[int, int]:
    periodo = periodo or periodo_actual()
    try:
        anio, mes = periodo.split('-', 1)
        return int(anio), int(mes)
    except Exception:
        hoy = datetime.now()
        return hoy.year, hoy.month


def mes_nombre(mes: int) -> str:
    return MESES.get(int(mes), str(mes)).capitalize()


def extract_text_from_file(path: str, filename: str | None = None) -> str:
    filename = filename or os.path.basename(path)
    ext = os.path.splitext(filename.lower())[1]
    try:
        if ext in {'.xlsx', '.xls', '.xlsm'}:
            xl = pd.ExcelFile(path)
            chunks: list[str] = []
            for sheet in xl.sheet_names[:10]:
                df = pd.read_excel(path, sheet_name=sheet, header=None, dtype=str)
                chunks.append(f'HOJA: {sheet}')
                chunks.extend(' | '.join(str(v) for v in row if str(v).strip() and str(v).lower() != 'nan') for row in df.fillna('').values.tolist()[:500])
            return '\n'.join(chunks)
        if ext == '.csv':
            df = pd.read_csv(path, dtype=str, encoding_errors='ignore')
            return df.fillna('').to_csv(index=False)
        if ext == '.txt':
            return Path(path).read_text(encoding='utf-8', errors='ignore')
        if ext == '.docx':
            from docx import Document
            doc = Document(path)
            parts: list[str] = []
            parts.extend(p.text for p in doc.paragraphs if p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    parts.append(' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip()))
            return '\n'.join(parts)
        if ext == '.pdf':
            try:
                from pypdf import PdfReader
                reader = PdfReader(path)
                return '\n'.join(page.extract_text() or '' for page in reader.pages[:30])
            except Exception:
                return ''
    except Exception as exc:
        return f'ERROR_EXTRACCION: {exc}'
    return ''


def _field_from_text(text: str, aliases: list[str]) -> str:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    alias_norm = [normalizar(a) for a in aliases]
    for line in lines:
        ln = normalizar(line)
        for alias in alias_norm:
            if alias and (ln.startswith(alias) or f'{alias} ' in ln):
                if ':' in line:
                    val = line.split(':', 1)[1].strip()
                    if val:
                        return val
                if '|' in line:
                    parts = [p.strip() for p in line.split('|')]
                    for idx, part in enumerate(parts):
                        if normalizar(part) == alias and idx + 1 < len(parts):
                            return parts[idx + 1]
                return line
    return ''


def _detect_date(text: str, fallback_periodo: str) -> str:
    patterns = [
        r'\b(\d{4}-\d{2}-\d{2})\b',
        r'\b(\d{1,2})[/-](\d{1,2})[/-](\d{4})\b',
        r'\b(\d{1,2})\s+de\s+([a-záéíóúñ]+)\s+de\s+(\d{4})\b',
    ]
    for pat in patterns:
        m = re.search(pat, text, flags=re.I)
        if not m:
            continue
        if len(m.groups()) == 1:
            return m.group(1)
        if len(m.groups()) == 3:
            a, b, c = m.groups()
            if b.isdigit():
                return f'{int(c):04d}-{int(b):02d}-{int(a):02d}'
            mes = MESES_INV.get(normalizar(b)) or 1
            return f'{int(c):04d}-{mes:02d}-{int(a):02d}'
    anio, mes = split_periodo(fallback_periodo)
    return f'{anio:04d}-{mes:02d}-01'


def classify_tipo_actividad(text: str, explicit: str = '') -> str:
    if explicit:
        return explicit
    n = normalizar(text)
    if 'encuentro comunitario' in n:
        return 'Encuentro comunitario'
    if 'hogar comunitario' in n:
        return 'Encuentro en hogar comunitario'
    if 'encuentro en el hogar' in n or 'encuentro hogar' in n:
        return 'Encuentro en el hogar'
    if 'visita al hogar' in n or 'visita hogar' in n:
        return 'Visita al hogar'
    if 'seguimiento familiar' in n:
        return 'Seguimiento familiar'
    if 'reunion de equipo' in n or 'reunión de equipo' in text.lower():
        return 'Reunión de equipo'
    if 'evidencia' in n:
        return 'Entrega de evidencia'
    return 'Actividad pedagógica'


def parse_planeacion_payload(text: str, form: dict[str, Any]) -> dict[str, Any]:
    periodo = form.get('periodo') or form.get('mes_anio') or ''
    if not periodo:
        # Busca mes/año en el texto.
        n = normalizar(text)
        for nombre, mes in MESES_INV.items():
            m = re.search(rf'{nombre}\s+(de\s+)?(20\d{{2}})', n)
            if m:
                periodo = f'{int(m.group(2)):04d}-{mes:02d}'
                break
    periodo = periodo or periodo_actual()
    anio, mes = split_periodo(periodo)
    tema = form.get('tema') or _field_from_text(text, ['tema', 'temática', 'tematica', 'nombre de la actividad'])
    objetivo = form.get('objetivo') or _field_from_text(text, ['objetivo', 'objetivo pedagógico', 'objetivo pedagogico'])
    actividad = form.get('actividad') or _field_from_text(text, ['actividad', 'desarrollo', 'estrategia'])
    evidencia = form.get('evidencia_requerida') or _field_from_text(text, ['evidencia', 'evidencia requerida', 'soporte'])
    poblacion = form.get('poblacion_objetivo') or _field_from_text(text, ['población objetivo', 'poblacion objetivo', 'población', 'poblacion'])
    observaciones = form.get('observaciones') or _field_from_text(text, ['observaciones', 'recomendaciones'])
    tipo = classify_tipo_actividad(text, form.get('tipo_encuentro') or form.get('tipo_actividad') or '')
    fecha = form.get('fecha_programada') or _detect_date(text, periodo)
    return {
        'periodo': periodo,
        'mes': mes,
        'anio': anio,
        'tema': tema or 'Planeación pedagógica mensual',
        'objetivo': objetivo or 'Fortalecer procesos de desarrollo integral en primera infancia.',
        'actividad': actividad or 'Actividad pedagógica programada según planeación mensual.',
        'fecha_programada': fecha,
        'poblacion_objetivo': poblacion or 'Niños, niñas, familias y comunidad participante.',
        'evidencia_requerida': evidencia or 'Acta, listado de asistencia, registro fotográfico e informe.',
        'tipo_encuentro': tipo,
        'observaciones': observaciones or '',
        'validacion': {
            'texto_detectado': bool(text.strip()),
            'campos_detectados': {
                'tema': bool(tema), 'objetivo': bool(objetivo), 'actividad': bool(actividad),
                'fecha_programada': bool(fecha), 'tipo_encuentro': bool(tipo)
            }
        }
    }


def variables_from_planeacion(planeacion: dict[str, Any], actividad: dict[str, Any] | None = None) -> dict[str, str]:
    actividad = actividad or {}
    periodo = planeacion.get('periodo') or periodo_actual()
    anio, mes = split_periodo(periodo)
    return {
        'nombreFundacion': str(planeacion.get('fundacion_nombre') or 'Fundación'),
        'nombreCoordinador': str(planeacion.get('coordinador_nombre') or ''),
        'nombreDocente': str(planeacion.get('docente_nombre') or ''),
        'unidad': str(planeacion.get('unidad') or actividad.get('unidad') or ''),
        'mes': mes_nombre(mes),
        'anio': str(anio),
        'tema': str(actividad.get('tema') or planeacion.get('tema') or ''),
        'objetivo': str(actividad.get('objetivo') or planeacion.get('objetivo') or ''),
        'actividad': str(actividad.get('actividad') or planeacion.get('actividad') or ''),
        'fecha': str(actividad.get('fecha_programada') or planeacion.get('fecha_programada') or ''),
        'poblacion': str(actividad.get('poblacion_objetivo') or planeacion.get('poblacion_objetivo') or ''),
        'evidencia': str(actividad.get('evidencia_requerida') or planeacion.get('evidencia_requerida') or ''),
        'observaciones': str(planeacion.get('observaciones') or ''),
        'tipoEncuentro': str(actividad.get('tipo_actividad') or planeacion.get('tipo_encuentro') or ''),
    }


def create_planeacion_from_file(repo: PlaneacionRepository, upload_path: str, original_name: str, form: dict[str, Any]) -> dict[str, Any]:
    text = extract_text_from_file(upload_path, original_name)
    parsed = parse_planeacion_payload(text, form)
    now = now_iso()
    new_id = repo.execute(
        """
        INSERT INTO pp_planeaciones
        (coordinador_id, docente_id, unidad, periodo, mes, anio, tema, objetivo, actividad,
         fecha_programada, poblacion_objetivo, evidencia_requerida, tipo_encuentro, observaciones,
         estado, ruta_archivo, nombre_original, nombre_guardado, texto_extraido, validacion_json,
         activo, fecha_carga, fecha_creacion, fecha_actualizacion)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'CARGADA', ?, ?, ?, ?, ?, 1, ?, ?, ?)
        """,
        (
            form.get('coordinador_id') or None,
            form.get('docente_id') or None,
            form.get('unidad') or '',
            parsed['periodo'], parsed['mes'], parsed['anio'], parsed['tema'], parsed['objetivo'], parsed['actividad'],
            parsed['fecha_programada'], parsed['poblacion_objetivo'], parsed['evidencia_requerida'], parsed['tipo_encuentro'], parsed['observaciones'],
            upload_path, original_name, os.path.basename(upload_path), text[:150000], json.dumps(parsed['validacion'], ensure_ascii=False),
            now, now, now,
        ),
    )
    planeacion = repo.get_planeacion(new_id) or {}
    create_activities_and_calendar(repo, planeacion)
    repo.log('CARGAR_PLANEACION', 'pp_planeaciones', new_id, nuevos=planeacion, planeacion_id=new_id)
    return repo.get_planeacion(new_id) or planeacion


def create_activities_and_calendar(repo: PlaneacionRepository, planeacion: dict[str, Any]) -> list[dict[str, Any]]:
    now = now_iso()
    # En Fase 3 se crea una actividad principal desde la planeación. Si después se carga una matriz
    # de varias actividades, este servicio se puede ampliar sin cambiar la tabla.
    act_id = repo.execute(
        """
        INSERT INTO pp_actividades
        (planeacion_id, coordinador_id, docente_id, unidad, tipo_actividad, titulo, tema, objetivo, actividad,
         fecha_programada, poblacion_objetivo, evidencia_requerida, estado, activo, fecha_creacion, fecha_actualizacion)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'PENDIENTE', 1, ?, ?)
        """,
        (
            planeacion.get('id'), planeacion.get('coordinador_id'), planeacion.get('docente_id'), planeacion.get('unidad') or '',
            planeacion.get('tipo_encuentro') or 'Actividad pedagógica', planeacion.get('tema') or 'Actividad de planeación',
            planeacion.get('tema') or '', planeacion.get('objetivo') or '', planeacion.get('actividad') or '',
            planeacion.get('fecha_programada') or f"{planeacion.get('periodo')}-01", planeacion.get('poblacion_objetivo') or '',
            planeacion.get('evidencia_requerida') or '', now, now,
        ),
    )
    actividad = repo.fetch_one('SELECT * FROM pp_actividades WHERE id=?', (act_id,)) or {}

    # Integra con calendario inteligente de fases previas sin cambiar su diseño.
    evento_id = repo.execute(
        """
        INSERT INTO gp_calendario_eventos
        (coordinador_id, titulo, tipo, fecha, hora, estado, descripcion, color, unidad, docente_id, docente_nombre,
         responsable, prioridad, evidencia_requerida, fecha_creacion, fecha_actualizacion)
        VALUES (?, ?, ?, ?, '', 'PENDIENTE', ?, 'AMARILLO', ?, ?, ?, ?, 'MEDIA', 1, ?, ?)
        """,
        (
            planeacion.get('coordinador_id'), planeacion.get('tema') or 'Planeación pedagógica',
            planeacion.get('tipo_encuentro') or 'Actividad pedagógica', planeacion.get('fecha_programada') or f"{planeacion.get('periodo')}-01",
            planeacion.get('actividad') or '', planeacion.get('unidad') or '', planeacion.get('docente_id'), planeacion.get('docente_nombre') or '',
            planeacion.get('coordinador_nombre') or '', now, now,
        ),
    )
    entregable_id = repo.execute(
        """
        INSERT INTO gp_entregables
        (coordinador_id, unidad, tipo, titulo, descripcion, periodo, fecha_limite, prioridad, estado, responsable,
         observaciones, activo, fecha_creacion, fecha_actualizacion)
        VALUES (?, ?, ?, ?, ?, ?, ?, 'media', 'Pendiente', ?, ?, 1, ?, ?)
        """,
        (
            planeacion.get('coordinador_id'), planeacion.get('unidad') or '', 'Planeación pedagógica',
            planeacion.get('tema') or 'Planeación pedagógica mensual', planeacion.get('actividad') or '', planeacion.get('periodo'),
            planeacion.get('fecha_programada'), planeacion.get('coordinador_nombre') or '', planeacion.get('evidencia_requerida') or '', now, now,
        ),
    )
    repo.execute_update('UPDATE pp_actividades SET calendario_evento_id=?, entregable_id=?, fecha_actualizacion=? WHERE id=?', (evento_id, entregable_id, now, act_id))
    repo.log('CREAR_ACTIVIDAD_CALENDARIO', 'pp_actividades', act_id, nuevos={'evento_id': evento_id, 'entregable_id': entregable_id}, planeacion_id=planeacion.get('id'))
    return repo.fetch_all('SELECT * FROM pp_actividades WHERE planeacion_id=? AND activo=1', (planeacion.get('id'),))


def replace_in_docx(document, variables: dict[str, str]) -> None:
    def repl(text: str) -> str:
        for key, value in variables.items():
            text = text.replace('{{' + key + '}}', str(value or ''))
        return text

    for p in document.paragraphs:
        if '{{' in p.text:
            # Preserva estructura del párrafo; puede compactar runs, pero no cambia tablas ni secciones.
            full = repl(p.text)
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = full
            else:
                p.add_run(full)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if '{{' in p.text:
                        full = repl(p.text)
                        for run in p.runs:
                            run.text = ''
                        if p.runs:
                            p.runs[0].text = full
                        else:
                            p.add_run(full)


def create_default_docx(path: str, tipo: str, variables: dict[str, str]) -> None:
    from docx import Document
    doc = Document()
    doc.add_heading(tipo, 0)
    doc.add_paragraph(f"Fundación: {variables.get('nombreFundacion','')}")
    doc.add_paragraph(f"Mes/Año: {variables.get('mes','')} {variables.get('anio','')}")
    doc.add_paragraph(f"Coordinador: {variables.get('nombreCoordinador','')}")
    doc.add_paragraph(f"Docente: {variables.get('nombreDocente','')}")
    doc.add_paragraph(f"Unidad: {variables.get('unidad','')}")
    doc.add_heading('Tema', level=1)
    doc.add_paragraph(variables.get('tema',''))
    doc.add_heading('Objetivo', level=1)
    doc.add_paragraph(variables.get('objetivo',''))
    doc.add_heading('Actividad / Desarrollo', level=1)
    doc.add_paragraph(variables.get('actividad',''))
    doc.add_heading('Fecha y población', level=1)
    doc.add_paragraph(f"Fecha: {variables.get('fecha','')}")
    doc.add_paragraph(f"Población objetivo: {variables.get('poblacion','')}")
    doc.add_heading('Evidencia requerida', level=1)
    doc.add_paragraph(variables.get('evidencia',''))
    doc.add_heading('Observaciones', level=1)
    doc.add_paragraph(variables.get('observaciones',''))
    doc.save(path)


def create_pdf(path: str, tipo: str, variables: dict[str, str]) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    doc = SimpleDocTemplate(path, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    story = [Paragraph(tipo, styles['Title']), Spacer(1, 12)]
    for label, key in [
        ('Fundación', 'nombreFundacion'), ('Coordinador', 'nombreCoordinador'), ('Docente', 'nombreDocente'),
        ('Unidad', 'unidad'), ('Mes', 'mes'), ('Año', 'anio'), ('Tema', 'tema'), ('Objetivo', 'objetivo'),
        ('Actividad', 'actividad'), ('Fecha', 'fecha'), ('Población', 'poblacion'), ('Evidencia', 'evidencia'),
        ('Observaciones', 'observaciones')
    ]:
        story.append(Paragraph(f'<b>{label}:</b> {variables.get(key, "")}', styles['BodyText']))
        story.append(Spacer(1, 6))
    doc.build(story)


def fill_xlsx_template(template_path: str, output_path: str, variables: dict[str, str]) -> None:
    from openpyxl import load_workbook
    wb = load_workbook(template_path)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and '{{' in cell.value:
                    value = cell.value
                    for key, replacement in variables.items():
                        value = value.replace('{{' + key + '}}', str(replacement or ''))
                    cell.value = value
    wb.save(output_path)


def generate_documents(repo: PlaneacionRepository, output_dir: str, planeacion_id: int, tipos: list[str] | None = None) -> list[dict[str, Any]]:
    planeacion = repo.get_planeacion(planeacion_id)
    if not planeacion:
        raise ValueError('Planeación no encontrada')
    os.makedirs(output_dir, exist_ok=True)
    tipos = tipos or TIPOS_DOCUMENTO_GENERABLES
    generated: list[dict[str, Any]] = []
    actividades = planeacion.get('actividades') or [None]
    # Para documentos mensuales se usa la primera actividad; cronograma/listados pueden usar datos de todas en fases siguientes.
    actividad = actividades[0] if actividades else None
    variables = variables_from_planeacion(planeacion, actividad)

    for tipo in tipos:
        plantillas = repo.list_plantillas(tipo)
        plantilla = plantillas[0] if plantillas else None
        base_name = secure_filename(f"{tipo}_{planeacion_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}")
        created_files: list[tuple[str, str]] = []
        if plantilla and plantilla.get('ruta_archivo') and os.path.exists(plantilla['ruta_archivo']):
            ext = os.path.splitext(plantilla['ruta_archivo'])[1].lower()
            if ext == '.docx':
                from docx import Document
                out_docx = os.path.join(output_dir, base_name + '.docx')
                doc = Document(plantilla['ruta_archivo'])
                replace_in_docx(doc, variables)
                doc.save(out_docx)
                created_files.append(('docx', out_docx))
            elif ext in {'.xlsx', '.xlsm'}:
                out_xlsx = os.path.join(output_dir, base_name + '.xlsx')
                fill_xlsx_template(plantilla['ruta_archivo'], out_xlsx, variables)
                created_files.append(('xlsx', out_xlsx))
            else:
                # PDF u otros no se alteran para no dañar estructura; se genera documento complementario.
                out_docx = os.path.join(output_dir, base_name + '.docx')
                create_default_docx(out_docx, tipo, variables)
                created_files.append(('docx', out_docx))
        else:
            out_docx = os.path.join(output_dir, base_name + '.docx')
            create_default_docx(out_docx, tipo, variables)
            created_files.append(('docx', out_docx))

        # Siempre se genera PDF adicional para gerencia.
        out_pdf = os.path.join(output_dir, base_name + '.pdf')
        create_pdf(out_pdf, tipo, variables)
        created_files.append(('pdf', out_pdf))

        for formato, path in created_files:
            new_id = repo.execute(
                """
                INSERT INTO pp_documentos_generados
                (planeacion_id, actividad_id, plantilla_id, tipo_documento, nombre, nombre_guardado, ruta_archivo, formato, estado, fecha_generacion)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'GENERADO', ?)
                """,
                (planeacion_id, actividad.get('id') if actividad else None, plantilla.get('id') if plantilla else None, tipo, os.path.basename(path), os.path.basename(path), path, formato, now_iso()),
            )
            row = repo.fetch_one('SELECT * FROM pp_documentos_generados WHERE id=?', (new_id,)) or {}
            repo.log('GENERAR_DOCUMENTO', 'pp_documentos_generados', new_id, nuevos=row, planeacion_id=planeacion_id)
            generated.append(row)
    return generated


def cambiar_estado_planeacion(repo: PlaneacionRepository, planeacion_id: int, accion: str, estado_nuevo: str, observacion: str = '') -> dict[str, Any]:
    planeacion = repo.get_planeacion(planeacion_id)
    if not planeacion:
        raise ValueError('Planeación no encontrada')
    anterior = planeacion.get('estado')
    repo.execute_update('UPDATE pp_planeaciones SET estado=?, fecha_actualizacion=? WHERE id=?', (estado_nuevo, now_iso(), planeacion_id))
    ctx = repo.context()
    repo.execute(
        """
        INSERT INTO pp_aprobaciones_planeacion
        (planeacion_id, accion, estado_anterior, estado_nuevo, observacion, usuario_aprueba, fundacion_id, usuario_creador_id, fecha_accion)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (planeacion_id, accion, anterior, estado_nuevo, observacion, ctx.get('username') or 'sistema', ctx.get('fundacion_id') or 1, ctx.get('usuario_id'), now_iso()),
    )
    repo.log(accion, 'pp_planeaciones', planeacion_id, anteriores={'estado': anterior}, nuevos={'estado': estado_nuevo, 'observacion': observacion}, planeacion_id=planeacion_id)
    return repo.get_planeacion(planeacion_id) or {}


def monthly_report(repo: PlaneacionRepository, periodo: str) -> dict[str, Any]:
    planeaciones = repo.list_planeaciones(periodo=periodo)
    estado_counts: dict[str, int] = {}
    for item in planeaciones:
        estado_counts[item.get('estado') or 'SIN ESTADO'] = estado_counts.get(item.get('estado') or 'SIN ESTADO', 0) + 1
    fid = repo.context().get('fundacion_id') or 1
    docs = repo.fetch_all(
        """
        SELECT tipo_documento, COUNT(*) total, fundacion_id
        FROM pp_documentos_generados
        WHERE (fundacion_id=? OR fundacion_id IS NULL) AND substr(fecha_generacion,1,7)=?
        GROUP BY tipo_documento, fundacion_id
        ORDER BY tipo_documento
        """,
        (fid, periodo),
    )
    acts = repo.fetch_all(
        """
        SELECT tipo_actividad, estado, COUNT(*) total, fundacion_id
        FROM pp_actividades
        WHERE (fundacion_id=? OR fundacion_id IS NULL) AND substr(COALESCE(fecha_programada,''),1,7)=?
        GROUP BY tipo_actividad, estado, fundacion_id
        ORDER BY tipo_actividad, estado
        """,
        (fid, periodo),
    )
    return {
        'periodo': periodo,
        'total_planeaciones': len(planeaciones),
        'estado_planeaciones': estado_counts,
        'documentos_generados': docs,
        'actividades': acts,
        'planeaciones': planeaciones,
    }
