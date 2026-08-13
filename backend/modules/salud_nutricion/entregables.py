"""ALPHA51: Gestión de entregables del componente Salud y Nutrición.

Este módulo es deliberadamente independiente del motor Pack35 de carga,
CoreCursor y formatos oficiales. Usa el repositorio ya registrado para crear
un catálogo mensual, evidencias, documentos de soporte y paquete ZIP sin tocar
la lógica estable de procesamiento.
"""
from __future__ import annotations

import json
import os
import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from werkzeug.utils import secure_filename

from modules.seguridad.tenant_context import current_tenant_context, current_tenant_id

from .services import now_iso

ALLOWED_EVIDENCE_EXT = {'.png', '.jpg', '.jpeg', '.webp', '.pdf', '.doc', '.docx', '.xlsx', '.xls'}

ENTREGABLES_SCHEMA_SQL = """
CREATE TABLE IF NOT EXISTS sn_entregables_catalogo (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    codigo TEXT NOT NULL UNIQUE,
    nombre TEXT NOT NULL,
    descripcion TEXT,
    actividad TEXT,
    evidencias_requeridas TEXT,
    forma_entrega TEXT DEFAULT 'Físico',
    responsable TEXT DEFAULT 'Enfermera/Nutricionista',
    periodicidad TEXT DEFAULT 'Mensual',
    aplica_por_uds INTEGER DEFAULT 1,
    requiere_acta INTEGER DEFAULT 0,
    requiere_listado INTEGER DEFAULT 0,
    requiere_fotos INTEGER DEFAULT 0,
    minimo_fotos INTEGER DEFAULT 0,
    requiere_oficio INTEGER DEFAULT 0,
    requiere_formato_excel INTEGER DEFAULT 0,
    requiere_word INTEGER DEFAULT 0,
    requiere_pdf INTEGER DEFAULT 0,
    requiere_firma INTEGER DEFAULT 0,
    plantilla_asociada TEXT,
    estado TEXT DEFAULT 'ACTIVO',
    observaciones TEXT,
    fecha_creacion TEXT,
    fecha_actualizacion TEXT
);

CREATE TABLE IF NOT EXISTS sn_entregables_mes (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    catalogo_id INTEGER NOT NULL,
    codigo TEXT NOT NULL,
    mes INTEGER NOT NULL,
    anio INTEGER NOT NULL,
    uds TEXT DEFAULT 'TODAS',
    coordinador TEXT,
    fundacion_id INTEGER DEFAULT 1,
    corporacion_id INTEGER DEFAULT 1,
    responsable TEXT,
    estado TEXT DEFAULT 'pendiente',
    observaciones TEXT,
    porcentaje INTEGER DEFAULT 0,
    fecha_creacion TEXT,
    fecha_actualizacion TEXT,
    usuario_id INTEGER,
    UNIQUE(catalogo_id, mes, anio, uds, fundacion_id),
    FOREIGN KEY (catalogo_id) REFERENCES sn_entregables_catalogo(id)
);

CREATE TABLE IF NOT EXISTS sn_entregables_evidencias (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    entregable_id INTEGER NOT NULL,
    nombre_original TEXT,
    nombre_guardado TEXT,
    ruta_archivo TEXT NOT NULL,
    tipo TEXT DEFAULT 'foto',
    actividad TEXT,
    fecha_actividad TEXT,
    uds TEXT,
    responsable TEXT,
    observaciones TEXT,
    fecha_carga TEXT,
    usuario_id INTEGER,
    FOREIGN KEY (entregable_id) REFERENCES sn_entregables_mes(id)
);

CREATE TABLE IF NOT EXISTS sn_entregables_archivos (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    entregable_id INTEGER,
    tipo TEXT NOT NULL,
    nombre_archivo TEXT NOT NULL,
    ruta_archivo TEXT NOT NULL,
    estado TEXT DEFAULT 'generado',
    metadata_json TEXT,
    fecha_generacion TEXT,
    usuario_id INTEGER,
    FOREIGN KEY (entregable_id) REFERENCES sn_entregables_mes(id)
);

CREATE TABLE IF NOT EXISTS sn_entregables_validaciones (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    entregable_id INTEGER NOT NULL,
    valido INTEGER DEFAULT 0,
    pendientes_json TEXT,
    resultado_json TEXT,
    fecha_validacion TEXT,
    usuario_id INTEGER,
    FOREIGN KEY (entregable_id) REFERENCES sn_entregables_mes(id)
);

CREATE TABLE IF NOT EXISTS sn_entregables_observaciones (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    entregable_id INTEGER NOT NULL,
    observacion TEXT NOT NULL,
    estado TEXT DEFAULT 'abierta',
    fecha_creacion TEXT,
    usuario_id INTEGER,
    FOREIGN KEY (entregable_id) REFERENCES sn_entregables_mes(id)
);

CREATE INDEX IF NOT EXISTS idx_sn_entregables_mes_periodo ON sn_entregables_mes(anio, mes, uds);
CREATE INDEX IF NOT EXISTS idx_sn_entregables_mes_estado ON sn_entregables_mes(estado);
CREATE INDEX IF NOT EXISTS idx_sn_entregables_evidencias_entregable ON sn_entregables_evidencias(entregable_id);
CREATE INDEX IF NOT EXISTS idx_sn_entregables_archivos_entregable ON sn_entregables_archivos(entregable_id);
"""

CATALOGO_BASE = [
    dict(codigo='E01_REGISTRO_NOVEDADES', nombre='Registro de novedades, articulaciones en salud y garantía de derechos',
         actividad='Acta de revisión de carpeta; identificación de garantías de derechos y canalizaciones.',
         evidencias='Acta de revisión, formato de novedades, oficio/correo de articulación y evidencias fotográficas.', acta=1, listado=0, fotos=1, oficio=1, excel=1, word=1),
    dict(codigo='E02_AAVN', nombre='Preparación y consumo de AAVN',
         actividad='Orientación a familias sobre preparación y consumo de AAVN.',
         evidencias='Acta, listado de asistencia y evidencias fotográficas.', acta=1, listado=1, fotos=1, oficio=0, excel=0, word=1),
    dict(codigo='E03_LACTANCIA', nombre='Lactancia materna / extracción de lactancia',
         actividad='Promover práctica de lactancia con gestantes, madres lactantes y familias.',
         evidencias='Acta, listado de asistencia y evidencias fotográficas.', acta=1, listado=1, fotos=1, oficio=0, excel=0, word=1),
    dict(codigo='E04_FICHA_LACTANCIA', nombre='Ficha de observación de lactancia materna',
         actividad='Aplicación y diligenciamiento de formato de lactancia materna.',
         evidencias='Formato ficha de observación y soportes fotográficos.', acta=0, listado=0, fotos=1, oficio=0, excel=1, word=0),
    dict(codigo='E05_ANTROPOMETRIA', nombre='Formato de captura de medidas antropométricas',
         actividad='Toma de peso, talla y perímetro braquial; usuarios nuevos, gestantes y seguimiento a riesgo/DNT.',
         evidencias='Formato de captura por grupo, reporte de Cuéntame y fotografías de valoración.', acta=0, listado=0, fotos=1, oficio=0, excel=1, word=0),
    dict(codigo='E06_SIGNOS_FISICOS', nombre='Formato de identificación de signos físicos',
         actividad='Aplicación y diligenciamiento de formato de signos físicos.',
         evidencias='Formato Excel magnético y muestra física en informe.', acta=0, listado=0, fotos=0, oficio=0, excel=1, word=0),
    dict(codigo='E07_MONITOREO_DNT', nombre='Formato de monitoreo de signos de alarma DNT',
         actividad='Monitoreo semanal para DNT, riesgo de DNT y signos físicos.',
         evidencias='Formato Excel magnético y anexo físico en informe.', acta=0, listado=0, fotos=0, oficio=0, excel=1, word=0),
    dict(codigo='E08_SOCIALIZACION_RPP', nombre='Socialización de minuta RPP',
         actividad='Socialización de minuta de raciones alimentarias.',
         evidencias='Acta, listado de asistencia y evidencias fotográficas.', acta=1, listado=1, fotos=1, oficio=0, excel=0, word=1),
    dict(codigo='E09_CONTROL_CALIDAD', nombre='Control de calidad de raciones y alimentos',
         actividad='Verificación de raciones alimentarias, almacenamiento y entrega de RPP/Bienestarina.',
         evidencias='Acta de verificación de alimentos y evidencias de verificación/almacenamiento.', acta=1, listado=0, fotos=1, oficio=0, excel=1, word=1),
    dict(codigo='E10_LIMPIEZA_DESINFECCION', nombre='Acta de limpieza y desinfección de UCAS',
         actividad='Verificación de limpieza y desinfección en UCAS.',
         evidencias='Acta, evidencias fotográficas y muestras de saneamiento básico.', acta=1, listado=0, fotos=1, oficio=0, excel=1, word=1),
    dict(codigo='E11_ENCUENTROS_HOGAR', nombre='Encuentros en el hogar',
         actividad='Encuentros en hogares priorizados: puerperio, DNT, riesgo, enfermedad o garantía de derechos.',
         evidencias='Formato de encuentro en el hogar, mínimo 10 muestras y evidencias fotográficas.', acta=0, listado=1, fotos=1, oficio=0, excel=1, word=1),
    dict(codigo='E12_ARTICULACION_INTERINSTITUCIONAL', nombre='Articulación interinstitucional',
         actividad='Oficio de acuerdo a la situación encontrada.',
         evidencias='Oficio físico/digital.', acta=0, listado=0, fotos=0, oficio=1, excel=0, word=1),
    dict(codigo='E13_ENTREGA_RPP', nombre='Entrega de RPP',
         actividad='Evidencias fotográficas de entrega de RPP.',
         evidencias='Evidencias fotográficas de entrega de RPP.', acta=0, listado=0, fotos=1, oficio=0, excel=0, word=0),
    dict(codigo='E14_OLLAS_REFRIGERIOS', nombre='Preparaciones de ollas comunitarias o entrega de refrigerios',
         actividad='Evidencias de preparación de olla comunitaria o entrega/consumo de refrigerios.',
         evidencias='Evidencias fotográficas con participación de familias.', acta=0, listado=1, fotos=1, oficio=0, excel=0, word=1),
    dict(codigo='E15_CUALIFICACION_TH', nombre='Acta de cualificación al talento humano',
         actividad='Cualificación en lactancia, almacenamiento de alimentos y control de plagas.',
         evidencias='Acta, listado de asistencia y evidencias fotográficas.', acta=1, listado=1, fotos=1, oficio=0, excel=0, word=1),
]


def _slug(text: str) -> str:
    text = str(text or '').upper()
    text = re.sub(r'[^A-Z0-9]+', '_', text)
    return text.strip('_') or 'SIN_DATO'


def _month_name(mes: int) -> str:
    meses = ['ENERO','FEBRERO','MARZO','ABRIL','MAYO','JUNIO','JULIO','AGOSTO','SEPTIEMBRE','OCTUBRE','NOVIEMBRE','DICIEMBRE']
    return meses[mes - 1] if 1 <= mes <= 12 else str(mes)


def _split_uds(value: Any) -> list[str]:
    if isinstance(value, list):
        raw = value
    else:
        raw = re.split(r'[,|;\n]+', str(value or ''))
    uds = []
    seen = set()
    for item in raw:
        cleaned = str(item or '').strip()
        if cleaned and cleaned.upper() not in seen:
            seen.add(cleaned.upper())
            uds.append(cleaned)
    return uds or ['TODAS']


class EntregablesSaludNutricionService:
    def __init__(self, repo, output_folder: str, upload_folder: str):
        self.repo = repo
        self._output_folder = output_folder
        self._upload_folder = upload_folder

    @property
    def output_folder(self) -> Path:
        path = Path(os.fspath(self._output_folder)) / 'salud_nutricion' / 'entregables'
        path.mkdir(parents=True, exist_ok=True)
        return path

    @property
    def upload_folder(self) -> Path:
        path = Path(os.fspath(self._upload_folder)) / 'salud_nutricion' / 'entregables'
        path.mkdir(parents=True, exist_ok=True)
        return path

    def init_schema(self) -> None:
        self.repo.execute_script(ENTREGABLES_SCHEMA_SQL)
        self.seed_catalogo()

    def seed_catalogo(self) -> None:
        existing = self.repo.fetch_one('SELECT COUNT(*) AS total FROM sn_entregables_catalogo') or {'total': 0}
        if int(existing.get('total') or 0) >= 15:
            return
        now = now_iso()
        for idx, item in enumerate(CATALOGO_BASE, start=1):
            found = self.repo.fetch_one('SELECT id FROM sn_entregables_catalogo WHERE codigo = ?', (item['codigo'],))
            if found:
                continue
            self.repo.execute(
                """
                INSERT INTO sn_entregables_catalogo
                (codigo, nombre, descripcion, actividad, evidencias_requeridas, forma_entrega, responsable,
                 periodicidad, aplica_por_uds, requiere_acta, requiere_listado, requiere_fotos, minimo_fotos,
                 requiere_oficio, requiere_formato_excel, requiere_word, requiere_pdf, requiere_firma,
                 plantilla_asociada, estado, observaciones, fecha_creacion, fecha_actualizacion)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    item['codigo'], item['nombre'], item.get('descripcion') or item['nombre'], item['actividad'],
                    item['evidencias'], 'Físico', 'Enfermera/Nutricionista', 'Mensual', 1,
                    item.get('acta', 0), item.get('listado', 0), item.get('fotos', 0), 4 if item.get('fotos', 0) else 0,
                    item.get('oficio', 0), item.get('excel', 0), item.get('word', 0), 0,
                    1 if item.get('acta') or item.get('listado') else 0,
                    item['codigo'], 'ACTIVO', 'Catálogo base ALPHA51', now, now,
                ),
            )

    def catalogo(self) -> list[dict[str, Any]]:
        return self.repo.fetch_all('SELECT * FROM sn_entregables_catalogo ORDER BY id ASC')

    def crear_mes(self, payload: dict[str, Any]) -> dict[str, Any]:
        mes = int(payload.get('mes') or datetime.now().month)
        anio = int(payload.get('anio') or datetime.now().year)
        uds_list = _split_uds(payload.get('uds') or payload.get('unidades') or 'TODAS')
        coordinador = str(payload.get('coordinador') or '').strip()
        responsable = str(payload.get('responsable') or payload.get('usuario') or 'Enfermera/Nutricionista').strip()
        context = current_tenant_context()
        requested_fundacion = payload.get('fundacion_id')
        if context.role == 'SUPERADMIN' and context.allow_global and requested_fundacion:
            fundacion_id = int(requested_fundacion)
        else:
            fundacion_id = int(context.tenant_id or 1)
        corporacion_id = int(payload.get('corporacion_id') or 1)
        now = now_iso()
        creados = 0
        existentes = 0
        for cat in self.catalogo():
            for uds in uds_list:
                found = self.repo.fetch_one(
                    'SELECT id FROM sn_entregables_mes WHERE catalogo_id = ? AND mes = ? AND anio = ? AND uds = ? AND fundacion_id = ?',
                    (cat['id'], mes, anio, uds, fundacion_id),
                )
                if found:
                    existentes += 1
                    continue
                self.repo.execute(
                    """
                    INSERT INTO sn_entregables_mes
                    (catalogo_id, codigo, mes, anio, uds, coordinador, fundacion_id, corporacion_id,
                     responsable, estado, observaciones, porcentaje, fecha_creacion, fecha_actualizacion, usuario_id)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (cat['id'], cat['codigo'], mes, anio, uds, coordinador, fundacion_id, corporacion_id,
                     responsable, 'pendiente', '', 0, now, now, payload.get('usuario_id')),
                )
                creados += 1
        return {'mes': mes, 'anio': anio, 'uds': uds_list, 'creados': creados, 'existentes': existentes}

    def listar(self, filtros: dict[str, Any]) -> dict[str, Any]:
        fundacion_id = int(current_tenant_id(1) or 1)
        where = [f'COALESCE(m.fundacion_id, 1) = {fundacion_id}']
        params: list[Any] = []
        if filtros.get('mes'):
            where.append('m.mes = ?')
            params.append(int(filtros['mes']))
        if filtros.get('anio'):
            where.append('m.anio = ?')
            params.append(int(filtros['anio']))
        if filtros.get('uds'):
            where.append('m.uds = ?')
            params.append(filtros['uds'])
        if filtros.get('estado'):
            where.append('m.estado = ?')
            params.append(filtros['estado'])
        rows = self.repo.fetch_all(
            f"""
            SELECT m.*, c.nombre, c.actividad, c.evidencias_requeridas, c.requiere_acta, c.requiere_listado,
                   c.requiere_fotos, c.minimo_fotos, c.requiere_oficio, c.requiere_formato_excel,
                   c.requiere_word, c.requiere_pdf, c.plantilla_asociada,
                   (SELECT COUNT(*) FROM sn_entregables_evidencias e
                     WHERE e.entregable_id = m.id
                       AND COALESCE(e.fundacion_id, 1) = {fundacion_id}) AS fotos_cargadas,
                   (SELECT COUNT(*) FROM sn_entregables_archivos a
                     WHERE a.entregable_id = m.id
                       AND COALESCE(a.fundacion_id, 1) = {fundacion_id}) AS archivos_generados
            FROM sn_entregables_mes m
            JOIN sn_entregables_catalogo c
              ON c.id = m.catalogo_id
             AND (c.fundacion_id IS NULL OR c.fundacion_id = {fundacion_id})
            WHERE {' AND '.join(where)}
            ORDER BY m.anio DESC, m.mes DESC, m.uds ASC, c.id ASC
            LIMIT 3000
            """,
            params,
        )
        resumen = self._resumen(rows)
        return {'entregables': rows, 'resumen': resumen}

    def detalle(self, entregable_id: int) -> dict[str, Any] | None:
        fundacion_id = int(current_tenant_id(1) or 1)
        row = self.repo.fetch_one(
            f"""
            SELECT m.*, c.nombre, c.actividad, c.evidencias_requeridas, c.requiere_acta, c.requiere_listado,
                   c.requiere_fotos, c.minimo_fotos, c.requiere_oficio, c.requiere_formato_excel,
                   c.requiere_word, c.requiere_pdf, c.plantilla_asociada
            FROM sn_entregables_mes m
            JOIN sn_entregables_catalogo c
              ON c.id = m.catalogo_id
             AND (c.fundacion_id IS NULL OR c.fundacion_id = {fundacion_id})
            WHERE m.id = ?
              AND COALESCE(m.fundacion_id, 1) = {fundacion_id}
            """,
            (entregable_id,),
        )
        if not row:
            return None
        row['evidencias'] = self.repo.fetch_all('SELECT * FROM sn_entregables_evidencias WHERE entregable_id = ? ORDER BY id', (entregable_id,))
        row['archivos'] = self.repo.fetch_all('SELECT * FROM sn_entregables_archivos WHERE entregable_id = ? ORDER BY id DESC', (entregable_id,))
        row['validaciones'] = self.repo.fetch_all('SELECT * FROM sn_entregables_validaciones WHERE entregable_id = ? ORDER BY id DESC LIMIT 10', (entregable_id,))
        return row

    def _resumen(self, rows: list[dict[str, Any]]) -> dict[str, Any]:
        total = len(rows)
        completos = sum(1 for r in rows if str(r.get('estado') or '').lower() == 'completo')
        pendientes = sum(1 for r in rows if str(r.get('estado') or '').lower() == 'pendiente')
        observados = sum(1 for r in rows if str(r.get('estado') or '').lower() == 'observado')
        fotos_faltantes = 0
        for r in rows:
            minimo = int(r.get('minimo_fotos') or 0) if int(r.get('requiere_fotos') or 0) else 0
            fotos = int(r.get('fotos_cargadas') or 0)
            fotos_faltantes += max(0, minimo - fotos)
        return {
            'total': total,
            'completos': completos,
            'pendientes': pendientes,
            'observados': observados,
            'fotos_faltantes': fotos_faltantes,
            'porcentaje': round((completos / total) * 100, 1) if total else 0,
        }

    def obtener_usuarios_base(self, uds: str | None = None, limit: int = 2000) -> list[dict[str, Any]]:
        # Fuente poblacional única: Base Maestra publicada.
        try:
            if self.repo.table_exists('master_ninos'):
                count = self.repo.fetch_one('SELECT COUNT(*) AS total FROM master_ninos') or {'total': 0}
                if int(count.get('total') or 0) > 0:
                    where = ['activo = 1']
                    params = []
                    if uds and uds != 'TODAS':
                        where.append('unidad_servicio = ?')
                        params.append(uds)
                    params.append(limit)
                    return self.repo.fetch_all(
                        f"""
                        SELECT documento, nombre_completo, unidad_servicio AS unidad, coordinador, grupo_etario,
                               estado, peso, talla, perimetro_braquial, diagnostico_nutricional, vacunas,
                               carne_salud, control_crecimiento, fecha_carga AS fecha_valoracion
                        FROM master_ninos
                        WHERE {' AND '.join(where)}
                        ORDER BY unidad_servicio, nombre_completo
                        LIMIT ?
                        """,
                        params,
                    )
        except Exception:
            pass
        return []

    def subir_evidencia(self, entregable_id: int, file, meta: dict[str, Any]) -> dict[str, Any]:
        ent = self.detalle(entregable_id)
        if not ent:
            raise ValueError('Entregable no encontrado.')
        nombre_original = file.filename or 'evidencia'
        ext = os.path.splitext(nombre_original.lower())[1]
        if ext not in ALLOWED_EVIDENCE_EXT:
            raise ValueError('Tipo de archivo no permitido para evidencia.')
        actividad = meta.get('actividad') or ent.get('codigo') or 'ENTREGABLE'
        fecha = meta.get('fecha') or datetime.now().strftime('%Y-%m-%d')
        uds = meta.get('uds') or ent.get('uds') or 'UDS'
        consecutivo = (self.repo.fetch_one('SELECT COUNT(*) AS total FROM sn_entregables_evidencias WHERE entregable_id = ?', (entregable_id,)) or {}).get('total') or 0
        nombre_guardado = f"{_slug(actividad)}_{fecha}_{_slug(uds)}_{int(consecutivo)+1:02d}{ext}"
        carpeta = self.upload_folder / f"{ent.get('anio')}_{int(ent.get('mes') or 0):02d}" / _slug(uds) / _slug(ent.get('codigo'))
        carpeta.mkdir(parents=True, exist_ok=True)
        ruta = carpeta / nombre_guardado
        file.save(ruta)
        evidencia_id = self.repo.execute(
            """
            INSERT INTO sn_entregables_evidencias
            (entregable_id, nombre_original, nombre_guardado, ruta_archivo, tipo, actividad, fecha_actividad,
             uds, responsable, observaciones, fecha_carga, usuario_id)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (entregable_id, nombre_original, nombre_guardado, str(ruta), 'evidencia', actividad, fecha, uds,
             meta.get('responsable') or ent.get('responsable'), meta.get('observaciones') or '', now_iso(), meta.get('usuario_id')),
        )
        if not evidencia_id:
            latest = self.repo.fetch_one('SELECT id FROM sn_entregables_evidencias WHERE entregable_id = ? ORDER BY id DESC LIMIT 1', (entregable_id,)) or {}
            evidencia_id = int(latest.get('id') or 0)
        self._touch(entregable_id, 'en proceso')
        return {'id': evidencia_id, 'nombre_guardado': nombre_guardado}

    def validar(self, entregable_id: int, usuario_id: int | None = None) -> dict[str, Any]:
        ent = self.detalle(entregable_id)
        if not ent:
            raise ValueError('Entregable no encontrado.')
        pendientes = []
        archivos = ent.get('archivos') or []
        evidencias = ent.get('evidencias') or []
        tipos = {a.get('tipo') for a in archivos}
        if int(ent.get('requiere_acta') or 0) and 'acta' not in tipos:
            pendientes.append('Falta acta generada.')
        if int(ent.get('requiere_listado') or 0) and 'listado' not in tipos:
            pendientes.append('Falta listado de asistencia.')
        if int(ent.get('requiere_oficio') or 0) and 'oficio' not in tipos:
            pendientes.append('Falta oficio de articulación/canalización.')
        if int(ent.get('requiere_formato_excel') or 0) and 'formato' not in tipos:
            pendientes.append('Falta formato Excel asociado.')
        if int(ent.get('requiere_fotos') or 0):
            minimo = int(ent.get('minimo_fotos') or 4)
            if len(evidencias) < minimo:
                pendientes.append(f'Faltan evidencias fotográficas: {len(evidencias)}/{minimo}.')
        valido = not pendientes
        estado = 'completo' if valido else 'pendiente'
        resultado = {'valido': valido, 'pendientes': pendientes, 'evidencias': len(evidencias), 'archivos': len(archivos)}
        self.repo.execute(
            'INSERT INTO sn_entregables_validaciones (entregable_id, valido, pendientes_json, resultado_json, fecha_validacion, usuario_id) VALUES (?, ?, ?, ?, ?, ?)',
            (entregable_id, 1 if valido else 0, json.dumps(pendientes, ensure_ascii=False), json.dumps(resultado, ensure_ascii=False), now_iso(), usuario_id),
        )
        self._touch(entregable_id, estado, 100 if valido else None)
        return resultado

    def generar_acta(self, entregable_id: int, usuario_id: int | None = None) -> dict[str, Any]:
        ent = self.detalle(entregable_id)
        if not ent:
            raise ValueError('Entregable no encontrado.')
        usuarios = self.obtener_usuarios_base(ent.get('uds'), limit=80)
        path = self._document_path(ent, 'ACTA', '.docx')
        doc = Document()
        doc.add_heading('ACTA DE ENTREGABLE SALUD Y NUTRICIÓN', 0)
        self._add_doc_context(doc, ent)
        doc.add_heading(ent.get('nombre') or '', level=1)
        doc.add_paragraph(f"Actividad: {ent.get('actividad') or ''}")
        doc.add_paragraph(f"Objetivo: dar cumplimiento al entregable del componente Salud y Nutrición para {_month_name(int(ent.get('mes') or 1))} de {ent.get('anio')}.")
        doc.add_paragraph('Desarrollo: se documenta la actividad, participantes, evidencias y compromisos asociados.')
        if usuarios:
            doc.add_heading('Participantes / usuarios relacionados', level=2)
            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = 'Documento'; hdr[1].text = 'Nombre'; hdr[2].text = 'UDS/UCA'; hdr[3].text = 'Observación'
            for u in usuarios[:40]:
                row = table.add_row().cells
                row[0].text = str(u.get('documento') or '')
                row[1].text = str(u.get('nombre_completo') or '')
                row[2].text = str(u.get('unidad') or ent.get('uds') or '')
                row[3].text = str(u.get('diagnostico_nutricional') or u.get('estado') or '')
        doc.add_heading('Compromisos', level=2)
        doc.add_paragraph('1. Revisar pendientes y cargar evidencias requeridas.\n2. Validar cierre del entregable en plataforma.')
        doc.add_paragraph('\nFirma responsable: ________________________________')
        doc.save(path)
        return self._register_file(entregable_id, 'acta', path, usuario_id, {'usuarios': len(usuarios)})

    def generar_listado(self, entregable_id: int, usuario_id: int | None = None) -> dict[str, Any]:
        ent = self.detalle(entregable_id)
        if not ent:
            raise ValueError('Entregable no encontrado.')
        usuarios = self.obtener_usuarios_base(ent.get('uds'), limit=500)
        path = self._document_path(ent, 'LISTADO_ASISTENCIA', '.xlsx')
        wb = Workbook()
        ws = wb.active
        ws.title = 'Listado asistencia'
        rows = [
            ['Actividad', ent.get('nombre') or ''],
            ['UDS/UCA', ent.get('uds') or ''],
            ['Mes/Año', f"{ent.get('mes')}/{ent.get('anio')}"],
            [],
            ['No.', 'Documento', 'Nombre completo', 'UDS/UCA', 'Rol', 'Teléfono', 'Firma', 'Observaciones']
        ]
        for row in rows:
            ws.append(row)
        for idx, u in enumerate(usuarios, start=1):
            ws.append([idx, u.get('documento') or '', u.get('nombre_completo') or '', u.get('unidad') or ent.get('uds') or '', 'Participante/Familia', '', '', ''])
        for col in range(1, 9):
            ws.column_dimensions[chr(64+col)].width = 22
        wb.save(path)
        return self._register_file(entregable_id, 'listado', path, usuario_id, {'usuarios': len(usuarios)})

    def generar_oficio(self, entregable_id: int, usuario_id: int | None = None) -> dict[str, Any]:
        ent = self.detalle(entregable_id)
        if not ent:
            raise ValueError('Entregable no encontrado.')
        usuarios = self._usuarios_priorizados(ent.get('uds'))
        path = self._document_path(ent, 'OFICIO_CANALIZACION', '.docx')
        doc = Document()
        doc.add_heading('OFICIO DE ARTICULACIÓN / CANALIZACIÓN EN SALUD', 0)
        self._add_doc_context(doc, ent)
        doc.add_paragraph('Asunto: Relación de usuarios para gestión, articulación o seguimiento en salud.')
        doc.add_paragraph('Cordial saludo. De acuerdo con la revisión del componente Salud y Nutrición, se relacionan los usuarios que requieren verificación, orientación o canalización según la información disponible en la Base Maestra.')
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = 'Documento'; hdr[1].text = 'Nombre'; hdr[2].text = 'UDS/UCA'; hdr[3].text = 'Situación'; hdr[4].text = 'Acción requerida'
        for u in usuarios[:100]:
            row = table.add_row().cells
            row[0].text = str(u.get('documento') or '')
            row[1].text = str(u.get('nombre_completo') or '')
            row[2].text = str(u.get('unidad') or ent.get('uds') or '')
            row[3].text = str(u.get('motivo') or 'Verificación de datos de salud')
            row[4].text = 'Realizar seguimiento y reportar soporte.'
        if not usuarios:
            doc.add_paragraph('No se identificaron usuarios con alertas automáticas para los filtros seleccionados. Este oficio queda como borrador editable.')
        doc.add_paragraph('\nFirma responsable: ________________________________')
        doc.save(path)
        return self._register_file(entregable_id, 'oficio', path, usuario_id, {'usuarios_priorizados': len(usuarios)})

    def generar_formato(self, entregable_id: int, usuario_id: int | None = None) -> dict[str, Any]:
        ent = self.detalle(entregable_id)
        if not ent:
            raise ValueError('Entregable no encontrado.')
        usuarios = self.obtener_usuarios_base(ent.get('uds'), limit=5000)
        path = self._document_path(ent, 'FORMATO', '.xlsx')
        wb = Workbook()
        ws = wb.active
        ws.title = 'Formato entregable'
        ws.append(['Entregable', ent.get('nombre') or ''])
        ws.append(['UDS/UCA', ent.get('uds') or ''])
        ws.append(['Mes/Año', f"{ent.get('mes')}/{ent.get('anio')}"])
        ws.append([])
        ws.append(['No.', 'Documento', 'Nombre completo', 'UDS/UCA', 'Grupo etario', 'Peso', 'Talla', 'Perímetro braquial', 'Diagnóstico', 'Vacunas', 'Carné salud', 'Observaciones'])
        for idx, u in enumerate(usuarios, start=1):
            ws.append([
                idx, u.get('documento') or '', u.get('nombre_completo') or '', u.get('unidad') or ent.get('uds') or '',
                u.get('grupo_etario') or '', u.get('peso') or '', u.get('talla') or '', u.get('perimetro_braquial') or '',
                u.get('diagnostico_nutricional') or '', u.get('vacunas') or '', u.get('carne_salud') or '', ''
            ])
        for col in range(1, 13):
            ws.column_dimensions[chr(64+col) if col <= 26 else 'A'].width = 20
        wb.save(path)
        return self._register_file(entregable_id, 'formato', path, usuario_id, {'usuarios': len(usuarios)})

    def generar_matriz(self, filtros: dict[str, Any], usuario_id: int | None = None) -> dict[str, Any]:
        data = self.listar(filtros)
        rows = data['entregables']
        mes = int(filtros.get('mes') or datetime.now().month)
        anio = int(filtros.get('anio') or datetime.now().year)
        path = self.output_folder / f"MATRIZ_CONTROL_ENTREGABLES_SALUD_NUTRICION_{_month_name(mes)}_{anio}.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = 'Matriz control'
        ws.append(['No.', 'Entregable', 'UDS/UCA', 'Responsable', 'Evidencia requerida', 'Evidencia cargada', 'Mínimo fotos', 'Fotos cargadas', 'Estado', 'Fecha', 'Observaciones', 'Archivos generados', 'Validación'])
        for idx, r in enumerate(rows, start=1):
            minimo = int(r.get('minimo_fotos') or 0) if int(r.get('requiere_fotos') or 0) else 0
            fotos = int(r.get('fotos_cargadas') or 0)
            ws.append([idx, r.get('nombre'), r.get('uds'), r.get('responsable'), r.get('evidencias_requeridas'), 'Sí' if fotos else 'No', minimo, fotos, r.get('estado'), r.get('fecha_actualizacion') or r.get('fecha_creacion'), r.get('observaciones'), r.get('archivos_generados'), 'OK' if r.get('estado') == 'completo' else 'Pendiente'])
        for col in range(1, 14):
            ws.column_dimensions[chr(64+col)].width = 22
        wb.save(path)
        return self._register_file(None, 'matriz', path, usuario_id, {'total_entregables': len(rows)})

    def generar_informe(self, filtros: dict[str, Any], usuario_id: int | None = None) -> dict[str, Any]:
        data = self.listar(filtros)
        rows = data['entregables']
        resumen = data['resumen']
        mes = int(filtros.get('mes') or datetime.now().month)
        anio = int(filtros.get('anio') or datetime.now().year)
        path = self.output_folder / f"INFORME_ENTREGABLES_SALUD_NUTRICION_{_month_name(mes)}_{anio}.docx"
        doc = Document()
        doc.add_heading('INFORME CONSOLIDADO DE ENTREGABLES', 0)
        doc.add_paragraph('Componente Salud y Nutrición')
        doc.add_paragraph(f"Periodo: {_month_name(mes)} {anio}")
        doc.add_paragraph(f"Cumplimiento: {resumen.get('porcentaje')}% · Completos: {resumen.get('completos')} · Pendientes: {resumen.get('pendientes')} · Fotos faltantes: {resumen.get('fotos_faltantes')}")
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = 'No.'; hdr[1].text = 'Entregable'; hdr[2].text = 'UDS/UCA'; hdr[3].text = 'Estado'; hdr[4].text = 'Observaciones'
        for idx, r in enumerate(rows, start=1):
            row = table.add_row().cells
            row[0].text = str(idx); row[1].text = str(r.get('nombre') or ''); row[2].text = str(r.get('uds') or ''); row[3].text = str(r.get('estado') or ''); row[4].text = str(r.get('observaciones') or '')
        doc.add_page_break()
        doc.add_heading('Pendientes y observaciones', level=1)
        for r in rows:
            if str(r.get('estado') or '').lower() != 'completo':
                doc.add_paragraph(f"{r.get('codigo')} - {r.get('nombre')} ({r.get('uds')}): {r.get('estado')}")
        doc.save(path)
        return self._register_file(None, 'informe', path, usuario_id, {'total_entregables': len(rows)})

    def generar_zip(self, filtros: dict[str, Any], usuario_id: int | None = None) -> dict[str, Any]:
        data = self.listar(filtros)
        rows = data['entregables']
        mes = int(filtros.get('mes') or datetime.now().month)
        anio = int(filtros.get('anio') or datetime.now().year)
        # Asegurar matriz e informe actualizados.
        matriz = self.generar_matriz(filtros, usuario_id)
        informe = self.generar_informe(filtros, usuario_id)
        zip_path = self.output_folder / f"ENTREGABLES_SALUD_NUTRICION_{_month_name(mes)}_{anio}.zip"
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(matriz['ruta_archivo'], arcname=os.path.basename(matriz['ruta_archivo']))
            zf.write(informe['ruta_archivo'], arcname=os.path.basename(informe['ruta_archivo']))
            for r in rows:
                folder = f"{str(r.get('codigo') or '').replace('E','').split('_')[0].zfill(2)}_{_slug(r.get('nombre'))}/"
                for a in self.repo.fetch_all('SELECT * FROM sn_entregables_archivos WHERE entregable_id = ?', (r['id'],)):
                    ruta = a.get('ruta_archivo')
                    if ruta and os.path.exists(ruta):
                        zf.write(ruta, arcname=folder + os.path.basename(ruta))
                for e in self.repo.fetch_all('SELECT * FROM sn_entregables_evidencias WHERE entregable_id = ?', (r['id'],)):
                    ruta = e.get('ruta_archivo')
                    if ruta and os.path.exists(ruta):
                        zf.write(ruta, arcname=folder + 'evidencias/' + os.path.basename(ruta))
        return self._register_file(None, 'zip', zip_path, usuario_id, {'total_entregables': len(rows)})

    def _register_file(self, entregable_id: int | None, tipo: str, path: Path | str, usuario_id: int | None, metadata: dict[str, Any]) -> dict[str, Any]:
        path = Path(path)
        archivo_id = self.repo.execute(
            """
            INSERT INTO sn_entregables_archivos
            (entregable_id, tipo, nombre_archivo, ruta_archivo, estado, metadata_json, fecha_generacion, usuario_id)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (entregable_id, tipo, path.name, str(path), 'generado', json.dumps(metadata, ensure_ascii=False), now_iso(), usuario_id),
        )
        if not archivo_id:
            if entregable_id:
                latest = self.repo.fetch_one('SELECT id FROM sn_entregables_archivos WHERE entregable_id = ? ORDER BY id DESC LIMIT 1', (entregable_id,)) or {}
            else:
                latest = self.repo.fetch_one('SELECT id FROM sn_entregables_archivos WHERE ruta_archivo = ? ORDER BY id DESC LIMIT 1', (str(path),)) or {}
            archivo_id = int(latest.get('id') or 0)
        if entregable_id:
            self._touch(entregable_id, 'en proceso')
        return {'archivo_id': archivo_id, 'tipo': tipo, 'nombre_archivo': path.name, 'ruta_archivo': str(path), 'download_url': f'/api/salud-nutricion/entregables/archivo/{archivo_id}'}

    def _document_path(self, ent: dict[str, Any], prefix: str, ext: str) -> Path:
        folder = self.output_folder / f"{ent.get('anio')}_{int(ent.get('mes') or 0):02d}" / _slug(ent.get('uds')) / _slug(ent.get('codigo'))
        folder.mkdir(parents=True, exist_ok=True)
        name = f"{prefix}_{_slug(ent.get('codigo'))}_{_slug(ent.get('uds'))}_{datetime.now().strftime('%Y%m%d%H%M%S')}{ext}"
        return folder / name

    def _touch(self, entregable_id: int, estado: str | None = None, porcentaje: int | None = None) -> None:
        if estado is not None and porcentaje is not None:
            self.repo.execute('UPDATE sn_entregables_mes SET estado = ?, porcentaje = ?, fecha_actualizacion = ? WHERE id = ?', (estado, porcentaje, now_iso(), entregable_id))
        elif estado is not None:
            self.repo.execute('UPDATE sn_entregables_mes SET estado = ?, fecha_actualizacion = ? WHERE id = ?', (estado, now_iso(), entregable_id))
        else:
            self.repo.execute('UPDATE sn_entregables_mes SET fecha_actualizacion = ? WHERE id = ?', (now_iso(), entregable_id))

    def _add_doc_context(self, doc: Document, ent: dict[str, Any]) -> None:
        doc.add_paragraph(f"Periodo: {_month_name(int(ent.get('mes') or 1))} {ent.get('anio')}")
        doc.add_paragraph(f"UDS/UCA: {ent.get('uds') or 'TODAS'}")
        doc.add_paragraph(f"Responsable: {ent.get('responsable') or 'Enfermera/Nutricionista'}")
        doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    def _usuarios_priorizados(self, uds: str | None = None) -> list[dict[str, Any]]:
        usuarios = self.obtener_usuarios_base(uds, limit=1000)
        priorizados = []
        for u in usuarios:
            motivos = []
            for campo, label in [('vacunas', 'Sin soporte de vacunas'), ('carne_salud', 'Sin carné de salud'), ('diagnostico_nutricional', 'Sin diagnóstico nutricional')]:
                val = str(u.get(campo) or '').strip().lower()
                if not val or val in {'no', 'sin dato', 'pendiente', 'none'}:
                    motivos.append(label)
            diag = str(u.get('diagnostico_nutricional') or '').lower()
            if any(t in diag for t in ['desnutric', 'riesgo', 'severa', 'moderada']):
                motivos.append('Alerta nutricional')
            if motivos:
                x = dict(u)
                x['motivo'] = '; '.join(motivos)
                priorizados.append(x)
        return priorizados

    def archivo(self, archivo_id: int) -> dict[str, Any] | None:
        return self.repo.fetch_one('SELECT * FROM sn_entregables_archivos WHERE id = ?', (archivo_id,))
