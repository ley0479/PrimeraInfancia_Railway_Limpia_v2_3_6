from __future__ import annotations

import calendar
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle


MESES_ES = {
    1: 'ENERO', 2: 'FEBRERO', 3: 'MARZO', 4: 'ABRIL',
    5: 'MAYO', 6: 'JUNIO', 7: 'JULIO', 8: 'AGOSTO',
    9: 'SEPTIEMBRE', 10: 'OCTUBRE', 11: 'NOVIEMBRE', 12: 'DICIEMBRE'
}

ALIAS_UNIDADES = {
    'UCA UNIDAD DEMO 01': 'UNIDAD DEMO 01', 'UNIDAD DEMO 01': 'UNIDAD DEMO 01', 'UCA UNIDAD DEMO 01': 'UNIDAD DEMO 01',
    'UCA UNIDAD DEMO 21': 'UNIDAD DEMO 21', 'UNIDAD DEMO 21': 'UNIDAD DEMO 21',
    'UCA UNIDAD DEMO 22': 'UNIDAD DEMO 22', 'UNIDAD DEMO 22': 'UNIDAD DEMO 22',
    'UCA UNIDAD DEMO 09': 'UNIDAD DEMO 09', 'UNIDAD DEMO 09': 'UNIDAD DEMO 09', 'UCA UNIDAD DEMO 12': 'UNIDAD DEMO 12', 'UNIDAD DEMO 12': 'UNIDAD DEMO 12',
    'UCA UNIDAD DEMO 02': 'UNIDAD DEMO 02', 'UNIDAD DEMO 02': 'UNIDAD DEMO 02', 'UNIDAD DEMO 02': 'UNIDAD DEMO 02',
    'UCA UNIDAD DEMO 03': 'UNIDAD DEMO 03', 'UNIDAD DEMO 03': 'UNIDAD DEMO 03',
    'UCA UNIDAD DEMO 16': 'UNIDAD DEMO 16', 'UNIDAD DEMO 16': 'UNIDAD DEMO 16',
    'UCA UNIDAD DEMO 17': 'UNIDAD DEMO 17', 'UNIDAD DEMO 17': 'UNIDAD DEMO 17',
    'UCA UNIDAD DEMO 18': 'UNIDAD DEMO 18', 'UNIDAD DEMO 18': 'UNIDAD DEMO 18', 'UNIDAD DEMO 18': 'UNIDAD DEMO 18', 'UNIDAD DEMO 18': 'UNIDAD DEMO 18',
    'UCA UNIDAD DEMO 11': 'UNIDAD DEMO 11', 'UCA UNIDAD DEMO 11': 'UNIDAD DEMO 11', 'UNIDAD DEMO 11': 'UNIDAD DEMO 11',
    'UCA UNIDAD DEMO 20': 'UNIDAD DEMO 20', 'UNIDAD DEMO 20': 'UNIDAD DEMO 20',
    'UCA UNIDAD DEMO 04': 'UNIDAD DEMO 04', 'UCA UNIDAD DEMO 05': 'UNIDAD DEMO 05', '15': 'UNIDAD DEMO 05',
}


def now_iso() -> str:
    return datetime.now().isoformat(timespec='seconds')


def limpiar_valor(valor: Any, default: str = '') -> str:
    if valor is None:
        return default
    try:
        if pd.isna(valor):
            return default
    except Exception:
        pass
    texto = str(valor).strip()
    if texto.lower() in {'nan', 'nat', 'none', 'null'}:
        return default
    return texto


def normalizar_texto(valor: Any) -> str:
    import unicodedata
    texto = limpiar_valor(valor).lower()
    texto = unicodedata.normalize('NFKD', texto)
    texto = ''.join(ch for ch in texto if not unicodedata.combining(ch))
    texto = re.sub(r'[^a-z0-9]+', ' ', texto)
    return ' '.join(texto.split())


def normalize_unidad(valor: Any) -> str:
    import unicodedata
    texto = limpiar_valor(valor).upper()
    if not texto:
        return ''
    texto = unicodedata.normalize('NFKD', texto)
    texto = ''.join(ch for ch in texto if not unicodedata.combining(ch))
    texto = re.sub(r'[,.;:]+$', '', texto)
    texto = re.sub(r'\s+', ' ', texto).strip()
    if texto in {'ACTIVO', 'INACTIVO', 'PENDIENTE', 'SIN UNIDAD', 'UNIDAD DE SERVICIO'}:
        return ''
    if texto in ALIAS_UNIDADES:
        return ALIAS_UNIDADES[texto]
    if texto.startswith('UCA '):
        sin = texto[4:].strip()
        return ALIAS_UNIDADES.get(sin, sin)
    return ALIAS_UNIDADES.get(texto, texto)


def parse_fecha(valor: Any) -> str:
    texto = limpiar_valor(valor)
    if not texto:
        return ''
    if isinstance(valor, datetime):
        return valor.date().isoformat()
    if re.fullmatch(r'\d+(\.\d+)?', texto):
        try:
            numero = float(texto)
            if numero > 20000:
                return pd.to_datetime(numero, unit='D', origin='1899-12-30').date().isoformat()
        except Exception:
            pass
    for dayfirst in (True, False):
        try:
            fecha = pd.to_datetime(texto, errors='coerce', dayfirst=dayfirst)
            if pd.notna(fecha):
                return fecha.date().isoformat()
        except Exception:
            pass
    return texto


def edad_texto_desde_fecha(fecha_nacimiento: str) -> str:
    fecha = parse_fecha(fecha_nacimiento)
    try:
        f = datetime.fromisoformat(fecha)
    except Exception:
        return ''
    hoy = datetime.now()
    meses = (hoy.year - f.year) * 12 + hoy.month - f.month
    if hoy.day < f.day:
        meses -= 1
    meses = max(0, meses)
    return f"{meses // 12} años {meses % 12} meses"


def buscar_columna(df: pd.DataFrame, aliases: list[str]) -> str | None:
    columnas = [(normalizar_texto(c), c) for c in df.columns]
    alias_norm = [normalizar_texto(a) for a in aliases]
    for alias in alias_norm:
        for col_norm, original in columnas:
            if col_norm == alias:
                return original
    for alias in alias_norm:
        for col_norm, original in columnas:
            if alias and alias in col_norm:
                return original
    return None


def serie(df: pd.DataFrame, columna: str | None, default: str = '') -> pd.Series:
    if columna and columna in df.columns:
        return df[columna].apply(lambda x: limpiar_valor(x, default))
    return pd.Series([default] * len(df), index=df.index)


def leer_tabla_desde_texto(texto: str) -> pd.DataFrame:
    lineas = [l.strip() for l in str(texto or '').splitlines() if l.strip()]
    if not lineas:
        raise ValueError('El archivo no contiene texto tabular legible.')
    muestra = '\n'.join(lineas[:80])
    from io import StringIO
    for sep in ['\t', ';', ',', '|']:
        if sep in muestra:
            df = pd.read_csv(StringIO('\n'.join(lineas)), sep=sep, dtype=str, engine='python')
            if not df.empty and len(df.columns) > 1:
                return df
    df = pd.read_csv(StringIO('\n'.join(lineas)), sep=r'\s{2,}', dtype=str, engine='python')
    if not df.empty and len(df.columns) > 1:
        return df
    raise ValueError('No se detectó una tabla válida en el texto.')


def read_tabular_file(path: str) -> pd.DataFrame:
    ext = os.path.splitext(path.lower())[1]
    if ext in {'.xlsx', '.xls', '.xlsm', '.ods'}:
        hojas = pd.read_excel(path, sheet_name=None, dtype=str)
        mejor = None
        mejor_score = -1
        for df in hojas.values():
            if df is None or df.empty:
                continue
            columnas = ' '.join(normalizar_texto(c) for c in df.columns)
            score = len(df) + len(df.columns) * 5
            for k in ['documento', 'beneficiario', 'unidad', 'acudiente', 'telefono', 'direccion']:
                if k in columnas:
                    score += 100
            if score > mejor_score:
                mejor = df
                mejor_score = score
        if mejor is not None:
            return mejor
        raise ValueError('No se encontró hoja con datos en el libro.')
    if ext in {'.csv', '.txt', '.tsv', '.tab', '.dat'}:
        for enc in ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']:
            try:
                sep = '\t' if ext in {'.tsv', '.tab'} else None
                df = pd.read_csv(path, sep=sep, dtype=str, engine='python', encoding=enc)
                if not df.empty:
                    return df
            except Exception:
                pass
        with open(path, 'r', encoding='latin-1', errors='ignore') as fh:
            return leer_tabla_desde_texto(fh.read())
    if ext in {'.html', '.htm'}:
        tablas = pd.read_html(path)
        if tablas:
            return tablas[0].astype(str)
    if ext == '.json':
        return pd.read_json(path, dtype=str)
    if ext == '.docx':
        from docx import Document
        doc = Document(path)
        tablas = []
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if len(rows) > 1 and len(rows[0]) > 1:
                tablas.append(pd.DataFrame(rows[1:], columns=rows[0]))
        if tablas:
            tablas.sort(key=lambda d: len(d) * len(d.columns), reverse=True)
            return tablas[0].astype(str)
        return leer_tabla_desde_texto('\n'.join(p.text for p in doc.paragraphs))
    if ext == '.pdf':
        from pypdf import PdfReader
        reader = PdfReader(path)
        return leer_tabla_desde_texto('\n'.join(page.extract_text() or '' for page in reader.pages))
    raise ValueError(f'Formato no soportado para cruce de bases: {ext}')


def normalizar_base(df: pd.DataFrame) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    df = df.copy()
    df.columns = [limpiar_valor(c) for c in df.columns]
    c_documento = buscar_columna(df, ['Documento del beneficiario', 'Número de documento del beneficiario', 'Numero de documento del beneficiario', 'documento', 'nui', 'nuip', 'identificacion', 'identificación'])
    c_tipo_doc = buscar_columna(df, ['Tipo de documento del beneficiario', 'tipo documento', 'tipo doc'])
    c_primer_nombre = buscar_columna(df, ['Primer Nombre del beneficiario', 'Primer Nombre'])
    c_segundo_nombre = buscar_columna(df, ['Segundo Nombre del beneficiario', 'Segundo Nombre'])
    c_primer_apellido = buscar_columna(df, ['Primer apellido del beneficiario', 'Primer Apellido'])
    c_segundo_apellido = buscar_columna(df, ['Segundo apellido del beneficiario', 'Segundo Apellido'])
    c_nombre = buscar_columna(df, ['Nombre completo del beneficiario', 'Nombre', 'Nombres y apellidos'])
    c_fecha = buscar_columna(df, ['Fecha de nacimiento del beneficiario', 'Fecha nacimiento', 'fecha_nacimiento'])
    c_unidad = buscar_columna(df, ['Nombre de la unidad de servicio', 'Unidad de servicio', 'Unidad', 'UCA', 'UDS'])
    c_grupo = buscar_columna(df, ['Grupo', 'Nombre Tipo de beneficiario', 'Tipo de beneficiario'])
    c_docente = buscar_columna(df, ['Docente', 'Agente educativo', 'Educador'])
    c_acudiente = buscar_columna(df, ['Nombre completo del acudiente', 'Nombre completo responsable', 'Nombre acudiente', 'Nombre responsable'])
    if c_acudiente and any(k in normalizar_texto(c_acudiente) for k in ['tipo documento', 'documento del acudiente', 'numero de documento']):
        c_acudiente = None
    c_doc_acudiente = buscar_columna(df, ['Número de documento del acudiente o responsable', 'Documento acudiente', 'Documento responsable'])
    c_pn_ac = buscar_columna(df, ['Primer nombre del acudiente o responsable'])
    c_sn_ac = buscar_columna(df, ['Segundo nombre del acudiente o responsable'])
    c_pa_ac = buscar_columna(df, ['Primer apellido del acudiente o responsable'])
    c_sa_ac = buscar_columna(df, ['Segundo apellido del acudiente o responsable'])
    c_tel = buscar_columna(df, ['Teléfono del beneficiario', 'Telefono del beneficiario', 'Teléfono', 'Telefono', 'Celular', 'Número de celular'])
    c_dir = buscar_columna(df, ['Direccion de residencia del beneficiario', 'Dirección de residencia', 'Direccion', 'Dirección'])
    c_sexo = buscar_columna(df, ['Sexo del beneficiario', 'Sexo'])
    c_estado = buscar_columna(df, ['Estado', 'Estado del beneficiario'])

    errores = []
    registros = []
    vistos = set()
    if not c_documento:
        errores.append({'fila': 0, 'campo': 'documento', 'error': 'No se detectó columna de documento/NUI.'})
        return [], errores
    if not c_unidad:
        errores.append({'fila': 0, 'campo': 'unidad', 'error': 'No se detectó columna de unidad de servicio.'})

    for idx, row in df.iterrows():
        documento = limpiar_valor(row.get(c_documento)).replace('.0', '')
        if not documento:
            errores.append({'fila': int(idx) + 2, 'campo': 'documento', 'error': 'Documento vacío.'})
            continue
        pn = limpiar_valor(row.get(c_primer_nombre)).upper() if c_primer_nombre else ''
        sn = limpiar_valor(row.get(c_segundo_nombre)).upper() if c_segundo_nombre else ''
        pa = limpiar_valor(row.get(c_primer_apellido)).upper() if c_primer_apellido else ''
        sa = limpiar_valor(row.get(c_segundo_apellido)).upper() if c_segundo_apellido else ''
        nombre = ' '.join(p for p in [pn, sn, pa, sa] if p).strip()
        if not nombre:
            nombre = limpiar_valor(row.get(c_nombre)).upper() if c_nombre else ''
        ac = ' '.join(limpiar_valor(row.get(c)).upper() for c in [c_pn_ac, c_sn_ac, c_pa_ac, c_sa_ac] if c).strip()
        if not ac and c_acudiente:
            ac = limpiar_valor(row.get(c_acudiente)).upper()
        unidad = normalize_unidad(row.get(c_unidad)) if c_unidad else ''
        registro = {
            'documento': documento,
            'tipo_documento': limpiar_valor(row.get(c_tipo_doc)).upper() if c_tipo_doc else '',
            'nombre': nombre,
            'primer_nombre': pn,
            'segundo_nombre': sn,
            'primer_apellido': pa,
            'segundo_apellido': sa,
            'fecha_nacimiento': parse_fecha(row.get(c_fecha)) if c_fecha else '',
            'unidad': unidad,
            'grupo': limpiar_valor(row.get(c_grupo)).upper() if c_grupo else '',
            'docente': limpiar_valor(row.get(c_docente)).upper() if c_docente else '',
            'acudiente': ac,
            'documento_acudiente': limpiar_valor(row.get(c_doc_acudiente)).replace('.0', '') if c_doc_acudiente else '',
            'telefono': limpiar_valor(row.get(c_tel)).replace('.0', '') if c_tel else '',
            'direccion': limpiar_valor(row.get(c_dir)).upper() if c_dir else '',
            'sexo': limpiar_valor(row.get(c_sexo)).upper() if c_sexo else '',
            'estado': limpiar_valor(row.get(c_estado)).upper() if c_estado else 'ACTIVO',
            'edad': edad_texto_desde_fecha(row.get(c_fecha)) if c_fecha else '',
        }
        if documento in vistos:
            errores.append({'fila': int(idx) + 2, 'campo': 'documento', 'error': f'Documento duplicado crítico: {documento}'})
        vistos.add(documento)
        registros.append(registro)
    return registros, errores


def key_text(value: Any) -> str:
    return normalizar_texto(value)


def comparar_bases(anterior: list[dict[str, Any]], actual: list[dict[str, Any]]) -> dict[str, Any]:
    ant = {r['documento']: r for r in anterior if r.get('documento')}
    act = {r['documento']: r for r in actual if r.get('documento')}
    docs_ant = set(ant)
    docs_act = set(act)
    nuevos = [act[d] for d in sorted(docs_act - docs_ant)]
    retirados = [ant[d] for d in sorted(docs_ant - docs_act)]
    comunes = sorted(docs_ant & docs_act)

    trasladados = []
    cambios = []
    cambios_unidad = []
    cambios_docente = []
    cambios_acudiente = []
    cambios_telefono = []
    cambios_direccion = []
    campos = [
        ('unidad', 'Unidad'), ('grupo', 'Grupo'), ('docente', 'Docente'),
        ('acudiente', 'Acudiente'), ('documento_acudiente', 'Documento acudiente'),
        ('telefono', 'Teléfono'), ('direccion', 'Dirección'),
        ('nombre', 'Nombre'), ('fecha_nacimiento', 'Fecha nacimiento'), ('sexo', 'Sexo')
    ]
    for doc in comunes:
        a = ant[doc]
        b = act[doc]
        cambios_doc = []
        for campo, etiqueta in campos:
            if key_text(a.get(campo)) != key_text(b.get(campo)):
                cambio = {
                    'documento': doc,
                    'nombre': b.get('nombre') or a.get('nombre'),
                    'campo': etiqueta,
                    'campo_clave': campo,
                    'anterior': a.get(campo, ''),
                    'actual': b.get(campo, ''),
                    'unidad_anterior': a.get('unidad', ''),
                    'unidad_actual': b.get('unidad', ''),
                    'docente_anterior': a.get('docente', ''),
                    'docente_actual': b.get('docente', ''),
                }
                cambios_doc.append(cambio)
                if campo == 'unidad':
                    cambios_unidad.append(cambio)
                elif campo == 'docente':
                    cambios_docente.append(cambio)
                elif campo in {'acudiente', 'documento_acudiente'}:
                    cambios_acudiente.append(cambio)
                elif campo == 'telefono':
                    cambios_telefono.append(cambio)
                elif campo == 'direccion':
                    cambios_direccion.append(cambio)
        if key_text(a.get('unidad')) != key_text(b.get('unidad')) or key_text(a.get('docente')) != key_text(b.get('docente')):
            trasladados.append({
                'documento': doc,
                'nombre': b.get('nombre') or a.get('nombre'),
                'unidad_anterior': a.get('unidad', ''),
                'unidad_actual': b.get('unidad', ''),
                'docente_anterior': a.get('docente', ''),
                'docente_actual': b.get('docente', ''),
            })
        if cambios_doc:
            cambios.append({
                'documento': doc,
                'nombre': b.get('nombre') or a.get('nombre'),
                'unidad_anterior': a.get('unidad', ''),
                'unidad_actual': b.get('unidad', ''),
                'docente_anterior': a.get('docente', ''),
                'docente_actual': b.get('docente', ''),
                'cambios': cambios_doc,
            })

    # Reemplazo operativo: un retiro y un ingreso en la misma unidad durante el mismo corte.
    reemplazados = []
    usados_nuevos = set()
    for r in retirados:
        mejor_idx = None
        for idx, n in enumerate(nuevos):
            if idx in usados_nuevos:
                continue
            if key_text(n.get('unidad')) == key_text(r.get('unidad')):
                mejor_idx = idx
                break
        if mejor_idx is not None:
            usados_nuevos.add(mejor_idx)
            n = nuevos[mejor_idx]
            reemplazados.append({
                'documento_retirado': r.get('documento'),
                'nombre_retirado': r.get('nombre'),
                'documento_nuevo': n.get('documento'),
                'nombre_nuevo': n.get('nombre'),
                'unidad': n.get('unidad') or r.get('unidad'),
                'docente': n.get('docente') or r.get('docente'),
                'fecha_corte': datetime.now().date().isoformat(),
                'observacion': 'Ingreso y retiro detectados en la misma unidad durante el corte mensual.',
                'unidad_retirado': r.get('unidad'),
                'unidad_nuevo': n.get('unidad'),
                'docente_retirado': r.get('docente'),
                'docente_nuevo': n.get('docente'),
            })

    resumen = {
        'total_anterior': len(anterior),
        'total_actual': len(actual),
        'nuevos': len(nuevos),
        'retirados': len(retirados),
        'reemplazados': len(reemplazados),
        'trasladados': len(trasladados),
        'cambios_unidad': len(cambios_unidad),
        'cambios_docente': len(cambios_docente),
        'cambios_acudiente': len(cambios_acudiente),
        'cambios_telefono': len(cambios_telefono),
        'cambios_direccion': len(cambios_direccion),
        'cambios_total': len(cambios),
    }
    return {
        'resumen': resumen,
        'nuevos': nuevos,
        'retirados': retirados,
        'reemplazados': reemplazados,
        'trasladados': trasladados,
        'cambios': cambios,
        'cambios_unidad': cambios_unidad,
        'cambios_docente': cambios_docente,
        'cambios_acudiente': cambios_acudiente,
        'cambios_telefono': cambios_telefono,
        'cambios_direccion': cambios_direccion,
        'errores': [],
    }


def resumen_a_tabla(item: dict[str, Any], tipo: str) -> list[Any]:
    if tipo == 'reemplazados':
        return [item.get('nombre_retirado', ''), item.get('documento_retirado', ''), item.get('nombre_nuevo', ''), item.get('documento_nuevo', ''), item.get('unidad', ''), item.get('docente', ''), item.get('fecha_corte', ''), item.get('observacion', '')]
    if tipo == 'trasladados':
        return [item.get('nombre', ''), item.get('documento', ''), item.get('unidad_anterior', ''), item.get('unidad_actual', ''), item.get('docente_anterior', ''), item.get('docente_actual', '')]
    if tipo == 'cambios':
        return [item.get('nombre', ''), item.get('documento', ''), '; '.join(f"{c.get('campo')}: {c.get('anterior')} → {c.get('actual')}" for c in item.get('cambios', [])), item.get('unidad_actual', ''), item.get('docente_actual', '')]
    if tipo.startswith('cambios_'):
        return [item.get('nombre', ''), item.get('documento', ''), item.get('campo', ''), item.get('anterior', ''), item.get('actual', ''), item.get('unidad_actual', ''), item.get('docente_actual', '')]
    return [item.get('documento', ''), item.get('nombre', ''), item.get('unidad', ''), item.get('docente', ''), item.get('acudiente', ''), item.get('telefono', ''), item.get('direccion', '')]


def headers_tipo(tipo: str) -> list[str]:
    if tipo == 'reemplazados':
        return ['Niño retirado', 'Documento retirado', 'Niño nuevo', 'Documento nuevo', 'Unidad', 'Docente', 'Fecha corte', 'Observación']
    if tipo == 'trasladados':
        return ['Niño', 'Documento', 'Unidad anterior', 'Unidad actual', 'Docente anterior', 'Docente actual']
    if tipo == 'cambios':
        return ['Niño', 'Documento', 'Cambios detectados', 'Unidad actual', 'Docente actual']
    if tipo.startswith('cambios_'):
        return ['Niño', 'Documento', 'Campo', 'Anterior', 'Actual', 'Unidad', 'Docente']
    return ['Documento', 'Nombre completo', 'Unidad', 'Docente', 'Acudiente', 'Teléfono', 'Dirección']


def generar_excel_resultado(resultado: dict[str, Any], output_folder: str, prefix: str = 'CRUCE_BASES') -> str:
    os.makedirs(output_folder, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = 'Resumen'
    headers = ['Indicador', 'Valor']
    ws.append(headers)
    for k, v in (resultado.get('resumen') or {}).items():
        ws.append([k.replace('_', ' ').title(), v])
    for cell in ws[1]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill('solid', fgColor='1F4E78')
        cell.alignment = Alignment(horizontal='center')
    ws.column_dimensions['A'].width = 32
    ws.column_dimensions['B'].width = 14

    for tipo in ['nuevos', 'retirados', 'reemplazados', 'trasladados', 'cambios', 'cambios_unidad', 'cambios_docente', 'cambios_acudiente', 'cambios_telefono', 'cambios_direccion']:
        data = resultado.get(tipo) or []
        sh = wb.create_sheet(tipo[:31])
        h = headers_tipo(tipo)
        sh.append(h)
        for item in data:
            sh.append(resumen_a_tabla(item, tipo))
        for cell in sh[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill('solid', fgColor='305496')
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        for col in range(1, len(h) + 1):
            sh.column_dimensions[chr(64 + col) if col <= 26 else 'Z'].width = 24
        sh.freeze_panes = 'A2'
    nombre = f"{prefix}_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
    path = os.path.join(output_folder, nombre)
    wb.save(path)
    return path


def generar_pdf_resultado(resultado: dict[str, Any], output_folder: str, tipo: str = 'resumen', prefix: str = 'CRUCE_BASES') -> str:
    os.makedirs(output_folder, exist_ok=True)
    nombre = f"{prefix}_{tipo}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
    path = os.path.join(output_folder, nombre)
    doc = SimpleDocTemplate(path, pagesize=landscape(letter), rightMargin=24, leftMargin=24, topMargin=24, bottomMargin=24)
    styles = getSampleStyleSheet()
    story = [Paragraph('Reporte de cruce mensual de bases', styles['Title']), Spacer(1, 10)]
    if tipo == 'resumen':
        rows = [['Indicador', 'Valor']] + [[k.replace('_', ' ').title(), str(v)] for k, v in (resultado.get('resumen') or {}).items()]
    else:
        rows = [headers_tipo(tipo)] + [resumen_a_tabla(item, tipo) for item in resultado.get(tipo, [])]
    rows = [[str(c or '')[:180] for c in row] for row in rows]
    table = Table(rows, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTSIZE', (0, 0), (-1, -1), 7),
    ]))
    story.append(table)
    doc.build(story)
    return path


def insertar_alertas_salud_nutricion(database_path: str, resultado: dict[str, Any], fundacion_id: int = 1) -> None:
    """Inserta alertas nutricionales desde el cruce usando SQLAlchemy Core."""
    from modules.sqlalchemy_compat import CoreCompatRepository

    repo = CoreCompatRepository()
    if not repo.table_exists('sn_alertas'):
        return
    cols = repo.columns('sn_alertas')
    extra = ', fundacion_id' if 'fundacion_id' in cols else ''
    extra_val = ', ?' if 'fundacion_id' in cols else ''
    params_extra = [fundacion_id] if 'fundacion_id' in cols else []

    for item in resultado.get('nuevos', []):
        repo.execute(
            f"""
            INSERT INTO sn_alertas (documento, tipo, nivel, mensaje, unidad, fecha_alerta, atendida, observaciones, fecha_creacion{extra})
            VALUES (?, 'PENDIENTE_VALORACION', 'AMARILLO', ?, ?, ?, 0, ?, ?{extra_val})
            """,
            [item.get('documento'), f"Nuevo ingreso pendiente por valoración nutricional: {item.get('nombre')}", item.get('unidad'), datetime.now().date().isoformat(), 'Generado desde cruce mensual de bases', now_iso()] + params_extra,
        )
    for item in resultado.get('retirados', []):
        repo.execute(
            f"""
            INSERT INTO sn_alertas (documento, tipo, nivel, mensaje, unidad, fecha_alerta, atendida, observaciones, fecha_creacion{extra})
            VALUES (?, 'RETIRADO_BASE', 'GRIS', ?, ?, ?, 0, ?, ?{extra_val})
            """,
            [item.get('documento'), f"Beneficiario retirado de la base actual. Conservar historial nutricional: {item.get('nombre')}", item.get('unidad'), datetime.now().date().isoformat(), 'Generado desde cruce mensual de bases', now_iso()] + params_extra,
        )


def docente_por_unidad(database_path: str, unidad: str, fundacion_id: int | None = None) -> dict[str, Any]:
    """Busca docente por unidad usando SQLAlchemy Core."""
    from modules.sqlalchemy_compat import CoreCompatRepository

    unidad_norm = normalize_unidad(unidad)
    repo = CoreCompatRepository()
    if not repo.table_exists('coordinadores'):
        return {}
    cols = repo.columns('coordinadores')
    where = "activo = 1"
    params: list[Any] = []
    if fundacion_id and 'fundacion_id' in cols:
        where += " AND (fundacion_id = ? OR fundacion_id IS NULL)"
        params.append(fundacion_id)
    rows = repo.fetch_all(
        f"""
        SELECT * FROM coordinadores
        WHERE {where}
        ORDER BY CASE
            WHEN upper(COALESCE(tipo_equipo,'')) LIKE '%DOCENTE%' THEN 0
            WHEN upper(COALESCE(cargo,'')) LIKE '%AGENTE%' THEN 1
            WHEN upper(COALESCE(cargo,'')) LIKE '%DOCENTE%' THEN 2
            ELSE 3
        END, nombre
        """,
        params,
    )
    for d in rows:
        posibles = {normalize_unidad(d.get('unidad'))}
        try:
            if d.get('unidades'):
                posibles.update(normalize_unidad(x) for x in json.loads(d.get('unidades')))
        except Exception:
            pass
        if unidad_norm and unidad_norm in posibles:
            return d
    return {}


def allowed_units_for_user(database_path: str, user: dict[str, Any]) -> set[str] | None:
    rol = user.get('rol') or ''
    if rol in {'SUPERADMIN', 'GERENTE', 'NUTRICIONISTA', 'AUXILIAR_ADMINISTRATIVO'}:
        return None
    unidades = set()
    raw = user.get('unidades') or ''
    try:
        if raw and raw.strip().startswith('['):
            unidades.update(normalize_unidad(x) for x in json.loads(raw))
        elif raw:
            unidades.update(normalize_unidad(x) for x in raw.split(',') if x.strip())
    except Exception:
        pass

    from modules.sqlalchemy_compat import CoreCompatRepository
    repo = CoreCompatRepository()
    if repo.table_exists('coordinadores'):
        rows = repo.fetch_all(
            """
            SELECT unidad, unidades FROM coordinadores
            WHERE activo = 1 AND (documento = ? OR lower(nombre) = lower(?) OR lower(nombre) LIKE lower(?))
            """,
            (user.get('username'), user.get('nombre_completo') or '', f"%{user.get('nombre_completo') or user.get('username') or ''}%"),
        )
        for r in rows:
            unidades.add(normalize_unidad(r.get('unidad')))
            try:
                unidades.update(normalize_unidad(x) for x in json.loads(r.get('unidades') or '[]'))
            except Exception:
                pass
    return {u for u in unidades if u}
