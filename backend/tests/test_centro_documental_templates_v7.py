from __future__ import annotations

import hashlib
import os
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document
from openpyxl import Workbook

from migrations.migrate_centro_documental_v7 import migrate
from modules.centro_documental.repository import CentroDocumentalRepository
from modules.centro_documental.template_inspector_service import inspect_template, propose_mapping


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def run() -> None:
    with tempfile.TemporaryDirectory() as folder:
        root=Path(folder); database=root/"documents.sqlite"; migrate(str(database)); repo=CentroDocumentalRepository(str(database))
        docx=root/"acta.docx"; document=Document(); document.sections[0].header.paragraphs[0].text="MEMBRETE OFICIAL"; table=document.add_table(rows=2,cols=2); table.cell(0,0).text="Tema"; table.cell(0,1).text="{{ actividad.tema }}"; document.save(docx)
        before=digest(docx); doc_inspection=inspect_template(docx); after=digest(docx)
        assert before == after and doc_inspection["tipo"] == "DOCX" and doc_inspection["tablas"][0]["columnas"] == 2
        assert propose_mapping(doc_inspection)["requiere_aprobacion"] is True

        xlsx=root/"capture.xlsx"; book=Workbook(); sheet=book.active; sheet.title="CAPTURE"; sheet.merge_cells("A1:D1"); sheet["A1"]="FORMATO OFICIAL CAPTURE"; sheet.print_area="A1:D20"; sheet["A3"]="=1+1"; book.save(xlsx)
        excel_inspection=inspect_template(xlsx)
        assert excel_inspection["tipo"] == "XLSX" and excel_inspection["hojas"][0]["formulas"] == 1

        version=repo.create_template_version(
            {"codigo":"ACTA_HOGAR","nombre":"Acta hogar","componente":"PEDAGOGICO","tipo_documento":"ACTA_HOGAR","scope":"FUNDACION","fundacion_id":1},
            {"version":"1","nombre_original":docx.name,"nombre_seguro":docx.name,"ruta_privada":str(docx),"mime_type":"application/vnd.openxmlformats-officedocument.wordprocessingml.document","extension":".docx","hash_sha256":before,"inspeccion":doc_inspection},1)
        mapping=repo.save_mapping(version["id"],1,propose_mapping(doc_inspection),1); assert mapping["estado"] == "PROPUESTO"
        approved=repo.approve_mapping(version["id"],1,1); assert approved["estado"] == "APROBADO"
        assert repo.get_version(version["id"],2) is None
        assert len(repo.list_templates(1)) == 1 and len(repo.list_templates(2)) == 0
        try:
            repo.create_template_version(
                {"codigo":"OTRA","nombre":"Otra","componente":"PEDAGOGICO","tipo_documento":"ACTA","scope":"FUNDACION","fundacion_id":1},
                {"version":"1","nombre_original":docx.name,"nombre_seguro":docx.name,"ruta_privada":str(docx),"extension":".docx","hash_sha256":before,"inspeccion":doc_inspection},1)
            raise AssertionError("El hash repetido debe bloquearse")
        except ValueError:
            pass
    print("PASS test_centro_documental_templates_v7")


if __name__ == "__main__": run()
