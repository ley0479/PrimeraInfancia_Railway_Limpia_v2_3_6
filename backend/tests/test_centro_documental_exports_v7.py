from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch
import zipfile

ROOT=Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path: sys.path.insert(0,str(ROOT))

from docx import Document
from modules.centro_documental.document_builder_service import build_docx,build_package,convert_pdf,sha256_file


def run():
    with tempfile.TemporaryDirectory() as folder:
        root=Path(folder); template=root/"oficial.docx"; output=root/"generado.docx"
        document=Document(); document.sections[0].header.paragraphs[0].text="MEMBRETE OFICIAL"; document.add_heading("ACTA INSTITUCIONAL",0); document.add_paragraph("Tema: {{ actividad.tema }}"); document.add_paragraph("{{ documento.narrativa }}"); table=document.add_table(rows=1,cols=2); table.cell(0,0).text="Firma profesional"; table.cell(0,1).text=""; document.save(template)
        before=sha256_file(template); long_text=" ".join(["Hecho confirmado por el profesional."]*300)
        result=build_docx(template,output,{"tema":"Vínculos afectivos","mapped_fields":["documento.narrativa"]},long_text)
        assert result["original_intacto"] and sha256_file(template)==before
        generated=Document(output); all_text="\n".join(p.text for p in generated.paragraphs)
        assert "Vínculos afectivos" in all_text and long_text in all_text
        assert generated.sections[0].header.paragraphs[0].text=="MEMBRETE OFICIAL"
        assert generated.tables[0].cell(0,1).text==""
        pdf_status="PASS"
        try:
            pdf=convert_pdf(output,root/"pdf"); assert pdf.read_bytes()[:5]==b"%PDF-"
        except RuntimeError as exc:
            assert "Word permanece disponible" in str(exc); pdf_status="PENDING"
        simulated_dir=root/"pdf-simulado"
        def fake_run(arguments,**kwargs):
            assert "--headless" in arguments
            assert any(str(value).startswith("-env:UserInstallation=file:") for value in arguments)
            assert kwargs.get("timeout")==90 and kwargs.get("check") is False
            simulated_dir.mkdir(parents=True,exist_ok=True)
            (simulated_dir/output.with_suffix(".pdf").name).write_bytes(b"%PDF-1.7\nSIMULADO")
            return SimpleNamespace(returncode=0,stdout="",stderr="")
        with patch("modules.centro_documental.document_builder_service.shutil.which",return_value="/usr/bin/soffice"), patch("modules.centro_documental.document_builder_service.subprocess.run",side_effect=fake_run):
            simulated_pdf=convert_pdf(output,simulated_dir)
        assert simulated_pdf.read_bytes().startswith(b"%PDF-")
        dockerfile=(ROOT.parent/"Dockerfile").read_text(encoding="utf-8")
        assert "libreoffice-writer" in dockerfile
        package=root/"paquete.zip"; packaged=build_package(package,[("01_ACTA.docx",output)],{"documento_id":1,"capture":"PENDIENTE_PLANTILLA"})
        assert packaged["archivos"] and zipfile.is_zipfile(package)
        with zipfile.ZipFile(package) as archive:
            manifest=json.loads(archive.read("08_MANIFIESTO.json")); assert manifest["capture"]=="PENDIENTE_PLANTILLA" and "01_ACTA.docx" in archive.namelist()
    print(f"PASS test_centro_documental_exports_v7 PDF={pdf_status}")


if __name__=="__main__": run()
