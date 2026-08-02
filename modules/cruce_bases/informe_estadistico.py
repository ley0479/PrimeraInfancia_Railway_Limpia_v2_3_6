from __future__ import annotations

import json
import math
import os
import re
import shutil
import sqlite3
import subprocess
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd

try:
    from .services import normalize_unidad
except Exception:  # pragma: no cover - compatibilidad fuera del paquete
    def normalize_unidad(valor: Any) -> str:
        texto = str(valor or '').strip().upper()
        texto = re.sub(r'\s+', ' ', texto)
        if texto.startswith('UCA '):
            texto = texto[4:].strip()
        return texto


COLOR_PRIMARY = '1F4E78'
COLOR_SECONDARY = '305496'
COLOR_ACCENT = '00A6D6'
COLOR_LIGHT = 'EAF2F8'
COLOR_WARNING = 'F59E0B'
COLOR_DANGER = 'DC2626'
MAX_ANEXO_ROWS = 1500
LOGO_ALLOWED_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.webp'}
DEFAULT_CORPORACION = 'Fundación Principal'


DOCUMENTAL_ALIASES = {
    'sin_carne_salud': [
        'carne_salud', 'carnet_salud', 'carné_salud', 'carne de salud', 'carnet de salud',
        'carné de salud', 'tiene carne de salud', 'tiene carnet de salud', 'salud carnet'
    ],
    'sin_control_crecimiento': [
        'control_crecimiento_desarrollo', 'control crecimiento desarrollo', 'crecimiento y desarrollo',
        'control cyd', 'control c y d', 'control de crecimiento', 'asiste control crecimiento'
    ],
    'sin_vacunas': [
        'vacunas', 'esquema_vacunacion', 'esquema vacunacion', 'carnet vacunas',
        'vacunacion', 'vacunación', 'tiene vacunas', 'vacunas al dia', 'vacunas al día'
    ],
    'sin_carne_crecimiento': [
        'carne_crecimiento_desarrollo', 'carnet_crecimiento_desarrollo', 'carné_crecimiento_desarrollo',
        'carne crecimiento desarrollo', 'carnet crecimiento desarrollo', 'carné crecimiento desarrollo',
        'carnet cyd', 'carne cyd'
    ],
}


def _safe_text(value: Any, default: str = '') -> str:
    if value is None:
        return default
    try:
        if pd.isna(value):
            return default
    except Exception:
        pass
    text = str(value).strip()
    if text.lower() in {'nan', 'nat', 'none', 'null'}:
        return default
    return text


def _norm_text(value: Any) -> str:
    import unicodedata

    text = _safe_text(value).lower()
    text = unicodedata.normalize('NFKD', text)
    text = ''.join(ch for ch in text if not unicodedata.combining(ch))
    text = re.sub(r'[^a-z0-9]+', ' ', text)
    return ' '.join(text.split())


def _doc_key(value: Any) -> str:
    text = _safe_text(value)
    if text.endswith('.0'):
        text = text[:-2]
    return re.sub(r'\s+', '', text)


def _documento_no_valido(value: Any) -> bool:
    doc = _doc_key(value)
    norm = _norm_text(doc)
    if not doc:
        return True
    etiquetas = {'registro civil', 'tarjeta identidad', 'cedula', 'cedula ciudadania', 'nui', 'sin documento', 'pendiente'}
    if norm in etiquetas:
        return True
    # Evita tomar como documento textos del tipo de documento. La mayoría de documentos válidos tienen dígitos.
    if not re.search(r'\d', doc):
        return True
    return len(re.sub(r'[^0-9A-Za-z]', '', doc)) < 5


def _upper(value: Any) -> str:
    return _safe_text(value).upper()


def _connect(database_path: str) -> sqlite3.Connection:
    conn = sqlite3.connect(database_path)
    conn.row_factory = sqlite3.Row
    return conn


def _table_exists(conn: sqlite3.Connection, table: str) -> bool:
    row = conn.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (table,)).fetchone()
    return bool(row)


def _columns(conn: sqlite3.Connection, table: str) -> set[str]:
    if not _table_exists(conn, table):
        return set()
    return {r['name'] for r in conn.execute(f"PRAGMA table_info({table})").fetchall()}


def _read_table(conn: sqlite3.Connection, table: str, fundacion_id: int | None, superadmin: bool = False) -> pd.DataFrame:
    if not _table_exists(conn, table):
        return pd.DataFrame()
    cols = _columns(conn, table)
    sql = f"SELECT * FROM {table}"
    params: list[Any] = []
    if fundacion_id and not superadmin and 'fundacion_id' in cols:
        sql += " WHERE COALESCE(fundacion_id, 1) = ?"
        params.append(fundacion_id)
    try:
        return pd.read_sql_query(sql, conn, params=params)
    except Exception:
        rows = [dict(r) for r in conn.execute(sql, params).fetchall()]
        return pd.DataFrame(rows)



def _backend_dir(database_path: str) -> Path:
    return Path(database_path).resolve().parent


def _project_dir_from_database(database_path: str) -> Path:
    return _backend_dir(database_path).parent


def _static_dir(database_path: str) -> Path:
    static = _backend_dir(database_path) / 'static'
    (static / 'uploads' / 'logos').mkdir(parents=True, exist_ok=True)
    (static / 'img').mkdir(parents=True, exist_ok=True)
    return static


def ensure_logo_schema(database_path: str) -> None:
    """Migración segura ALPHA32: agrega logo_path sin tocar datos existentes."""
    conn = _connect(database_path)
    try:
        for table in ['corporaciones', 'fundaciones']:
            if _table_exists(conn, table) and 'logo_path' not in _columns(conn, table):
                conn.execute(f"ALTER TABLE {table} ADD COLUMN logo_path TEXT")
        conn.commit()
    finally:
        conn.close()


def _resolve_existing_path(database_path: str, raw_path: Any) -> str:
    value = _safe_text(raw_path)
    if not value:
        return ''
    candidates = []
    p = Path(value)
    if p.is_absolute():
        candidates.append(p)
    else:
        backend_dir = _backend_dir(database_path)
        project_dir = _project_dir_from_database(database_path)
        candidates.extend([backend_dir / value, project_dir / value, Path.cwd() / value])
    for candidate in candidates:
        try:
            if candidate.exists() and candidate.is_file():
                return str(candidate)
        except Exception:
            continue
    return ''


def _crear_logo_default(database_path: str, output_folder: str | None = None) -> str:
    """Devuelve un logo default estable en backend/static/img, o crea uno si hace falta."""
    static = _static_dir(database_path)
    default_logo = static / 'img' / 'logo_default.png'
    if default_logo.exists():
        return str(default_logo)

    template_logo = _backend_dir(database_path) / 'templates_originales' / 'informe_estadistico' / 'logo_primera_infancia.png'
    if template_logo.exists():
        try:
            shutil.copyfile(template_logo, default_logo)
            return str(default_logo)
        except Exception:
            return str(template_logo)

    # Respaldo adicional para ambientes donde no exista backend/static todavía.
    fallback_dir = Path(output_folder or _backend_dir(database_path)) / 'assets_informe_estadistico'
    fallback_dir.mkdir(parents=True, exist_ok=True)
    fallback_logo = fallback_dir / 'logo_primera_infancia.png'
    try:
        from PIL import Image, ImageDraw, ImageFont

        img = Image.new('RGB', (760, 220), 'white')
        draw = ImageDraw.Draw(img)
        draw.rounded_rectangle((28, 28, 172, 172), radius=32, fill=(31, 78, 120))
        draw.text((74, 78), 'PI', fill='white', font=ImageFont.load_default())
        draw.text((205, 70), 'PrimeraInfancia', fill=(31, 78, 120), font=ImageFont.load_default())
        draw.text((205, 110), 'Gestión institucional y seguimiento integral', fill=(80, 80, 80), font=ImageFont.load_default())
        img.save(default_logo)
        return str(default_logo)
    except Exception:
        return str(fallback_logo) if fallback_logo.exists() else ''


def obtener_corporacion_contexto(database_path: str, fundacion_id: int | None, corporacion_id: int | None = None) -> dict[str, Any]:
    """Obtiene datos institucionales para portada, encabezado y pie de página."""
    ensure_logo_schema(database_path)
    fundacion_id = int(fundacion_id or 1)
    conn = _connect(database_path)
    try:
        row = None
        if _table_exists(conn, 'corporaciones'):
            cols = _columns(conn, 'corporaciones')
            select_cols = ', '.join([c for c in ['id', 'fundacion_id', 'nombre', 'nit', 'representante', 'logo_path'] if c in cols]) or '*'
            if corporacion_id:
                row = conn.execute(f"SELECT {select_cols} FROM corporaciones WHERE id = ? LIMIT 1", (corporacion_id,)).fetchone()
            if row is None and 'fundacion_id' in cols:
                row = conn.execute(f"SELECT {select_cols} FROM corporaciones WHERE fundacion_id = ? LIMIT 1", (fundacion_id,)).fetchone()
        if row is None and _table_exists(conn, 'fundaciones'):
            cols = _columns(conn, 'fundaciones')
            select_cols = ', '.join([c for c in ['id', 'nombre', 'nit', 'representante', 'logo_path'] if c in cols]) or '*'
            row = conn.execute(f"SELECT {select_cols} FROM fundaciones WHERE id = ? LIMIT 1", (fundacion_id,)).fetchone()
        data = dict(row) if row else {}
        nombre = _safe_text(data.get('nombre'), DEFAULT_CORPORACION)
        return {
            'id': int(data.get('id') or corporacion_id or fundacion_id or 1),
            'fundacion_id': fundacion_id,
            'nombre': nombre,
            'nit': _safe_text(data.get('nit'), 'NIT no registrado'),
            'representante': _safe_text(data.get('representante'), ''),
            'logo_path': _safe_text(data.get('logo_path'), ''),
        }
    finally:
        conn.close()


def obtener_logo_corporacion(database_path: str, fundacion_id: int | None = None, corporacion_id: int | None = None, output_folder: str | None = None) -> str:
    corp = obtener_corporacion_contexto(database_path, fundacion_id, corporacion_id)
    logo = _resolve_existing_path(database_path, corp.get('logo_path'))
    if logo:
        return logo
    return _crear_logo_default(database_path, output_folder)


def registrar_logo_corporacion(database_path: str, file_storage: Any, upload_folder: str | None, ctx: dict[str, Any] | None = None) -> dict[str, Any]:
    """Guarda el logo de corporación/fundación y actualiza logo_path con migración segura."""
    if file_storage is None or not getattr(file_storage, 'filename', ''):
        raise ValueError('Debes seleccionar un archivo de logo.')
    filename = secure_filename(file_storage.filename) if 'secure_filename' in globals() else re.sub(r'[^A-Za-z0-9_.-]+', '_', file_storage.filename)
    ext = Path(filename).suffix.lower()
    if ext not in LOGO_ALLOWED_EXTENSIONS:
        raise ValueError('Formato de logo no permitido. Usa PNG, JPG, JPEG o WEBP.')
    ctx = ctx or {}
    fundacion_id = int(ctx.get('fundacion_id') or 1)
    corporacion_id = ctx.get('corporacion_id')
    static = _static_dir(database_path)
    logo_dir = static / 'uploads' / 'logos'
    safe_name = f"corporacion_{corporacion_id or fundacion_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:6]}{ext}"
    out_path = logo_dir / safe_name
    file_storage.save(str(out_path))

    # Valida que realmente sea imagen legible. Si Pillow no está disponible, se conserva el archivo.
    try:
        from PIL import Image
        with Image.open(out_path) as img:
            img.verify()
    except Exception as exc:
        try:
            out_path.unlink(missing_ok=True)
        except Exception:
            pass
        raise ValueError(f'El archivo cargado no parece un logo válido: {exc}')

    relative_path = str(Path('static') / 'uploads' / 'logos' / safe_name)
    ensure_logo_schema(database_path)
    conn = _connect(database_path)
    try:
        updated = False
        if _table_exists(conn, 'corporaciones'):
            if corporacion_id:
                cur = conn.execute("UPDATE corporaciones SET logo_path = ?, fecha_actualizacion = ? WHERE id = ?", (relative_path, datetime.now().isoformat(timespec='seconds'), corporacion_id))
                updated = cur.rowcount > 0
            if not updated:
                cur = conn.execute("UPDATE corporaciones SET logo_path = ?, fecha_actualizacion = ? WHERE fundacion_id = ?", (relative_path, datetime.now().isoformat(timespec='seconds'), fundacion_id))
                updated = cur.rowcount > 0
            if not updated:
                conn.execute(
                    "INSERT INTO corporaciones (fundacion_id, nombre, estado, fecha_creacion, fecha_actualizacion, logo_path) VALUES (?, ?, 'ACTIVA', ?, ?, ?)",
                    (fundacion_id, DEFAULT_CORPORACION, datetime.now().isoformat(timespec='seconds'), datetime.now().isoformat(timespec='seconds'), relative_path)
                )
                updated = True
        if not updated and _table_exists(conn, 'fundaciones'):
            conn.execute("UPDATE fundaciones SET logo_path = ?, fecha_actualizacion = ? WHERE id = ?", (relative_path, datetime.now().isoformat(timespec='seconds'), fundacion_id))
        conn.commit()
    finally:
        conn.close()
    return {'message': 'Logo institucional actualizado correctamente.', 'logo_path': relative_path, 'filename': safe_name}

def _to_numeric(series: pd.Series | Any) -> pd.Series:
    if isinstance(series, pd.Series):
        return pd.to_numeric(series, errors='coerce')
    return pd.Series([], dtype=float)


def _fecha_orden(df: pd.DataFrame) -> pd.Series:
    candidatos = ['fecha_actualizacion', 'fecha_carga', 'fecha_creacion', 'fecha_ingreso', 'fecha_valoracion', 'fecha_toma', 'fecha_medicion']
    out = pd.Series(pd.NaT, index=df.index)
    for col in candidatos:
        if col in df.columns:
            parsed = pd.to_datetime(df[col], errors='coerce')
            out = out.fillna(parsed)
    return out


def _edad_meses(row: pd.Series) -> int:
    for col in ['edad_meses', 'edad_mes']:
        if col in row and _safe_text(row.get(col)):
            try:
                return int(float(row.get(col)))
            except Exception:
                pass
    fecha = _safe_text(row.get('fecha_nacimiento'))
    if fecha:
        try:
            f = pd.to_datetime(fecha, errors='coerce')
            if pd.notna(f):
                hoy = pd.Timestamp.now()
                meses = (hoy.year - f.year) * 12 + (hoy.month - f.month)
                if hoy.day < f.day:
                    meses -= 1
                return max(0, int(meses))
        except Exception:
            pass
    return 0


def _grupo_etario(row: pd.Series) -> str:
    raw = _safe_text(row.get('grupo_edad') or row.get('grupo') or row.get('tipo_beneficiario'))
    if raw:
        return _upper(raw)
    meses = _edad_meses(row)
    tipo = _norm_text(row.get('tipo_beneficiario'))
    if 'gestante' in tipo:
        return '0 A 6 MESES Y GESTANTES'
    if meses <= 6:
        return '0 A 6 MESES Y GESTANTES'
    if meses <= 11:
        return '6 A 11 MESES 29 DIAS'
    if meses <= 35:
        return '1 A 2 ANOS 11 MESES'
    if meses <= 71:
        return '3 A 5 ANOS 11 MESES'
    return 'MAYOR DE 5 ANOS / SIN CLASIFICAR'


def _nombre_completo(row: pd.Series) -> str:
    direct = _safe_text(row.get('nombre') or row.get('nombre_completo'))
    if direct:
        return direct.upper()
    nombres = _safe_text(row.get('nombres') or ' '.join(_safe_text(row.get(c)) for c in ['primer_nombre', 'segundo_nombre']).strip())
    apellidos = _safe_text(row.get('apellidos') or ' '.join(_safe_text(row.get(c)) for c in ['primer_apellido', 'segundo_apellido']).strip())
    return ' '.join([nombres, apellidos]).strip().upper()


def _version_activa_id(conn: sqlite3.Connection, fundacion_id: int | None, superadmin: bool = False) -> int | None:
    if not (_table_exists(conn, 'master_versiones') and _table_exists(conn, 'master_ninos')):
        return None
    params: list[Any] = []
    sql = "SELECT id FROM master_versiones WHERE activa = 1"
    if fundacion_id and not superadmin:
        sql += " AND COALESCE(fundacion_id, 1) = ?"
        params.append(fundacion_id)
    elif fundacion_id:
        sql += " AND COALESCE(fundacion_id, 1) = ?"
        params.append(fundacion_id)
    sql += " ORDER BY fecha_publicacion DESC, id DESC LIMIT 1"
    row = conn.execute(sql, params).fetchone()
    return int(row['id']) if row else None


def _preparar_df_maestro(df: pd.DataFrame, fuente: str) -> tuple[pd.DataFrame, list[dict[str, Any]], str]:
    if df.empty:
        base_cols = ['documento', 'nombre_completo', 'unidad', 'estado', 'sexo', 'grupo_etario', 'coordinador']
        return pd.DataFrame(columns=base_cols), [], fuente
    df = df.copy()
    if 'documento' not in df.columns and 'nui' in df.columns:
        df['documento'] = df['nui']
    if 'unidad' not in df.columns:
        if 'unidad_servicio' in df.columns:
            df['unidad'] = df['unidad_servicio']
        elif 'nombre_unidad' in df.columns:
            df['unidad'] = df['nombre_unidad']
    if 'grupo_edad' not in df.columns and 'grupo_etario' in df.columns:
        df['grupo_edad'] = df['grupo_etario']
    if 'peso_kg' not in df.columns and 'peso' in df.columns:
        df['peso_kg'] = df['peso']
    if 'talla_cm' not in df.columns and 'talla' in df.columns:
        df['talla_cm'] = df['talla']
    if 'perimetro_braquial_cm' not in df.columns and 'perimetro_braquial' in df.columns:
        df['perimetro_braquial_cm'] = df['perimetro_braquial']

    df['documento'] = df.get('documento', pd.Series([''] * len(df), index=df.index)).apply(_doc_key)
    # Se conserva la regla histórica: si no hay documento, no se cuenta como niño maestro.
    df = df[df['documento'].astype(str).str.len() > 0].copy()
    if df.empty:
        return pd.DataFrame(columns=['documento', 'nombre_completo', 'unidad', 'estado', 'sexo', 'grupo_etario', 'coordinador']), [], fuente
    df['nombre_completo'] = df.apply(_nombre_completo, axis=1)
    df['unidad'] = df.get('unidad', pd.Series([''] * len(df), index=df.index)).apply(normalize_unidad)
    df['estado'] = df.get('estado', pd.Series(['ACTIVO'] * len(df), index=df.index)).apply(lambda x: _upper(x) or 'ACTIVO')
    df['sexo'] = df.get('sexo', pd.Series(['SIN DATO'] * len(df), index=df.index)).apply(lambda x: _upper(x) or 'SIN DATO')
    df['grupo_etario'] = df.apply(_grupo_etario, axis=1)
    df['_fecha_orden'] = _fecha_orden(df)
    if 'id' in df.columns:
        df['_id_orden'] = pd.to_numeric(df['id'], errors='coerce').fillna(0)
    else:
        df['_id_orden'] = range(1, len(df) + 1)

    duplicated_mask = df.duplicated('documento', keep=False)
    duplicados: list[dict[str, Any]] = []
    if duplicated_mask.any():
        dup_df = df[duplicated_mask].sort_values(['documento', '_fecha_orden', '_id_orden'])
        for _, r in dup_df.iterrows():
            duplicados.append({
                'documento': r.get('documento', ''),
                'nombre': r.get('nombre_completo', ''),
                'nombre_completo': r.get('nombre_completo', ''),
                'unidad': r.get('unidad', ''),
                'estado': r.get('estado', ''),
                'fecha_carga': _safe_text(r.get('fecha_carga') or r.get('fecha_actualizacion') or r.get('fecha_creacion')),
                'origen': fuente,
                'observacion': 'Documento repetido en fuente maestra; se conservó el registro más reciente.',
            })

    df = df.sort_values(['documento', '_fecha_orden', '_id_orden']).drop_duplicates('documento', keep='last')
    df['_fuente_maestra'] = fuente
    return df.reset_index(drop=True), duplicados, fuente


def _cargar_base_maestra(conn: sqlite3.Connection, fundacion_id: int | None, superadmin: bool = False) -> tuple[pd.DataFrame, list[dict[str, Any]], str]:
    """Carga oficial del informe.

    ALPHA32: primero intenta leer la Base Maestra publicada (master_ninos +
    master_versiones). Solo si no existe una publicación activa se usa el
    comportamiento histórico como respaldo, para no dejar sin informe a una
    instalación que aún no haya publicado Base Maestra.
    """
    version_id = _version_activa_id(conn, fundacion_id, superadmin)
    if version_id:
        cols = _columns(conn, 'master_ninos')
        sql = "SELECT * FROM master_ninos WHERE version_id = ? AND COALESCE(activo, 0) = 1"
        params: list[Any] = [version_id]
        if fundacion_id and not superadmin and 'fundacion_id' in cols:
            sql += " AND COALESCE(fundacion_id, 1) = ?"
            params.append(fundacion_id)
        try:
            df_master = pd.read_sql_query(sql, conn, params=params)
        except Exception:
            df_master = pd.DataFrame([dict(r) for r in conn.execute(sql, params).fetchall()])
        if not df_master.empty:
            return _preparar_df_maestro(df_master, 'master_ninos')

    # Respaldo histórico: beneficiarios / usuarios, solo si no hay Base Maestra publicada.
    fuente = 'beneficiarios'
    df = _read_table(conn, 'beneficiarios', fundacion_id, superadmin)
    if df.empty:
        fuente = 'usuarios'
        df = _read_table(conn, 'usuarios', fundacion_id, superadmin)
    return _preparar_df_maestro(df, fuente)

def _mapa_coordinadores(conn: sqlite3.Connection, fundacion_id: int | None, superadmin: bool) -> dict[str, str]:
    mapa: dict[str, str] = {}

    for table in ['unidades', 'th_asignaciones', 'coordinadores', 'gp_unidades_asignadas']:
        df = _read_table(conn, table, fundacion_id, superadmin)
        if df.empty:
            continue
        if table == 'unidades':
            for _, r in df.iterrows():
                unidad = normalize_unidad(r.get('nombre') or r.get('unidad'))
                coord = _safe_text(r.get('coordinador_nombre') or r.get('coordinador'))
                if unidad and coord:
                    mapa.setdefault(unidad, coord.upper())
        elif table == 'th_asignaciones':
            for _, r in df.iterrows():
                unidad = normalize_unidad(r.get('unidad'))
                coord = _safe_text(r.get('coordinador_nombre'))
                if unidad and coord:
                    mapa.setdefault(unidad, coord.upper())
        elif table == 'coordinadores':
            for _, r in df.iterrows():
                unidad = normalize_unidad(r.get('unidad'))
                coord = _safe_text(r.get('coordinador'))
                cargo = _norm_text(r.get('cargo') or r.get('tipo_equipo') or r.get('perfil'))
                nombre = _safe_text(r.get('nombre') or ' '.join([_safe_text(r.get('nombres')), _safe_text(r.get('apellidos'))]).strip())
                if not coord and ('coordinador' in cargo or 'coord' in cargo):
                    coord = nombre
                if unidad and coord:
                    mapa.setdefault(unidad, coord.upper())
                if r.get('unidades') and (coord or ('coordinador' in cargo and nombre)):
                    final_coord = (coord or nombre).upper()
                    try:
                        unidades = json.loads(r.get('unidades') or '[]')
                        if isinstance(unidades, str):
                            unidades = [unidades]
                        for u in unidades:
                            unidad_json = normalize_unidad(u)
                            if unidad_json:
                                mapa.setdefault(unidad_json, final_coord)
                    except Exception:
                        pass
        elif table == 'gp_unidades_asignadas':
            for _, r in df.iterrows():
                unidad = normalize_unidad(r.get('unidad') or r.get('nombre_unidad'))
                coord = _safe_text(r.get('coordinador_nombre') or r.get('coordinador'))
                if unidad and coord:
                    mapa.setdefault(unidad, coord.upper())
    return mapa


def _enriquecer_coordinador(df: pd.DataFrame, conn: sqlite3.Connection, fundacion_id: int | None, superadmin: bool) -> tuple[pd.DataFrame, dict[str, str]]:
    mapa = _mapa_coordinadores(conn, fundacion_id, superadmin)
    if df.empty:
        df['coordinador'] = []
        return df, mapa
    df = df.copy()
    if 'coordinador' in df.columns:
        base_coord = df['coordinador'].apply(lambda x: _upper(x) or '')
    else:
        base_coord = pd.Series([''] * len(df), index=df.index)
    df['coordinador'] = [base_coord.iloc[i] or mapa.get(normalize_unidad(unidad), 'SIN COORDINADOR') for i, unidad in enumerate(df['unidad'])]
    return df, mapa


def _latest_by_document(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    df = df.copy()
    if 'documento' not in df.columns and 'nui' in df.columns:
        df['documento'] = df['nui']
    df['documento'] = df.get('documento', pd.Series([''] * len(df), index=df.index)).apply(_doc_key)
    df = df[df['documento'].astype(str).str.len() > 0].copy()
    if df.empty:
        return df
    df['_fecha_orden'] = _fecha_orden(df)
    df['_id_orden'] = pd.to_numeric(df['id'], errors='coerce').fillna(0) if 'id' in df.columns else range(1, len(df) + 1)
    return df.sort_values(['documento', '_fecha_orden', '_id_orden']).drop_duplicates('documento', keep='last')


def _valor_faltante(value: Any) -> bool:
    text = _norm_text(value)
    if not text:
        return True
    if text in {'no', 'n', 'na', 'n a', 'sin dato', 'sin soporte', 'pendiente', 'no registra', 'no aplica', 'ninguno', 'null'}:
        return True
    if text.startswith('no '):
        return True
    return False


def _valor_afirmativo(value: Any) -> bool:
    text = _norm_text(value)
    return text in {'si', 's', 'yes', 'true', '1', 'completo', 'al dia', 'al dia vigente', 'vigente', 'cumple'}


def _buscar_columna_alias(df: pd.DataFrame, aliases: list[str]) -> str | None:
    if df.empty:
        return None
    normalized = [(_norm_text(c), c) for c in df.columns]
    alias_norm = [_norm_text(a) for a in aliases]
    for alias in alias_norm:
        for col_norm, original in normalized:
            if col_norm == alias:
                return original
    for alias in alias_norm:
        for col_norm, original in normalized:
            if alias and alias in col_norm:
                return original
    return None


def _faltante_por_alias(df: pd.DataFrame, aliases: list[str], notas: list[str], etiqueta: str) -> pd.Series:
    col = _buscar_columna_alias(df, aliases)
    if not col:
        notas.append(f'No se encontró columna para {etiqueta}; el indicador se deja en 0 y no se infiere para evitar conteos falsos.')
        return pd.Series([False] * len(df), index=df.index)
    return df[col].apply(lambda x: (not _valor_afirmativo(x)) and _valor_faltante(x))


def _contiene(series: pd.Series, *tokens: str) -> pd.Series:
    normalized = series.fillna('').astype(str).apply(_norm_text)
    mask = pd.Series([True] * len(series), index=series.index)
    for token in tokens:
        mask &= normalized.str.contains(_norm_text(token), regex=False)
    return mask


def _enriquecer_salud_nutricion(df: pd.DataFrame, conn: sqlite3.Connection, fundacion_id: int | None, superadmin: bool) -> tuple[pd.DataFrame, list[str], pd.DataFrame]:
    notas: list[str] = []
    if df.empty:
        return df, notas, pd.DataFrame()
    out = df.copy()
    usa_base_maestra_publicada = '_fuente_maestra' in out.columns and out['_fuente_maestra'].astype(str).eq('master_ninos').all()

    if not usa_base_maestra_publicada:
        sn = _latest_by_document(_read_table(conn, 'sn_valoraciones', fundacion_id, superadmin))
        if not sn.empty:
            columnas = ['documento', 'peso_kg', 'talla_cm', 'perimetro_braquial_cm', 'diagnostico_global', 'diag_peso_edad', 'diag_talla_edad', 'diag_peso_talla', 'diag_imc_edad', 'diag_braquial_edad', 'nivel_alerta', 'fecha_valoracion', 'periodo', 'trimestre']
            columnas = [c for c in columnas if c in sn.columns]
            out = out.merge(sn[columnas], on='documento', how='left', suffixes=('', '_sn'))
        else:
            notas.append('No hay registros en sn_valoraciones; se usa peso_talla como respaldo cuando exista.')

        pt = _latest_by_document(_read_table(conn, 'peso_talla', fundacion_id, superadmin))
        if not pt.empty:
            columnas = ['documento', 'peso', 'talla', 'estado_nutricional', 'fecha_toma', 'fecha_medicion']
            columnas = [c for c in columnas if c in pt.columns]
            out = out.merge(pt[columnas], on='documento', how='left', suffixes=('', '_pt'))
    else:
        notas.append('Informe alimentado desde Base Maestra publicada; no se consultaron Excel originales ni fuentes staging.')

    if 'peso_kg' not in out.columns:
        out['peso_kg'] = pd.NA
    if 'talla_cm' not in out.columns:
        out['talla_cm'] = pd.NA
    if 'peso' in out.columns:
        out['peso_kg'] = pd.to_numeric(out['peso_kg'], errors='coerce')
        out['peso'] = pd.to_numeric(out['peso'], errors='coerce')
        out['peso_kg'] = out['peso_kg'].where(out['peso_kg'].notna(), out['peso'])
    if 'talla' in out.columns:
        out['talla_cm'] = pd.to_numeric(out['talla_cm'], errors='coerce')
        out['talla'] = pd.to_numeric(out['talla'], errors='coerce')
        out['talla_cm'] = out['talla_cm'].where(out['talla_cm'].notna(), out['talla'])
    if 'perimetro_braquial_cm' not in out.columns:
        out['perimetro_braquial_cm'] = pd.NA
    if 'perimetro_braquial' in out.columns:
        out['perimetro_braquial_cm'] = pd.to_numeric(out['perimetro_braquial_cm'], errors='coerce')
        out['perimetro_braquial'] = pd.to_numeric(out['perimetro_braquial'], errors='coerce')
        out['perimetro_braquial_cm'] = out['perimetro_braquial_cm'].where(out['perimetro_braquial_cm'].notna(), out['perimetro_braquial'])

    diag_cols = [c for c in ['diagnostico_global', 'diag_peso_edad', 'diag_talla_edad', 'diag_peso_talla', 'diag_imc_edad', 'diag_braquial_edad', 'diagnostico_nutricional', 'estado_nutricional'] if c in out.columns]
    if diag_cols:
        diag = out[diag_cols].apply(lambda r: next((_safe_text(v).upper() for v in r if _safe_text(v)), ''), axis=1)
    else:
        diag = pd.Series([''] * len(out), index=out.index)
        notas.append('No se encontraron columnas de diagnóstico nutricional.')
    out['diagnostico_nutricional'] = diag.replace('', 'SIN DIAGNOSTICO')

    out['sin_peso'] = _to_numeric(out.get('peso_kg')).isna() | (_to_numeric(out.get('peso_kg')).fillna(0) <= 0)
    out['sin_talla'] = _to_numeric(out.get('talla_cm')).isna() | (_to_numeric(out.get('talla_cm')).fillna(0) <= 0)
    out['sin_peso_talla'] = out['sin_peso'] & out['sin_talla']
    if 'perimetro_braquial_cm' in out.columns:
        out['sin_perimetro_braquial'] = _to_numeric(out['perimetro_braquial_cm']).isna() | (_to_numeric(out['perimetro_braquial_cm']).fillna(0) <= 0)
    else:
        out['sin_perimetro_braquial'] = True
        notas.append('No se encontró perímetro braquial en la fuente nutricional; se marca como faltante para seguimiento.')
    out['sin_registro_civil'] = out.get('documento', pd.Series([''] * len(out), index=out.index)).apply(_documento_no_valido)
    out['sin_diagnostico_nutricional'] = out['diagnostico_nutricional'].apply(lambda x: _norm_text(x) in {'', 'pendiente', 'sin diagnostico', 'sin diagnostico nutricional', 'no registra'})

    for key, aliases in DOCUMENTAL_ALIASES.items():
        etiqueta = key.replace('sin_', '').replace('_', ' ')
        out[key] = _faltante_por_alias(out, aliases, notas, etiqueta)

    diag_series = out['diagnostico_nutricional']
    out['bajo_peso'] = _contiene(diag_series, 'bajo peso') & ~_contiene(diag_series, 'riesgo')
    out['riesgo_bajo_peso'] = _contiene(diag_series, 'riesgo', 'bajo peso') | _contiene(diag_series, 'riesgo de desnutricion')
    out['sobrepeso'] = _contiene(diag_series, 'sobrepeso')
    out['obesidad'] = _contiene(diag_series, 'obesidad')
    out['talla_baja'] = _contiene(diag_series, 'talla baja') | _contiene(diag_series, 'retraso talla')

    alertas = _read_table(conn, 'sn_alertas', fundacion_id, superadmin)
    if not alertas.empty:
        alertas = alertas.copy()
        alertas['documento'] = alertas.get('documento', pd.Series([''] * len(alertas), index=alertas.index)).apply(_doc_key)
        if 'atendida' in alertas.columns:
            abiertas = alertas[pd.to_numeric(alertas['atendida'], errors='coerce').fillna(0) == 0]
        else:
            abiertas = alertas
        counts = abiertas.groupby('documento').size().rename('alertas_abiertas') if not abiertas.empty else pd.Series(dtype=int, name='alertas_abiertas')
        out = out.merge(counts.reset_index(), on='documento', how='left')
        out['alertas_abiertas'] = pd.to_numeric(out['alertas_abiertas'], errors='coerce').fillna(0).astype(int)
    else:
        alertas = pd.DataFrame()
        out['alertas_abiertas'] = 0

    nivel = out.get('nivel_alerta', pd.Series([''] * len(out), index=out.index)).fillna('').astype(str).apply(_norm_text)
    out['alertas_nutricionales'] = (
        out['alertas_abiertas'].fillna(0).astype(int).gt(0)
        | out['bajo_peso'] | out['riesgo_bajo_peso'] | out['sobrepeso'] | out['obesidad'] | out['talla_baja']
        | (~nivel.isin({'', 'verde', 'normal', 'sin alerta'}))
    )

    faltante_cols = [
        'sin_carne_salud', 'sin_control_crecimiento', 'sin_vacunas', 'sin_registro_civil', 'sin_carne_crecimiento',
        'sin_peso', 'sin_talla', 'sin_peso_talla', 'sin_perimetro_braquial', 'sin_diagnostico_nutricional'
    ]
    out['datos_faltantes'] = out[faltante_cols].any(axis=1)
    return out, notas, alertas


def _apply_filters(df: pd.DataFrame, filtros: dict[str, Any]) -> pd.DataFrame:
    if df.empty:
        return df
    out = df.copy()
    unidad = normalize_unidad(filtros.get('unidad'))
    if unidad:
        out = out[out['unidad'].apply(normalize_unidad) == unidad]
    coordinador = _norm_text(filtros.get('coordinador'))
    if coordinador:
        out = out[out['coordinador'].apply(_norm_text).str.contains(coordinador, regex=False, na=False)]
    grupo = _norm_text(filtros.get('grupo_etario'))
    if grupo:
        out = out[out['grupo_etario'].apply(_norm_text) == grupo]
    diagnostico = _norm_text(filtros.get('estado_nutricional') or filtros.get('diagnostico'))
    if diagnostico:
        out = out[out['diagnostico_nutricional'].apply(_norm_text).str.contains(diagnostico, regex=False, na=False)]
    if str(filtros.get('alertas') or '').lower() in {'1', 'true', 'si', 'sí'}:
        out = out[out['alertas_nutricionales']]
    if str(filtros.get('faltantes') or '').lower() in {'1', 'true', 'si', 'sí'}:
        out = out[out['datos_faltantes']]
    return out.reset_index(drop=True)


def _item_unidades(item: dict[str, Any]) -> set[str]:
    fields = ['unidad', 'unidad_actual', 'unidad_anterior', 'unidad_nuevo', 'unidad_retirado']
    return {normalize_unidad(item.get(f)) for f in fields if normalize_unidad(item.get(f))}


def _item_documentos(item: dict[str, Any]) -> set[str]:
    fields = ['documento', 'documento_actual', 'documento_anterior', 'documento_nuevo', 'documento_retirado']
    return {_doc_key(item.get(f)) for f in fields if _doc_key(item.get(f))}


def _filtrar_items_movimiento(items: list[dict[str, Any]], docs_scope: set[str], filtros: dict[str, Any], unidad_coord: dict[str, str]) -> list[dict[str, Any]]:
    unidad = normalize_unidad(filtros.get('unidad'))
    coordinador = _norm_text(filtros.get('coordinador'))
    doc_scope_required = bool(filtros.get('_doc_scope_required'))
    if not unidad and not coordinador and not filtros.get('_scope_active'):
        return items
    if doc_scope_required and not docs_scope:
        return []
    filtered = []
    for item in items:
        item_units = _item_unidades(item)
        item_docs = _item_documentos(item)
        ok = True
        if unidad:
            ok = unidad in item_units
        if ok and coordinador:
            coords = {_norm_text(unidad_coord.get(u, '')) for u in item_units}
            ok = any(coordinador in c for c in coords)
        if ok and doc_scope_required:
            ok = bool(item_docs & docs_scope)
        if ok:
            filtered.append(item)
    return filtered


def _movimientos_contexto(resultado: dict[str, Any], df_scope: pd.DataFrame, filtros: dict[str, Any], unidad_coord: dict[str, str]) -> dict[str, list[dict[str, Any]]]:
    docs = set(df_scope['documento'].astype(str)) if not df_scope.empty and 'documento' in df_scope else set()
    scope_active = any(_safe_text(filtros.get(k)) for k in ['unidad', 'coordinador', 'grupo_etario', 'estado_nutricional', 'diagnostico']) or str(filtros.get('alertas') or '').lower() in {'1', 'true', 'si', 'sí'} or str(filtros.get('faltantes') or '').lower() in {'1', 'true', 'si', 'sí'}
    doc_scope_required = any(_safe_text(filtros.get(k)) for k in ['grupo_etario', 'estado_nutricional', 'diagnostico']) or str(filtros.get('alertas') or '').lower() in {'1', 'true', 'si', 'sí'} or str(filtros.get('faltantes') or '').lower() in {'1', 'true', 'si', 'sí'}
    filtros = dict(filtros, _scope_active=scope_active, _doc_scope_required=doc_scope_required)
    tipos = ['nuevos', 'retirados', 'reemplazados', 'trasladados', 'cambios', 'cambios_unidad', 'cambios_docente', 'cambios_acudiente', 'cambios_telefono', 'cambios_direccion']
    return {tipo: _filtrar_items_movimiento(resultado.get(tipo, []) or [], docs, filtros, unidad_coord) for tipo in tipos}


def _conteo_estado_activo(df: pd.DataFrame) -> int:
    if df.empty or 'estado' not in df:
        return 0
    estados = df['estado'].fillna('').astype(str).apply(_norm_text)
    return int(estados.str.contains('activo', na=False).sum())


def _table_counts(df: pd.DataFrame, group_col: str, extra_cols: list[str] | None = None, label: str | None = None, top: int | None = None) -> list[dict[str, Any]]:
    if df.empty or group_col not in df.columns:
        return []
    extra_cols = extra_cols or []
    rows = []
    grouped = df.groupby(df[group_col].fillna('SIN DATO').astype(str).replace('', 'SIN DATO'))
    for name, g in grouped:
        row: dict[str, Any] = {label or group_col: name, 'total': int(len(g))}
        if 'estado' in df.columns:
            row['activos'] = _conteo_estado_activo(g)
        if 'alertas_nutricionales' in df.columns:
            row['alertas'] = int(g['alertas_nutricionales'].sum())
        if 'datos_faltantes' in df.columns:
            row['datos_faltantes'] = int(g['datos_faltantes'].sum())
        for col in extra_cols:
            if col in df.columns:
                row[col] = int(g[col].sum()) if g[col].dtype == bool else int(g[col].notna().sum())
        rows.append(row)
    rows.sort(key=lambda r: r.get('total', 0), reverse=True)
    return rows[:top] if top else rows



def _estado_mask(df: pd.DataFrame, token: str) -> pd.Series:
    if df.empty or 'estado' not in df.columns:
        return pd.Series([False] * len(df), index=df.index)
    return df['estado'].fillna('').astype(str).apply(_norm_text).str.contains(_norm_text(token), regex=False, na=False)


def _movement_counts(items: list[dict[str, Any]], key_candidates: list[str]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for item in items or []:
        value = ''
        for key in key_candidates:
            value = normalize_unidad(item.get(key)) if 'unidad' in key else _safe_text(item.get(key)).upper()
            if value:
                break
        if value:
            counts[value] = counts.get(value, 0) + 1
    return counts


def _tabla_por_unidad(df: pd.DataFrame, movimientos: dict[str, list[dict[str, Any]]] | None = None) -> list[dict[str, Any]]:
    if df.empty or 'unidad' not in df.columns:
        return []
    movimientos = movimientos or {}
    nuevos = _movement_counts(movimientos.get('nuevos', []), ['unidad', 'unidad_actual', 'unidad_nueva'])
    rows = []
    for unidad, g in df.groupby(df['unidad'].fillna('SIN UNIDAD').astype(str).replace('', 'SIN UNIDAD')):
        unidad_norm = normalize_unidad(unidad)
        rows.append({
            'Unidad de servicio': unidad,
            'total': int(len(g)),
            'activos': _conteo_estado_activo(g),
            'retirados': int(_estado_mask(g, 'retirado').sum()),
            'nuevos': int(nuevos.get(unidad_norm, 0)),
            'sin_peso': int(g['sin_peso'].sum()) if 'sin_peso' in g.columns else 0,
            'sin_talla': int(g['sin_talla'].sum()) if 'sin_talla' in g.columns else 0,
            'sin_peso_talla': int(g['sin_peso_talla'].sum()) if 'sin_peso_talla' in g.columns else 0,
            'alertas_nutricionales': int(g['alertas_nutricionales'].sum()) if 'alertas_nutricionales' in g.columns else 0,
            'datos_faltantes': int(g['datos_faltantes'].sum()) if 'datos_faltantes' in g.columns else 0,
        })
    rows.sort(key=lambda r: r.get('total', 0), reverse=True)
    return rows


def _tabla_por_coordinador(df: pd.DataFrame) -> list[dict[str, Any]]:
    if df.empty or 'coordinador' not in df.columns:
        return []
    rows = []
    for coord, g in df.groupby(df['coordinador'].fillna('SIN COORDINADOR').astype(str).replace('', 'SIN COORDINADOR')):
        unidades = int(g['unidad'].nunique()) if 'unidad' in g.columns else 0
        rows.append({
            'Coordinador': coord,
            'unidades_asignadas': unidades,
            'total': int(len(g)),
            'activos': _conteo_estado_activo(g),
            'ninos_con_pendientes': int(g['datos_faltantes'].sum()) if 'datos_faltantes' in g.columns else 0,
            'alertas_prioritarias': int(g['alertas_nutricionales'].sum()) if 'alertas_nutricionales' in g.columns else 0,
        })
    rows.sort(key=lambda r: r.get('total', 0), reverse=True)
    return rows


def _tabla_generica_estadistica(df: pd.DataFrame, group_col: str, label: str) -> list[dict[str, Any]]:
    if df.empty or group_col not in df.columns:
        return []
    rows = []
    for name, g in df.groupby(df[group_col].fillna('SIN DATO').astype(str).replace('', 'SIN DATO')):
        rows.append({
            label: name,
            'total': int(len(g)),
            'activos': _conteo_estado_activo(g),
            'retirados': int(_estado_mask(g, 'retirado').sum()),
            'alertas': int(g['alertas_nutricionales'].sum()) if 'alertas_nutricionales' in g.columns else 0,
            'datos_faltantes': int(g['datos_faltantes'].sum()) if 'datos_faltantes' in g.columns else 0,
        })
    rows.sort(key=lambda r: r.get('total', 0), reverse=True)
    return rows


def _tabla_movimientos_resumen(resumen: dict[str, Any]) -> list[dict[str, Any]]:
    return [
        {'movimiento': 'Nuevos', 'total': int(resumen.get('total_nuevos') or 0)},
        {'movimiento': 'Retirados', 'total': int(resumen.get('total_retirados') or 0)},
        {'movimiento': 'Permanecen', 'total': int(resumen.get('total_permanecen') or 0)},
        {'movimiento': 'Cambiaron de unidad', 'total': int(resumen.get('total_cambios_unidad') or 0)},
    ]


def _series_from_records(records: list[dict[str, Any]], label_key: str, value_key: str = 'total', top: int | None = None) -> pd.Series:
    if not records:
        return pd.Series(dtype=int)
    data = {str(r.get(label_key) or 'SIN DATO'): int(r.get(value_key) or 0) for r in records if int(r.get(value_key) or 0) >= 0}
    s = pd.Series(data, dtype='int64') if data else pd.Series(dtype=int)
    s = s.sort_values(ascending=False)
    return s.head(top) if top else s

def _faltantes_table(df: pd.DataFrame) -> list[dict[str, Any]]:
    mapping = [
        ('Sin carné de salud', 'sin_carne_salud'),
        ('Sin vacunas', 'sin_vacunas'),
        ('Sin control de crecimiento y desarrollo', 'sin_control_crecimiento'),
        ('Sin registro civil / documento válido', 'sin_registro_civil'),
        ('Sin carné de crecimiento y desarrollo', 'sin_carne_crecimiento'),
        ('Sin peso', 'sin_peso'),
        ('Sin talla', 'sin_talla'),
        ('Sin peso y talla', 'sin_peso_talla'),
        ('Sin perímetro braquial', 'sin_perimetro_braquial'),
        ('Sin diagnóstico nutricional', 'sin_diagnostico_nutricional'),
    ]
    rows = []
    for label, col in mapping:
        rows.append({'indicador': label, 'total': int(df[col].sum()) if col in df.columns else 0})
    return rows


def _indicadores_salud(df: pd.DataFrame) -> dict[str, int]:
    mapping = {
        'sin_carne_salud': 'sin_carne_salud',
        'sin_control_crecimiento': 'sin_control_crecimiento',
        'sin_vacunas': 'sin_vacunas',
        'sin_registro_civil': 'sin_registro_civil',
        'sin_carne_crecimiento': 'sin_carne_crecimiento',
        'sin_peso': 'sin_peso',
        'sin_talla': 'sin_talla',
        'sin_peso_talla': 'sin_peso_talla',
        'sin_perimetro_braquial': 'sin_perimetro_braquial',
        'sin_diagnostico_nutricional': 'sin_diagnostico_nutricional',
        'bajo_peso': 'bajo_peso',
        'riesgo_bajo_peso': 'riesgo_bajo_peso',
        'sobrepeso': 'sobrepeso',
        'obesidad': 'obesidad',
        'talla_baja': 'talla_baja',
        'alertas_nutricionales': 'alertas_nutricionales',
    }
    return {key: int(df[col].sum()) if col in df.columns and not df.empty else 0 for key, col in mapping.items()}


def _alertas_prioritarias(alertas: pd.DataFrame, df_scope: pd.DataFrame) -> list[dict[str, Any]]:
    if alertas.empty:
        return []
    docs = set(df_scope['documento'].astype(str)) if not df_scope.empty and 'documento' in df_scope.columns else set()
    tmp = alertas.copy()
    if 'documento' in tmp.columns and docs:
        tmp = tmp[tmp['documento'].astype(str).isin(docs)]
    if tmp.empty:
        return []
    if 'nivel' not in tmp.columns:
        tmp['nivel'] = 'SIN NIVEL'
    if 'tipo' not in tmp.columns:
        tmp['tipo'] = 'ALERTA'
    grouped = tmp.groupby([tmp['nivel'].fillna('SIN NIVEL').astype(str), tmp['tipo'].fillna('ALERTA').astype(str)]).size().reset_index(name='total')
    grouped = grouped.sort_values('total', ascending=False)
    return grouped.to_dict('records')


def _listado(df: pd.DataFrame, mask: pd.Series | None = None, motivo: str = '') -> list[dict[str, Any]]:
    if df.empty:
        return []
    tmp = df[mask].copy() if mask is not None else df.copy()
    if tmp.empty:
        return []
    cols = ['documento', 'nombre_completo', 'unidad', 'coordinador', 'grupo_etario', 'sexo', 'estado', 'diagnostico_nutricional']
    cols = [c for c in cols if c in tmp.columns]
    rows = []
    for _, r in tmp.head(MAX_ANEXO_ROWS).iterrows():
        row = {c: _safe_text(r.get(c)) for c in cols}
        if motivo:
            row['observacion'] = motivo
        rows.append(row)
    return rows


def _listado_movimiento(items: list[dict[str, Any]], tipo: str) -> list[dict[str, Any]]:
    rows = []
    for item in items[:MAX_ANEXO_ROWS]:
        if tipo == 'cambios_unidad' or tipo == 'trasladados':
            rows.append({
                'documento': item.get('documento', ''),
                'nombre_completo': item.get('nombre', ''),
                'unidad_anterior': item.get('unidad_anterior', ''),
                'unidad_actual': item.get('unidad_actual', ''),
                'observacion': 'Cambio de unidad detectado',
            })
        elif tipo == 'retirados':
            rows.append({
                'documento': item.get('documento', ''),
                'nombre_completo': item.get('nombre', ''),
                'unidad': item.get('unidad', ''),
                'observacion': 'Retirado de la base actual',
            })
        elif tipo == 'nuevos':
            rows.append({
                'documento': item.get('documento', ''),
                'nombre_completo': item.get('nombre', ''),
                'unidad': item.get('unidad', ''),
                'observacion': 'Nuevo ingreso detectado',
            })
    return rows


def _crear_logo_si_no_existe(output_folder: str) -> str:
    assets = Path(output_folder) / 'assets_informe_estadistico'
    assets.mkdir(parents=True, exist_ok=True)
    logo = assets / 'logo_primera_infancia.png'
    if logo.exists():
        return str(logo)
    try:
        from PIL import Image, ImageDraw, ImageFont

        img = Image.new('RGB', (640, 180), 'white')
        draw = ImageDraw.Draw(img)
        draw.rounded_rectangle((20, 20, 150, 150), radius=28, fill=(31, 78, 120))
        draw.text((54, 58), 'PI', fill='white', font=ImageFont.load_default())
        draw.text((180, 55), 'PrimeraInfancia', fill=(31, 78, 120), font=ImageFont.load_default())
        draw.text((180, 90), 'Gestión institucional y seguimiento integral', fill=(80, 80, 80), font=ImageFont.load_default())
        img.save(logo)
    except Exception:
        # Pillow está en requirements, pero si no está disponible se continúa sin imagen.
        return ''
    return str(logo)


def _figure_available() -> bool:
    try:
        import matplotlib  # noqa: F401
        return True
    except Exception:
        return False


def _serie_counts(df: pd.DataFrame, col: str, top: int | None = None) -> pd.Series:
    if df.empty or col not in df.columns:
        return pd.Series(dtype=int)
    s = df[col].fillna('SIN DATO').astype(str).replace('', 'SIN DATO').value_counts()
    return s.head(top) if top else s


def _generar_grafica_barras(series: pd.Series, title: str, path: Path, xlabel: str = '', ylabel: str = 'Total') -> str:
    if series.empty:
        return ''
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    labels = [str(x)[:35] for x in series.index]
    values = series.astype(int).tolist()
    fig = plt.figure(figsize=(10, 5.2))
    ax = fig.add_subplot(111)
    ax.bar(labels, values, color=f'#{COLOR_ACCENT}', edgecolor=f'#{COLOR_PRIMARY}', linewidth=0.6)
    ax.set_title(title, fontweight='bold')
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    ax.tick_params(axis='x', rotation=45)
    for i, v in enumerate(values):
        ax.text(i, v, str(v), ha='center', va='bottom', fontsize=8)
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches='tight')
    plt.close(fig)
    return str(path)


def _generar_grafica_pie(series: pd.Series, title: str, path: Path) -> str:
    if series.empty or int(series.sum()) <= 0:
        return ''
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    fig = plt.figure(figsize=(7.5, 5.5))
    ax = fig.add_subplot(111)
    palette = [f'#{COLOR_PRIMARY}', f'#{COLOR_ACCENT}', f'#{COLOR_SECONDARY}', f'#{COLOR_WARNING}', f'#{COLOR_DANGER}', '#64748B', '#22C55E']
    ax.pie(series.astype(int).tolist(), labels=[str(x)[:35] for x in series.index], autopct='%1.1f%%', startangle=90, colors=palette[:len(series)])
    ax.set_title(title, fontweight='bold')
    ax.axis('equal')
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches='tight')
    plt.close(fig)
    return str(path)


def _generar_grafica_linea(df: pd.DataFrame, path: Path) -> str:
    if df.empty:
        return ''
    fecha_col = None
    for col in ['fecha_carga', 'fecha_ingreso', 'fecha_valoracion', 'fecha_actualizacion']:
        if col in df.columns:
            fecha_col = col
            break
    if not fecha_col:
        return ''
    fechas = pd.to_datetime(df[fecha_col], errors='coerce')
    tmp = pd.DataFrame({'fecha': fechas}).dropna()
    if tmp.empty:
        return ''
    series = tmp.groupby(tmp['fecha'].dt.to_period('M')).size().sort_index()
    if series.empty:
        return ''
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    labels = [str(x) for x in series.index]
    values = series.astype(int).tolist()
    fig = plt.figure(figsize=(9.5, 4.8))
    ax = fig.add_subplot(111)
    ax.plot(labels, values, marker='o', color=f'#{COLOR_PRIMARY}')
    ax.fill_between(range(len(values)), values, alpha=0.15, color=f'#{COLOR_ACCENT}')
    ax.set_title('Tendencia por fecha o periodo disponible', fontweight='bold')
    ax.set_xlabel('Periodo')
    ax.set_ylabel('Registros')
    ax.tick_params(axis='x', rotation=45)
    fig.tight_layout()
    fig.savefig(path, dpi=160, bbox_inches='tight')
    plt.close(fig)
    return str(path)


def _prefer_series(primary: pd.Series, fallback: pd.Series) -> pd.Series:
    return primary if isinstance(primary, pd.Series) and not primary.empty else fallback


def _generar_graficas(df: pd.DataFrame, output_folder: str, tablas: dict[str, Any] | None = None, resumen: dict[str, Any] | None = None) -> list[dict[str, str]]:
    charts: list[dict[str, str]] = []
    if not _figure_available():
        return charts
    tablas = tablas or {}
    resumen = resumen or {}
    folder = Path(output_folder) / f'informe_estadistico_graficas_{datetime.now().strftime("%Y%m%d%H%M%S")}_{uuid.uuid4().hex[:6]}'
    folder.mkdir(parents=True, exist_ok=True)

    pendientes_salud = pd.Series({
        'Sin peso': int(resumen.get('sin_peso') or 0),
        'Sin talla': int(resumen.get('sin_talla') or 0),
        'Sin peso y talla': int(resumen.get('sin_peso_talla') or 0),
        'Sin perímetro braquial': int(resumen.get('sin_perimetro_braquial') or 0),
        'Sin diagnóstico': int(resumen.get('sin_diagnostico_nutricional') or 0),
    })
    alertas_unidad = _series_from_records(tablas.get('por_unidad', []), 'Unidad de servicio', 'alertas_nutricionales', 20)
    movimientos = _series_from_records(tablas.get('movimientos', []), 'movimiento', 'total')

    specs = [
        ('por_unidad', 'Gráfica de barras por unidad de servicio', lambda: _generar_grafica_barras(_prefer_series(_series_from_records(tablas.get('por_unidad', []), 'Unidad de servicio', 'total', 20), _serie_counts(df, 'unidad', 20)), 'Niños por unidad de servicio', folder / 'bar_unidad.png', 'Unidad de servicio')),
        ('por_grupo_etario', 'Gráfica de barras por grupo etario', lambda: _generar_grafica_barras(_prefer_series(_series_from_records(tablas.get('por_grupo_etario', []), 'Grupo etario', 'total'), _serie_counts(df, 'grupo_etario')), 'Niños por grupo etario', folder / 'bar_grupo_etario.png', 'Grupo etario')),
        ('por_sexo', 'Gráfica circular por sexo', lambda: _generar_grafica_pie(_prefer_series(_series_from_records(tablas.get('por_sexo', []), 'Sexo', 'total'), _serie_counts(df, 'sexo')), 'Distribución por sexo', folder / 'pie_sexo.png')),
        ('por_diagnostico', 'Gráfica circular por diagnóstico nutricional', lambda: _generar_grafica_pie(_prefer_series(_series_from_records(tablas.get('por_diagnostico', []), 'Diagnóstico nutricional', 'total', 8), _serie_counts(df, 'diagnostico_nutricional', 8)), 'Distribución por diagnóstico nutricional', folder / 'pie_diagnostico.png')),
        ('faltantes', 'Gráfica de barras de documentos faltantes', lambda: _generar_grafica_barras(_series_from_records(tablas.get('faltantes', []), 'indicador', 'total'), 'Niños con documentos o datos faltantes', folder / 'bar_faltantes.png', 'Indicador')),
        ('pendientes_salud', 'Gráfica de barras de pendientes de salud/nutrición', lambda: _generar_grafica_barras(pendientes_salud[pendientes_salud > 0], 'Pendientes de salud y nutrición', folder / 'bar_pendientes_salud.png', 'Pendiente')),
        ('alertas_unidad', 'Gráfica de alertas críticas por unidad', lambda: _generar_grafica_barras(alertas_unidad[alertas_unidad > 0], 'Alertas críticas por unidad', folder / 'bar_alertas_unidad.png', 'Unidad')),
        ('movimientos', 'Gráfica de movimientos: nuevos, retirados y permanencias', lambda: _generar_grafica_barras(movimientos, 'Movimientos entre bases', folder / 'bar_movimientos.png', 'Movimiento')),
        ('linea_periodo', 'Gráfica de línea por fecha o periodo', lambda: _generar_grafica_linea(df, folder / 'linea_periodo.png')),
    ]
    for key, title, fn in specs:
        try:
            path = fn()
            if path and os.path.exists(path):
                charts.append({'key': key, 'titulo': title, 'path': path})
        except Exception:
            continue
    return charts

def _filtros_label(filtros: dict[str, Any]) -> str:
    labels = []
    if _safe_text(filtros.get('unidad')):
        labels.append(f"Unidad: {_safe_text(filtros.get('unidad'))}")
    if _safe_text(filtros.get('coordinador')):
        labels.append(f"Coordinador: {_safe_text(filtros.get('coordinador'))}")
    if _safe_text(filtros.get('grupo_etario')):
        labels.append(f"Grupo etario: {_safe_text(filtros.get('grupo_etario'))}")
    if _safe_text(filtros.get('estado_nutricional') or filtros.get('diagnostico')):
        labels.append(f"Diagnóstico: {_safe_text(filtros.get('estado_nutricional') or filtros.get('diagnostico'))}")
    if str(filtros.get('alertas') or '').lower() in {'1', 'true', 'si', 'sí'}:
        labels.append('Solo niños con alertas')
    if str(filtros.get('faltantes') or '').lower() in {'1', 'true', 'si', 'sí'}:
        labels.append('Solo niños con datos faltantes')
    return ' | '.join(labels) if labels else 'General'


def construir_contexto_informe(database_path: str, row_cruce: dict[str, Any], resultado: dict[str, Any], filtros: dict[str, Any] | None, usuario_ctx: dict[str, Any] | None, output_folder: str) -> dict[str, Any]:
    filtros = dict(filtros or {})
    usuario_ctx = dict(usuario_ctx or {})
    fundacion_id = int(usuario_ctx.get('fundacion_id') or row_cruce.get('fundacion_id') or 1)
    superadmin = str(usuario_ctx.get('rol') or '').upper() == 'SUPERADMIN'

    conn = _connect(database_path)
    try:
        df_master, duplicados, fuente_maestra = _cargar_base_maestra(conn, fundacion_id, superadmin)
        df_master, unidad_coord = _enriquecer_coordinador(df_master, conn, fundacion_id, superadmin)
        df_master, notas_salud, alertas_df = _enriquecer_salud_nutricion(df_master, conn, fundacion_id, superadmin)
        df_scope = _apply_filters(df_master, filtros)
        movimientos = _movimientos_contexto(resultado, df_scope, filtros, unidad_coord)
        alertas_prioritarias = _alertas_prioritarias(alertas_df, df_scope)
    finally:
        conn.close()

    salud = _indicadores_salud(df_scope)
    resumen_cruce = resultado.get('resumen') or {}
    total_actual = int(resumen_cruce.get('total_actual') or len(df_scope))
    nuevos = len(movimientos.get('nuevos', []))
    retirados = len(movimientos.get('retirados', []))
    permanecen = max(0, int(total_actual) - int(resumen_cruce.get('nuevos') or nuevos)) if not filtros else max(0, len(df_scope) - nuevos)
    cambiaron_unidad = len(movimientos.get('cambios_unidad', [])) or len(movimientos.get('trasladados', []))

    resumen_ejecutivo = {
        'total_ninos': int(len(df_scope)),
        'total_unidades': int(df_scope['unidad'].nunique()) if 'unidad' in df_scope.columns and not df_scope.empty else 0,
        'total_coordinadores': int(df_scope['coordinador'].replace('', 'SIN COORDINADOR').nunique()) if 'coordinador' in df_scope.columns and not df_scope.empty else 0,
        'total_activos': _conteo_estado_activo(df_scope),
        'total_nuevos': nuevos,
        'total_retirados': retirados,
        'total_permanecen': permanecen,
        'total_cambios_unidad': cambiaron_unidad,
        'diferencia_bases': int(resumen_cruce.get('total_actual') or 0) - int(resumen_cruce.get('total_anterior') or 0),
        'duplicados_detectados': len({d.get('documento') for d in duplicados if d.get('documento')}),
    }

    tablas = {
        'por_unidad': _tabla_por_unidad(df_scope, movimientos),
        'por_coordinador': _tabla_por_coordinador(df_scope),
        'por_grupo_etario': _tabla_generica_estadistica(df_scope, 'grupo_etario', 'Grupo etario'),
        'por_sexo': _tabla_generica_estadistica(df_scope, 'sexo', 'Sexo'),
        'por_estado': _tabla_generica_estadistica(df_scope, 'estado', 'Estado del niño'),
        'por_diagnostico': _tabla_generica_estadistica(df_scope, 'diagnostico_nutricional', 'Diagnóstico nutricional'),
        'faltantes': _faltantes_table(df_scope),
        'alertas_prioritarias': alertas_prioritarias,
        'movimientos': _tabla_movimientos_resumen(resumen_ejecutivo),
    }

    anexos = {
        'sin_carne_salud': _listado(df_scope, df_scope.get('sin_carne_salud', pd.Series([False] * len(df_scope), index=df_scope.index)), 'Sin carné de salud'),
        'sin_control_crecimiento': _listado(df_scope, df_scope.get('sin_control_crecimiento', pd.Series([False] * len(df_scope), index=df_scope.index)), 'Sin control de crecimiento y desarrollo'),
        'sin_vacunas': _listado(df_scope, df_scope.get('sin_vacunas', pd.Series([False] * len(df_scope), index=df_scope.index)), 'Sin vacunas'),
        'sin_registro_civil': _listado(df_scope, df_scope.get('sin_registro_civil', pd.Series([False] * len(df_scope), index=df_scope.index)), 'Sin registro civil / documento válido'),
        'sin_peso_talla': _listado(df_scope, df_scope.get('sin_peso_talla', pd.Series([False] * len(df_scope), index=df_scope.index)), 'Sin peso y talla'),
        'sin_perimetro_braquial': _listado(df_scope, df_scope.get('sin_perimetro_braquial', pd.Series([False] * len(df_scope), index=df_scope.index)), 'Sin perímetro braquial'),
        'sin_diagnostico': _listado(df_scope, df_scope.get('sin_diagnostico_nutricional', pd.Series([False] * len(df_scope), index=df_scope.index)), 'Sin diagnóstico nutricional'),
        'nuevos': _listado_movimiento(movimientos.get('nuevos', []), 'nuevos'),
        'retirados': _listado_movimiento(movimientos.get('retirados', []), 'retirados'),
        'cambios_unidad': _listado_movimiento((movimientos.get('cambios_unidad', []) or movimientos.get('trasladados', [])), 'cambios_unidad'),
        'duplicados': duplicados[:MAX_ANEXO_ROWS],
    }

    resumen_para_graficas = {**resumen_ejecutivo, **salud}
    graficas = _generar_graficas(df_scope, output_folder, tablas=tablas, resumen=resumen_para_graficas)
    filtros_label = _filtros_label(filtros)
    corporacion_id = filtros.get('corporacion_id') or usuario_ctx.get('corporacion_id')
    try:
        corporacion_id = int(corporacion_id) if corporacion_id else None
    except Exception:
        corporacion_id = None
    corporacion = obtener_corporacion_contexto(database_path, fundacion_id, corporacion_id)
    ctx = {
        'fecha_generacion': datetime.now().strftime('%d/%m/%Y %I:%M %p'),
        'usuario': usuario_ctx.get('usuario') or usuario_ctx.get('username') or usuario_ctx.get('email') or 'sistema',
        'rol': usuario_ctx.get('rol') or 'SUPERADMIN',
        'fundacion_id': fundacion_id,
        'filtros': filtros,
        'filtros_label': filtros_label,
        'row_cruce': row_cruce,
        'resultado': resultado,
        'resumen_cruce': resumen_cruce,
        'resumen_ejecutivo': resumen_ejecutivo,
        'salud': salud,
        'tablas': tablas,
        'anexos': anexos,
        'graficas': graficas,
        'notas': notas_salud,
        'fuente_maestra': fuente_maestra,
        'corporacion': corporacion,
        'logo_path': obtener_logo_corporacion(database_path, fundacion_id, corporacion_id, output_folder),
    }
    return ctx


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def _format_table(table, header_fill: str = COLOR_PRIMARY) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    table.style = 'Table Grid'
    table.autofit = True

    def ensure_row_property(row, tag: str, attrs: dict[str, str] | None = None) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn(tag)) is not None:
            return
        el = OxmlElement(tag)
        for key, value in (attrs or {}).items():
            el.set(qn(key), value)
        tr_pr.append(el)

    for row_idx, row in enumerate(table.rows):
        ensure_row_property(row, 'w:cantSplit')
        if row_idx == 0:
            ensure_row_property(row, 'w:tblHeader', {'w:val': 'true'})
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
            if row_idx == 0:
                _set_cell_shading(cell, header_fill)
                from docx.shared import RGBColor
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def _add_page_field(paragraph, field_code: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = field_code
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _configure_docx(doc, ctx: dict[str, Any]) -> None:
    from docx.enum.section import WD_SECTION_START
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(10)
    for style_name, size, color in [('Title', 21, COLOR_PRIMARY), ('Heading 1', 15, COLOR_PRIMARY), ('Heading 2', 12, COLOR_SECONDARY)]:
        style = styles[style_name]
        style.font.name = 'Arial'
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    corporacion = ctx.get('corporacion') or {}
    corp_name = _safe_text(corporacion.get('nombre'), DEFAULT_CORPORACION)
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = ''
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ctx.get('logo_path') and os.path.exists(ctx['logo_path']):
        try:
            header_para.add_run().add_picture(ctx['logo_path'], width=Cm(1.25))
            header_para.add_run('  ')
        except Exception:
            pass
    header_para.add_run(f'{corp_name} | Informe estadístico institucional de Primera Infancia')
    for run in header_para.runs:
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLOR_PRIMARY)

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = ''
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run(f'{corp_name} | Generado: {ctx.get("fecha_generacion", "")} | Página ')
    _add_page_field(footer_para, 'PAGE')
    footer_para.add_run(' de ')
    _add_page_field(footer_para, 'NUMPAGES')
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string('666666')


def _add_paragraph(doc, text: str, style: str | None = None, justify: bool = True):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph(text, style=style)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _add_heading(doc, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def _add_key_value_table(doc, rows: list[tuple[str, Any]], title: str | None = None):
    if title:
        _add_heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'Indicador'
    table.rows[0].cells[1].text = 'Valor'
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)
    _format_table(table)
    return table


def _add_records_table(doc, records: list[dict[str, Any]], headers: list[tuple[str, str]], empty_text: str = 'Sin registros para mostrar.', max_rows: int = MAX_ANEXO_ROWS):
    if not records:
        _add_paragraph(doc, empty_text, justify=False)
        return None
    table = doc.add_table(rows=1, cols=len(headers))
    for i, (_, label) in enumerate(headers):
        table.rows[0].cells[i].text = label
    for record in records[:max_rows]:
        cells = table.add_row().cells
        for i, (key, _) in enumerate(headers):
            cells[i].text = _safe_text(record.get(key))[:280]
    _format_table(table)
    if len(records) > max_rows:
        _add_paragraph(doc, f'Se muestran {max_rows} de {len(records)} registros. El archivo conserva el resumen completo en las tablas estadísticas.', justify=False)
    return table


def _add_chart_image(doc, chart: dict[str, str] | None, width_cm: float = 16.4) -> bool:
    if not chart or not chart.get('path') or not os.path.exists(chart.get('path', '')):
        return False
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(chart['path'], width=Cm(width_cm))
        return True
    except Exception:
        _add_paragraph(doc, 'La gráfica no pudo insertarse en el documento.', justify=False)
        return False


def _chart_by_key(ctx: dict[str, Any], key: str) -> dict[str, str] | None:
    for chart in ctx.get('graficas', []) or []:
        if chart.get('key') == key:
            return chart
    return None


def _crear_docx(ctx: dict[str, Any], output_folder: str, prefix: str) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    Path(output_folder).mkdir(parents=True, exist_ok=True)
    doc = Document()
    _configure_docx(doc, ctx)

    # Portada institucional
    corporacion = ctx.get('corporacion') or {}
    corp_name = _safe_text(corporacion.get('nombre'), DEFAULT_CORPORACION)
    if ctx.get('logo_path') and os.path.exists(ctx['logo_path']):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(ctx['logo_path'], width=Cm(6.5))
    p_corp = doc.add_paragraph()
    p_corp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = p_corp.add_run(corp_name.upper())
    rc.bold = True
    rc.font.size = Pt(15)
    rc.font.color.rgb = RGBColor.from_string(COLOR_SECONDARY)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('INFORME ESTADÍSTICO INSTITUCIONAL\n')
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(COLOR_PRIMARY)
    r2 = title.add_run('Cruce de bases y Base Maestra de Primera Infancia')
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor.from_string(COLOR_SECONDARY)
    _add_key_value_table(doc, [
        ('Nombre de la corporación', corp_name),
        ('NIT', _safe_text(corporacion.get('nit'), 'NIT no registrado')),
        ('Nombre de la plataforma', 'PrimeraInfancia'),
        ('Nombre del informe', 'Informe estadístico posterior al cruce de bases'),
        ('Fecha de generación', ctx['fecha_generacion']),
        ('Usuario que genera', ctx['usuario']),
        ('Alcance / filtro aplicado', ctx['filtros_label']),
        ('Periodo del cruce', ctx.get('row_cruce', {}).get('periodo') or ''),
        ('Fuente maestra consultada', ctx.get('fuente_maestra', 'Base Maestra publicada')),
    ], title='Datos institucionales del informe')
    _add_paragraph(doc, 'Documento generado automáticamente desde la Base Maestra consolidada. El identificador principal usado para consolidar y evitar dobles conteos es el documento del niño o beneficiario.', justify=True)
    doc.add_page_break()

    resumen = ctx['resumen_ejecutivo']
    salud = ctx['salud']
    _add_heading(doc, '1. Resumen ejecutivo', 1)
    _add_key_value_table(doc, [
        ('Total de niños', resumen['total_ninos']),
        ('Total de unidades', resumen['total_unidades']),
        ('Total de coordinadores', resumen['total_coordinadores']),
        ('Total de niños activos', resumen['total_activos']),
        ('Total de niños nuevos', resumen['total_nuevos']),
        ('Total de niños retirados', resumen['total_retirados']),
        ('Total de niños que permanecen', resumen['total_permanecen']),
        ('Niños que cambiaron de unidad', resumen['total_cambios_unidad']),
        ('Diferencia entre bases cargadas', resumen['diferencia_bases']),
        ('Documentos duplicados detectados', resumen['duplicados_detectados']),
    ])
    _add_paragraph(doc, 'El resumen toma como fuente los resultados del último cruce seleccionado y la Base Maestra vigente. Los duplicados se consolidan por documento y se reportan en anexos de inconsistencias.', justify=True)

    _add_heading(doc, '2. Análisis de movimientos', 1)
    _add_key_value_table(doc, [
        ('Niños que ingresaron', len(ctx['anexos']['nuevos'])),
        ('Niños que salieron', len(ctx['anexos']['retirados'])),
        ('Niños que cambiaron de unidad', len(ctx['anexos']['cambios_unidad'])),
        ('Niños que permanecen', resumen['total_permanecen']),
        ('Diferencia entre base anterior y base actual', resumen['diferencia_bases']),
    ])

    _add_heading(doc, '3. Indicadores de salud y nutrición', 1)
    _add_key_value_table(doc, [
        ('Sin carné de salud', salud['sin_carne_salud']),
        ('Sin control de crecimiento y desarrollo', salud['sin_control_crecimiento']),
        ('Sin vacunas', salud['sin_vacunas']),
        ('Sin registro civil / documento válido', salud.get('sin_registro_civil', 0)),
        ('Sin carné de crecimiento y desarrollo', salud['sin_carne_crecimiento']),
        ('Sin peso', salud['sin_peso']),
        ('Sin talla', salud['sin_talla']),
        ('Sin peso y talla', salud['sin_peso_talla']),
        ('Sin perímetro braquial', salud['sin_perimetro_braquial']),
        ('Sin diagnóstico nutricional', salud['sin_diagnostico_nutricional']),
        ('Con bajo peso', salud['bajo_peso']),
        ('Con riesgo de bajo peso', salud['riesgo_bajo_peso']),
        ('Con sobrepeso', salud['sobrepeso']),
        ('Con obesidad', salud['obesidad']),
        ('Con talla baja', salud['talla_baja']),
        ('Con alertas nutricionales', salud['alertas_nutricionales']),
    ])
    if ctx.get('notas'):
        _add_heading(doc, 'Notas de disponibilidad de datos', 2)
        for nota in ctx['notas'][:10]:
            _add_paragraph(doc, f'- {nota}', justify=False)

    _add_heading(doc, '4. Tablas estadísticas', 1)
    table_specs = [
        ('Unidad de servicio', 'por_unidad', [('Unidad de servicio', 'Unidad de servicio'), ('total', 'Total'), ('activos', 'Activos'), ('retirados', 'Retirados'), ('nuevos', 'Nuevos'), ('sin_peso', 'Sin peso'), ('sin_talla', 'Sin talla'), ('sin_peso_talla', 'Sin peso/talla'), ('alertas_nutricionales', 'Alertas')]),
        ('Coordinador', 'por_coordinador', [('Coordinador', 'Coordinador'), ('unidades_asignadas', 'Unidades'), ('total', 'Total niños'), ('ninos_con_pendientes', 'Con pendientes'), ('alertas_prioritarias', 'Alertas prioritarias')]),
        ('Grupo etario', 'por_grupo_etario', [('Grupo etario', 'Grupo etario'), ('total', 'Total'), ('activos', 'Activos'), ('retirados', 'Retirados'), ('alertas', 'Alertas'), ('datos_faltantes', 'Datos faltantes')]),
        ('Sexo', 'por_sexo', [('Sexo', 'Sexo'), ('total', 'Total'), ('activos', 'Activos'), ('retirados', 'Retirados'), ('alertas', 'Alertas'), ('datos_faltantes', 'Datos faltantes')]),
        ('Estado del niño', 'por_estado', [('Estado del niño', 'Estado'), ('total', 'Total'), ('activos', 'Activos'), ('retirados', 'Retirados'), ('alertas', 'Alertas'), ('datos_faltantes', 'Datos faltantes')]),
        ('Diagnóstico nutricional', 'por_diagnostico', [('Diagnóstico nutricional', 'Diagnóstico'), ('total', 'Total'), ('activos', 'Activos'), ('retirados', 'Retirados'), ('alertas', 'Alertas'), ('datos_faltantes', 'Datos faltantes')]),
        ('Faltantes documentales y de salud', 'faltantes', [('indicador', 'Indicador'), ('total', 'Total')]),
        ('Alertas prioritarias', 'alertas_prioritarias', [('nivel', 'Nivel'), ('tipo', 'Tipo'), ('total', 'Total')]),
        ('Movimientos', 'movimientos', [('movimiento', 'Movimiento'), ('total', 'Total')]),
    ]
    inserted_charts = set()
    for titulo, key, headers in table_specs:
        _add_heading(doc, titulo, 2)
        _add_records_table(doc, ctx['tablas'].get(key, []), headers, empty_text=f'Sin datos para {titulo.lower()}.', max_rows=80)
        chart = _chart_by_key(ctx, key)
        if chart:
            _add_chart_image(doc, chart)
            inserted_charts.add(chart.get('key'))
        if key == 'faltantes':
            chart = _chart_by_key(ctx, 'pendientes_salud')
            if chart:
                _add_chart_image(doc, chart)
                inserted_charts.add(chart.get('key'))
        if key == 'alertas_prioritarias':
            chart = _chart_by_key(ctx, 'alertas_unidad')
            if chart:
                _add_chart_image(doc, chart)
                inserted_charts.add(chart.get('key'))

    _add_heading(doc, '5. Gráficas complementarias', 1)
    remaining = [c for c in ctx.get('graficas', []) if c.get('key') not in inserted_charts]
    if not remaining and not ctx.get('graficas'):
        _add_paragraph(doc, 'No fue posible generar gráficas automáticas porque matplotlib no está disponible o no hay datos suficientes.', justify=False)
    for chart in remaining:
        _add_heading(doc, chart['titulo'], 2)
        _add_chart_image(doc, chart)

    _add_heading(doc, '6. Anexos detallados', 1)
    anexo_specs = [
        ('Niños sin carné de salud', 'sin_carne_salud'),
        ('Niños sin crecimiento y desarrollo', 'sin_control_crecimiento'),
        ('Niños sin vacunas', 'sin_vacunas'),
        ('Niños sin registro civil / documento válido', 'sin_registro_civil'),
        ('Niños sin peso y talla', 'sin_peso_talla'),
        ('Niños sin perímetro braquial', 'sin_perimetro_braquial'),
        ('Niños sin diagnóstico', 'sin_diagnostico'),
        ('Niños nuevos', 'nuevos'),
        ('Niños retirados', 'retirados'),
        ('Niños cambiados de unidad', 'cambios_unidad'),
        ('Niños duplicados o inconsistentes', 'duplicados'),
    ]
    default_headers = [('documento', 'Documento'), ('nombre_completo', 'Nombre'), ('unidad', 'Unidad'), ('grupo_etario', 'Grupo'), ('estado', 'Estado'), ('observacion', 'Observación')]
    movement_headers = [('documento', 'Documento'), ('nombre_completo', 'Nombre'), ('unidad', 'Unidad'), ('unidad_anterior', 'Unidad anterior'), ('unidad_actual', 'Unidad actual'), ('observacion', 'Observación')]
    for titulo, key in anexo_specs:
        _add_heading(doc, titulo, 2)
        headers = movement_headers if key in {'nuevos', 'retirados', 'cambios_unidad', 'duplicados'} else default_headers
        _add_records_table(doc, ctx['anexos'].get(key, []), headers, empty_text='Sin registros en este anexo.', max_rows=MAX_ANEXO_ROWS)

    filename = f"{prefix}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    path = os.path.join(output_folder, filename)
    doc.save(path)
    return path


def _crear_pdf_reportlab(ctx: dict[str, Any], output_folder: str, prefix: str) -> str:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    Path(output_folder).mkdir(parents=True, exist_ok=True)
    path = os.path.join(output_folder, f"{prefix}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Small', parent=styles['BodyText'], fontSize=8, leading=10))

    def kv_table(rows):
        data = [['Indicador', 'Valor']] + [[str(k), str(v)] for k, v in rows]
        t = Table(data, repeatRows=1, colWidths=[260, 190])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        return t

    story = []
    corporacion = ctx.get('corporacion') or {}
    corp_name = _safe_text(corporacion.get('nombre'), DEFAULT_CORPORACION)
    if ctx.get('logo_path') and os.path.exists(ctx['logo_path']):
        story.append(Image(ctx['logo_path'], width=210, height=65))
    story.append(Paragraph(corp_name.upper(), styles['Heading2']))
    story.append(Paragraph('INFORME ESTADÍSTICO INSTITUCIONAL', styles['Title']))
    story.append(Paragraph('Cruce de bases y Base Maestra de Primera Infancia', styles['Heading2']))
    story.append(Spacer(1, 8))
    story.append(kv_table([
        ('Nombre de la corporación', corp_name),
        ('NIT', _safe_text(corporacion.get('nit'), 'NIT no registrado')),
        ('Fecha de generación', ctx['fecha_generacion']),
        ('Usuario que genera', ctx['usuario']),
        ('Alcance / filtro aplicado', ctx['filtros_label']),
        ('Periodo del cruce', ctx.get('row_cruce', {}).get('periodo') or ''),
        ('Fuente maestra consultada', ctx.get('fuente_maestra', 'Base Maestra publicada')),
    ]))
    story.append(PageBreak())

    resumen = ctx['resumen_ejecutivo']
    salud = ctx['salud']
    story.append(Paragraph('1. Resumen ejecutivo', styles['Heading1']))
    story.append(kv_table([
        ('Total de niños', resumen['total_ninos']),
        ('Total de unidades', resumen['total_unidades']),
        ('Total de coordinadores', resumen['total_coordinadores']),
        ('Total de niños activos', resumen['total_activos']),
        ('Total de niños nuevos', resumen['total_nuevos']),
        ('Total de niños retirados', resumen['total_retirados']),
        ('Total de niños que permanecen', resumen['total_permanecen']),
        ('Niños que cambiaron de unidad', resumen['total_cambios_unidad']),
        ('Diferencia entre bases cargadas', resumen['diferencia_bases']),
        ('Documentos duplicados detectados', resumen['duplicados_detectados']),
    ]))
    story.append(Spacer(1, 10))
    story.append(Paragraph('2. Indicadores de salud y nutrición', styles['Heading1']))
    story.append(kv_table([
        ('Sin carné de salud', salud['sin_carne_salud']),
        ('Sin control de crecimiento y desarrollo', salud['sin_control_crecimiento']),
        ('Sin vacunas', salud['sin_vacunas']),
        ('Sin registro civil / documento válido', salud.get('sin_registro_civil', 0)),
        ('Sin carné de crecimiento y desarrollo', salud['sin_carne_crecimiento']),
        ('Sin peso', salud['sin_peso']),
        ('Sin talla', salud['sin_talla']),
        ('Sin peso y talla', salud['sin_peso_talla']),
        ('Sin perímetro braquial', salud['sin_perimetro_braquial']),
        ('Sin diagnóstico nutricional', salud['sin_diagnostico_nutricional']),
        ('Con bajo peso', salud['bajo_peso']),
        ('Con riesgo de bajo peso', salud['riesgo_bajo_peso']),
        ('Con sobrepeso', salud['sobrepeso']),
        ('Con obesidad', salud['obesidad']),
        ('Con talla baja', salud['talla_baja']),
        ('Con alertas nutricionales', salud['alertas_nutricionales']),
    ]))
    story.append(PageBreak())

    story.append(Paragraph('3. Gráficas', styles['Heading1']))
    for chart in ctx.get('graficas', [])[:6]:
        story.append(Paragraph(chart['titulo'], styles['Heading2']))
        try:
            story.append(Image(chart['path'], width=500, height=260))
        except Exception:
            story.append(Paragraph('No fue posible insertar la gráfica.', styles['Small']))
        story.append(Spacer(1, 8))

    story.append(PageBreak())
    story.append(Paragraph('4. Tablas estadísticas', styles['Heading1']))
    for titulo, key, headers in [
        ('Unidad de servicio', 'por_unidad', ['Unidad de servicio', 'total', 'activos', 'retirados', 'nuevos', 'sin_peso', 'sin_talla', 'sin_peso_talla', 'alertas_nutricionales']),
        ('Coordinador', 'por_coordinador', ['Coordinador', 'unidades_asignadas', 'total', 'ninos_con_pendientes', 'alertas_prioritarias']),
        ('Grupo etario', 'por_grupo_etario', ['Grupo etario', 'total', 'activos', 'retirados', 'alertas', 'datos_faltantes']),
        ('Faltantes documentales y de salud', 'faltantes', ['indicador', 'total']),
        ('Movimientos', 'movimientos', ['movimiento', 'total']),
    ]:
        records = ctx['tablas'].get(key, [])[:80]
        story.append(Paragraph(titulo, styles['Heading2']))
        if records:
            data = [headers] + [[str(r.get(h, '')) for h in headers] for r in records]
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E78')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
                ('FONTSIZE', (0, 0), (-1, -1), 7),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(t)
        else:
            story.append(Paragraph('Sin registros.', styles['Small']))
        story.append(Spacer(1, 8))

    doc.build(story)
    return path


def _convertir_docx_a_pdf(docx_path: str, output_folder: str, prefix: str) -> str | None:
    candidates = [shutil.which('soffice'), shutil.which('libreoffice')]
    candidates = [c for c in candidates if c]
    if not candidates:
        return None
    tmp_dir = Path(output_folder) / f'pdf_tmp_{uuid.uuid4().hex[:8]}'
    tmp_dir.mkdir(parents=True, exist_ok=True)
    for cmd in candidates:
        try:
            subprocess.run([cmd, '--headless', '--convert-to', 'pdf', '--outdir', str(tmp_dir), docx_path], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=90)
            converted = tmp_dir / (Path(docx_path).stem + '.pdf')
            if converted.exists() and converted.stat().st_size > 0:
                final = Path(output_folder) / f"{prefix}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
                shutil.move(str(converted), final)
                shutil.rmtree(tmp_dir, ignore_errors=True)
                return str(final)
        except Exception:
            continue
    shutil.rmtree(tmp_dir, ignore_errors=True)
    return None


def crear_informe_estadistico(database_path: str, output_folder: str, row_cruce: dict[str, Any], resultado: dict[str, Any], filtros: dict[str, Any] | None, usuario_ctx: dict[str, Any] | None, formato: str = 'docx') -> str:
    formato = (formato or 'docx').lower().strip()
    safe_scope = re.sub(r'[^A-Z0-9]+', '_', _upper(_filtros_label(filtros or {}))).strip('_')[:48] or 'GENERAL'
    prefix = f"INFORME_ESTADISTICO_{safe_scope}"
    ctx = construir_contexto_informe(database_path, row_cruce, resultado, filtros, usuario_ctx, output_folder)
    docx_path = _crear_docx(ctx, output_folder, prefix)
    if formato in {'docx', 'word'}:
        return docx_path
    if formato == 'pdf':
        converted = _convertir_docx_a_pdf(docx_path, output_folder, prefix)
        if converted:
            return converted
        return _crear_pdf_reportlab(ctx, output_folder, prefix)
    raise ValueError('Formato de informe no soportado. Usa docx o pdf.')


def obtener_opciones_informe(database_path: str, fundacion_id: int | None = None, superadmin: bool = False) -> dict[str, list[str]]:
    conn = _connect(database_path)
    try:
        df, _, _ = _cargar_base_maestra(conn, fundacion_id, superadmin)
        df, _ = _enriquecer_coordinador(df, conn, fundacion_id, superadmin)
        df, _, _ = _enriquecer_salud_nutricion(df, conn, fundacion_id, superadmin)
    finally:
        conn.close()
    if df.empty:
        return {'unidades': [], 'coordinadores': [], 'grupos_etarios': [], 'diagnosticos': []}

    def uniq(col: str) -> list[str]:
        if col not in df.columns:
            return []
        vals = sorted({_safe_text(v) for v in df[col].dropna().tolist() if _safe_text(v)})
        return vals[:300]

    return {
        'unidades': uniq('unidad'),
        'coordinadores': uniq('coordinador'),
        'grupos_etarios': uniq('grupo_etario'),
        'diagnosticos': uniq('diagnostico_nutricional'),
    }
