"""
Rutas Flask del módulo independiente Gestión Pedagógica.

Se registra como Blueprint para no tocar la lógica existente de la plataforma.
"""
from __future__ import annotations

import os
import re
import csv
import zipfile
from datetime import datetime

from flask import Blueprint, jsonify, request, send_from_directory

from .documentos import save_uploaded_document
from .repository import GestionPedagogicaRepository, now_iso
from .services import (
    ESTADOS_ENTREGABLE,
    calendario_operativo_default,
    dashboard_pedagogico,
    generar_alertas,
    periodo_actual,
    reporte_mensual,
)


def json_data() -> dict:
    return request.get_json(silent=True) or {}


def register_gestion_pedagogica(app, database_path: str, upload_folder: str) -> None:
    repo = GestionPedagogicaRepository(database_path)
    repo.init_schema()

    bp = Blueprint('gestion_pedagogica', __name__, url_prefix='/api/gestion-pedagogica')

    @bp.before_request
    def _ensure_schema():
        repo.init_schema()

    @bp.route('/dashboard', methods=['GET'])
    def dashboard():
        periodo = request.args.get('periodo') or periodo_actual()
        return jsonify(dashboard_pedagogico(repo, periodo)), 200

    @bp.route('/coordinadores', methods=['GET', 'POST'])
    def coordinadores():
        if request.method == 'GET':
            incluir_inactivos = str(request.args.get('incluir_inactivos', '')).lower() in {'1', 'true', 'si', 'sí'}
            return jsonify({'coordinadores': repo.listar_coordinadores(incluir_inactivos)}), 200

        data = json_data()
        if not data.get('nombre'):
            return jsonify({'error': 'El nombre del coordinador es obligatorio.'}), 400
        coord = repo.crear_coordinador(data)
        return jsonify({'message': 'Coordinador creado correctamente.', 'coordinador': coord}), 201

    @bp.route('/coordinadores/<int:coordinador_id>', methods=['GET', 'PUT', 'PATCH', 'DELETE'])
    def coordinador_detalle(coordinador_id: int):
        if request.method == 'GET':
            coord = repo.obtener_coordinador(coordinador_id)
            if not coord:
                return jsonify({'error': 'Coordinador no encontrado.'}), 404
            detalle = {
                'coordinador': coord,
                'equipo': repo.listar_equipos(coordinador_id),
                'docentes': repo.listar_docentes(coordinador_id),
                'entregables': repo.listar_entregables(coordinador_id=coordinador_id),
                'documentos': repo.fetch_all(
                    "SELECT * FROM gp_documentos WHERE coordinador_id = ? AND activo = 1 ORDER BY fecha_carga DESC",
                    (coordinador_id,),
                ),
                'calendario': repo.fetch_all(
                    "SELECT * FROM gp_calendario_eventos WHERE coordinador_id = ? ORDER BY fecha DESC",
                    (coordinador_id,),
                )
            }
            return jsonify(detalle), 200

        if request.method in {'PUT', 'PATCH'}:
            updated = repo.actualizar_coordinador(coordinador_id, json_data())
            if not updated:
                return jsonify({'error': 'Coordinador no encontrado.'}), 404
            return jsonify({'message': 'Coordinador actualizado.', 'coordinador': updated}), 200

        if not repo.desactivar_coordinador(coordinador_id):
            return jsonify({'error': 'Coordinador no encontrado.'}), 404
        return jsonify({'message': 'Coordinador eliminado/desactivado.'}), 200

    @bp.route('/equipos', methods=['GET', 'POST'])
    def equipos():
        if request.method == 'GET':
            coordinador_id = request.args.get('coordinador_id', type=int)
            return jsonify({'equipos': repo.listar_equipos(coordinador_id)}), 200

        data = json_data()
        if not data.get('nombre'):
            return jsonify({'error': 'El nombre del integrante es obligatorio.'}), 400
        if not data.get('rol'):
            return jsonify({'error': 'El rol del integrante es obligatorio.'}), 400
        equipo = repo.crear_equipo(data)
        return jsonify({'message': 'Integrante creado correctamente.', 'integrante': equipo}), 201

    @bp.route('/equipos/<int:equipo_id>', methods=['PUT', 'PATCH', 'DELETE'])
    def equipo_detalle(equipo_id: int):
        if request.method in {'PUT', 'PATCH'}:
            updated = repo.actualizar_equipo(equipo_id, json_data())
            if not updated:
                return jsonify({'error': 'Integrante no encontrado.'}), 404
            return jsonify({'message': 'Integrante actualizado.', 'integrante': updated}), 200

        if not repo.desactivar_equipo(equipo_id):
            return jsonify({'error': 'Integrante no encontrado.'}), 404
        return jsonify({'message': 'Integrante eliminado/desactivado.'}), 200

    @bp.route('/docentes', methods=['GET', 'POST'])
    def docentes():
        if request.method == 'GET':
            coordinador_id = request.args.get('coordinador_id', type=int)
            return jsonify({'docentes': repo.listar_docentes(coordinador_id)}), 200

        data = json_data()
        if not data.get('nombre'):
            return jsonify({'error': 'El nombre del docente es obligatorio.'}), 400
        docente = repo.crear_docente(data)
        return jsonify({'message': 'Docente creado correctamente.', 'docente': docente}), 201

    @bp.route('/docentes/<int:docente_id>', methods=['PUT', 'PATCH', 'DELETE'])
    def docente_detalle(docente_id: int):
        if request.method in {'PUT', 'PATCH'}:
            updated = repo.actualizar_docente(docente_id, json_data())
            if not updated:
                return jsonify({'error': 'Docente no encontrado.'}), 404
            return jsonify({'message': 'Docente actualizado.', 'docente': updated}), 200

        if not repo.desactivar_docente(docente_id):
            return jsonify({'error': 'Docente no encontrado.'}), 404
        return jsonify({'message': 'Docente eliminado/desactivado.'}), 200

    @bp.route('/entregables', methods=['GET', 'POST'])
    def entregables():
        if request.method == 'GET':
            periodo = request.args.get('periodo')
            coordinador_id = request.args.get('coordinador_id', type=int)
            estado = request.args.get('estado')
            entregables_data = repo.listar_entregables(periodo=periodo, coordinador_id=coordinador_id, estado=estado)
            return jsonify({'entregables': entregables_data, 'estados': ESTADOS_ENTREGABLE}), 200

        data = json_data()
        if not data.get('tipo'):
            return jsonify({'error': 'El tipo de entregable es obligatorio.'}), 400
        data.setdefault('periodo', periodo_actual())
        data.setdefault('titulo', data.get('tipo'))
        entregable = repo.crear_entregable(data)
        return jsonify({'message': 'Entregable creado correctamente.', 'entregable': entregable}), 201

    @bp.route('/entregables/<int:entregable_id>', methods=['PUT', 'PATCH', 'DELETE'])
    def entregable_detalle(entregable_id: int):
        if request.method in {'PUT', 'PATCH'}:
            updated = repo.actualizar_entregable(entregable_id, json_data())
            if not updated:
                return jsonify({'error': 'Entregable no encontrado.'}), 404
            return jsonify({'message': 'Entregable actualizado.', 'entregable': updated}), 200

        if not repo.desactivar_entregable(entregable_id):
            return jsonify({'error': 'Entregable no encontrado.'}), 404
        return jsonify({'message': 'Entregable eliminado/desactivado.'}), 200

    @bp.route('/calendario', methods=['GET'])
    def calendario():
        periodo = request.args.get('periodo')
        coordinador_id = request.args.get('coordinador_id', type=int)
        estado = request.args.get('estado')
        tipo = request.args.get('tipo')

        where = ["1=1"]
        params = []
        if periodo:
            where.append("substr(fecha, 1, 7) = ?")
            params.append(periodo)
        if coordinador_id:
            where.append("coordinador_id = ?")
            params.append(coordinador_id)
        if estado:
            where.append("estado = ?")
            params.append(estado)
        if tipo:
            where.append("tipo = ?")
            params.append(tipo)

        eventos = repo.fetch_all(
            f"""
            SELECT ev.*, c.nombre AS coordinador_nombre
            FROM gp_calendario_eventos ev
            LEFT JOIN gp_coordinadores c ON c.id = ev.coordinador_id
            WHERE {' AND '.join(where)}
            ORDER BY ev.fecha, ev.hora
            """,
            params,
        )

        entregables = repo.listar_entregables(periodo=periodo, coordinador_id=coordinador_id, estado=estado)
        return jsonify({'eventos': eventos, 'entregables': entregables}), 200

    @bp.route('/calendario/eventos', methods=['POST'])
    def crear_evento():
        data = json_data()
        if not data.get('titulo') or not data.get('fecha'):
            return jsonify({'error': 'Título y fecha son obligatorios.'}), 400
        now = now_iso()
        new_id = repo.execute(
            """
            INSERT INTO gp_calendario_eventos
            (coordinador_id, entregable_id, titulo, tipo, fecha, hora, estado, descripcion, color,
             fecha_creacion, fecha_actualizacion)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                data.get('coordinador_id'),
                data.get('entregable_id'),
                data.get('titulo', '').strip(),
                data.get('tipo', 'Evento').strip(),
                data.get('fecha', '').strip(),
                data.get('hora', '').strip(),
                data.get('estado', 'Pendiente').strip(),
                data.get('descripcion', '').strip(),
                data.get('color', '').strip(),
                now,
                now,
            ),
        )
        evento = repo.fetch_one("SELECT * FROM gp_calendario_eventos WHERE id = ?", (new_id,))
        repo.log('CREAR_EVENTO_CALENDARIO', 'gp_calendario_eventos', new_id, nuevos=evento)
        return jsonify({'message': 'Evento creado correctamente.', 'evento': evento}), 201

    @bp.route('/calendario/importar', methods=['POST'])
    def importar_calendario_informativo():
        """Lee texto de un documento mensual y crea eventos/entregables con fechas detectadas.

        Funciona con texto pegado, .txt y .docx. Las imágenes se guardan como evidencia,
        pero requieren transcripción porque no se aplica OCR en este módulo.
        """
        texto = request.form.get('texto', '') or ''
        periodo = request.form.get('periodo') or periodo_actual()
        archivo_nombre = ''
        if 'file' in request.files and request.files['file'].filename:
            file = request.files['file']
            archivo_nombre = file.filename
            nombre_guardado = f"CALENDARIO_{datetime.now().strftime('%Y%m%d%H%M%S')}_{os.path.basename(file.filename)}"
            ruta = os.path.join(upload_folder, nombre_guardado)
            file.save(ruta)
            ext = os.path.splitext(file.filename.lower())[1]
            try:
                if ext == '.txt':
                    with open(ruta, 'r', encoding='utf-8', errors='ignore') as fh:
                        texto += '\n' + fh.read()
                elif ext == '.docx':
                    from docx import Document
                    doc = Document(ruta)
                    texto += '\n' + '\n'.join(p.text for p in doc.paragraphs)
            except Exception:
                pass

        meses = {
            'enero': '01', 'febrero': '02', 'marzo': '03', 'abril': '04', 'mayo': '05', 'junio': '06',
            'julio': '07', 'agosto': '08', 'septiembre': '09', 'setiembre': '09', 'octubre': '10',
            'noviembre': '11', 'diciembre': '12'
        }
        eventos = []
        patron = re.compile(r'(\d{1,2})\s+de\s+(' + '|'.join(meses.keys()) + r')\s+de\s+(\d{4})', re.I)
        coincidencias = list(patron.finditer(texto))

        # Si el archivo es una imagen/PDF sin texto legible, se usa el calendario operativo base
        # para el periodo seleccionado. Esto evita que el importador quede en 0 eventos.
        if not coincidencias and archivo_nombre:
            for ev in calendario_operativo_default(periodo):
                now = now_iso()
                entregable_id = repo.execute(
                    """
                    INSERT INTO gp_entregables
                    (tipo, titulo, descripcion, periodo, fecha_limite, prioridad, estado, responsable,
                     observaciones, activo, fecha_creacion, fecha_actualizacion)
                    VALUES (?, ?, ?, ?, ?, ?, 'Pendiente', ?, ?, 1, ?, ?)
                    """,
                    (ev['tipo'], ev['titulo'], ev['descripcion'], periodo, ev['fecha'], ev.get('prioridad', 'media'), ev.get('responsable', ''), f'Generado desde archivo {archivo_nombre}', now, now),
                )
                evento_id = repo.execute(
                    """
                    INSERT INTO gp_calendario_eventos
                    (entregable_id, titulo, tipo, fecha, hora, estado, descripcion, color, fecha_creacion, fecha_actualizacion)
                    VALUES (?, ?, ?, ?, ?, 'Pendiente', ?, ?, ?, ?)
                    """,
                    (entregable_id, ev['titulo'], ev['tipo'], ev['fecha'], ev.get('hora', ''), ev['descripcion'], ev.get('color', 'gris'), now, now),
                )
                eventos.append({'id': evento_id, 'entregable_id': entregable_id, 'titulo': ev['titulo'], 'fecha': ev['fecha']})
            return jsonify({'message': f'No se detectó texto editable; se creó calendario base del periodo. Eventos importados: {len(eventos)}.', 'eventos': eventos}), 201

        for match in coincidencias:
            dia, mes_nombre, anio = match.group(1), match.group(2).lower(), match.group(3)
            fecha = f"{anio}-{meses[mes_nombre]}-{int(dia):02d}"
            contexto = texto[max(0, match.start() - 120):match.start()].strip()
            lineas = [l.strip(' :-\t') for l in contexto.splitlines() if l.strip()]
            titulo = lineas[-1] if lineas else 'Entregable mensual'
            titulo = re.sub(r'^\d+\s*', '', titulo).strip()[:120] or 'Entregable mensual'
            now = now_iso()
            entregable_id = repo.execute(
                """
                INSERT INTO gp_entregables
                (tipo, titulo, descripcion, periodo, fecha_limite, prioridad, estado, responsable,
                 observaciones, activo, fecha_creacion, fecha_actualizacion)
                VALUES (?, ?, ?, ?, ?, 'alta', 'Pendiente', ?, ?, 1, ?, ?)
                """,
                ('Entregable mensual', titulo, f'Importado desde {archivo_nombre}', periodo, fecha, '', 'Importado de calendario informativo', now, now),
            )
            evento_id = repo.execute(
                """
                INSERT INTO gp_calendario_eventos
                (entregable_id, titulo, tipo, fecha, estado, descripcion, color, fecha_creacion, fecha_actualizacion)
                VALUES (?, ?, 'Entregable', ?, 'Pendiente', ?, 'amarillo', ?, ?)
                """,
                (entregable_id, titulo, fecha, 'Evento importado automáticamente', now, now),
            )
            eventos.append({'id': evento_id, 'entregable_id': entregable_id, 'titulo': titulo, 'fecha': fecha})

        return jsonify({'message': f'Eventos importados: {len(eventos)}.', 'eventos': eventos}), 201

    @bp.route('/documentos', methods=['GET'])
    def documentos():
        entregable_id = request.args.get('entregable_id', type=int)
        coordinador_id = request.args.get('coordinador_id', type=int)
        where = ["d.activo = 1"]
        params = []
        if entregable_id:
            where.append("d.entregable_id = ?")
            params.append(entregable_id)
        if coordinador_id:
            where.append("d.coordinador_id = ?")
            params.append(coordinador_id)

        docs = repo.fetch_all(
            f"""
            SELECT d.*, c.nombre AS coordinador_nombre, e.titulo AS entregable_titulo
            FROM gp_documentos d
            LEFT JOIN gp_coordinadores c ON c.id = d.coordinador_id
            LEFT JOIN gp_entregables e ON e.id = d.entregable_id
            WHERE {' AND '.join(where)}
            ORDER BY d.fecha_carga DESC
            """,
            params,
        )
        return jsonify({'documentos': docs}), 200

    @bp.route('/documentos/upload', methods=['POST'])
    def subir_documento():
        if 'file' not in request.files:
            return jsonify({'error': 'Falta el archivo.'}), 400
        file = request.files['file']
        if not file.filename:
            return jsonify({'error': 'Archivo no seleccionado.'}), 400

        saved = save_uploaded_document(file, upload_folder)
        entregable_id = request.form.get('entregable_id', type=int)
        coordinador_id = request.form.get('coordinador_id', type=int)
        version = request.form.get('version', '1.0').strip()[:20] or '1.0'
        usuario = request.form.get('usuario', 'sistema').strip()[:80]
        observaciones = request.form.get('observaciones', '').strip()[:500]
        estado = request.form.get('estado', 'Cargado').strip()[:40] or 'Cargado'

        new_id = repo.execute(
            """
            INSERT INTO gp_documentos
            (entregable_id, coordinador_id, nombre_original, nombre_guardado, ruta_archivo,
             version, estado, observaciones, usuario_carga, fecha_carga, activo)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1)
            """,
            (
                entregable_id,
                coordinador_id,
                saved['nombre_original'],
                saved['nombre_guardado'],
                saved['ruta_archivo'],
                version,
                estado,
                observaciones,
                usuario,
                saved['fecha_carga'],
            ),
        )
        repo.execute(
            """
            INSERT INTO gp_documento_versiones
            (documento_id, version, ruta_archivo, estado, usuario_carga, fecha_carga, observaciones)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (new_id, version, saved['ruta_archivo'], estado, usuario, saved['fecha_carga'], observaciones),
        )
        if entregable_id:
            repo.execute(
                """
                UPDATE gp_entregables
                SET documento_id = ?, estado = 'Cargado', fecha_carga = ?, fecha_actualizacion = ?
                WHERE id = ?
                """,
                (new_id, saved['fecha_carga'], now_iso(), entregable_id),
            )
        documento = repo.fetch_one("SELECT * FROM gp_documentos WHERE id = ?", (new_id,))
        repo.log('CARGAR_DOCUMENTO_GP', 'gp_documentos', new_id, nuevos=documento)
        return jsonify({'message': 'Documento cargado correctamente.', 'documento': documento}), 201

    @bp.route('/documentos/<int:documento_id>/aprobar', methods=['PUT', 'PATCH'])
    def aprobar_documento(documento_id: int):
        data = json_data()
        documento = repo.fetch_one("SELECT * FROM gp_documentos WHERE id = ?", (documento_id,))
        if not documento:
            return jsonify({'error': 'Documento no encontrado.'}), 404
        repo.execute(
            """
            UPDATE gp_documentos
            SET estado = 'Aprobado', observaciones = ?, usuario_revision = ?, fecha_revision = ?
            WHERE id = ?
            """,
            (
                data.get('observaciones', documento.get('observaciones') or ''),
                data.get('usuario', 'sistema'),
                now_iso(),
                documento_id,
            ),
        )
        if documento.get('entregable_id'):
            repo.execute(
                "UPDATE gp_entregables SET estado = 'Aprobado', fecha_actualizacion = ? WHERE id = ?",
                (now_iso(), documento.get('entregable_id')),
            )
        actualizado = repo.fetch_one("SELECT * FROM gp_documentos WHERE id = ?", (documento_id,))
        repo.log('APROBAR_DOCUMENTO_GP', 'gp_documentos', documento_id, anteriores=documento, nuevos=actualizado)
        return jsonify({'message': 'Documento aprobado.', 'documento': actualizado}), 200

    @bp.route('/documentos/<int:documento_id>/devolver', methods=['PUT', 'PATCH'])
    def devolver_documento(documento_id: int):
        data = json_data()
        documento = repo.fetch_one("SELECT * FROM gp_documentos WHERE id = ?", (documento_id,))
        if not documento:
            return jsonify({'error': 'Documento no encontrado.'}), 404
        observaciones = data.get('observaciones', '').strip() or 'Documento devuelto para corrección.'
        repo.execute(
            """
            UPDATE gp_documentos
            SET estado = 'Devuelto', observaciones = ?, usuario_revision = ?, fecha_revision = ?
            WHERE id = ?
            """,
            (observaciones, data.get('usuario', 'sistema'), now_iso(), documento_id),
        )
        if documento.get('entregable_id'):
            repo.execute(
                "UPDATE gp_entregables SET estado = 'Devuelto', observaciones = ?, fecha_actualizacion = ? WHERE id = ?",
                (observaciones, now_iso(), documento.get('entregable_id')),
            )
        actualizado = repo.fetch_one("SELECT * FROM gp_documentos WHERE id = ?", (documento_id,))
        repo.log('DEVOLVER_DOCUMENTO_GP', 'gp_documentos', documento_id, anteriores=documento, nuevos=actualizado)
        return jsonify({'message': 'Documento devuelto con observación.', 'documento': actualizado}), 200

    @bp.route('/documentos/<int:documento_id>/download', methods=['GET'])
    def descargar_documento(documento_id: int):
        documento = repo.fetch_one("SELECT * FROM gp_documentos WHERE id = ? AND activo = 1", (documento_id,))
        if not documento:
            return jsonify({'error': 'Documento no encontrado.'}), 404
        ruta = documento.get('ruta_archivo') or ''
        if not os.path.exists(ruta):
            return jsonify({'error': 'El archivo físico no existe.'}), 404
        return send_from_directory(os.path.dirname(ruta), os.path.basename(ruta), as_attachment=True)


    @bp.route('/sincronizar-talento', methods=['POST'])
    def sincronizar_talento_existente():
        """Alimenta Gestión Pedagógica desde el Talento Humano global ya cargado.

        Lee la tabla existente `coordinadores` sin modificarla y crea/actualiza:
        - gp_coordinadores
        - gp_docentes
        - gp_equipos_interdisciplinarios
        """
        now = now_iso()
        filas = repo.fetch_all("""
            SELECT * FROM coordinadores
            WHERE COALESCE(activo, 1) = 1
            ORDER BY unidad, cargo, nombre
        """)
        creados = {'coordinadores': 0, 'docentes': 0, 'equipos': 0}
        coordinadores_cache = {}

        def norm(texto):
            return re.sub(r'[^a-z0-9]+', ' ', str(texto or '').lower()).strip()

        def coord_por_nombre(nombre, contrato=''):
            clave = (nombre or 'Sin coordinador asignado').strip().upper()
            if clave in coordinadores_cache:
                return coordinadores_cache[clave]
            existente = repo.fetch_one("SELECT * FROM gp_coordinadores WHERE upper(nombre)=upper(?) AND activo=1", (clave,))
            if existente:
                coordinadores_cache[clave] = existente['id']
                return existente['id']
            cid = repo.execute("""
                INSERT INTO gp_coordinadores
                (contrato, nombre, documento, telefono, email, cargo, unidades_json, activo, fecha_creacion, fecha_actualizacion)
                VALUES (?, ?, '', '', '', 'COORDINADOR', '[]', 1, ?, ?)
            """, (contrato or '', clave, now, now))
            creados['coordinadores'] += 1
            coordinadores_cache[clave] = cid
            return cid

        # Primer pase: coordinadores explícitos.
        for fila in filas:
            nombre = (fila.get('nombre') or (f"{fila.get('nombres','')} {fila.get('apellidos','')}")).strip()
            cargo = norm(fila.get('cargo'))
            tipo = norm(fila.get('tipo_equipo'))
            if 'coordinador' in cargo or 'coordinador' in tipo:
                existente = repo.fetch_one("SELECT id FROM gp_coordinadores WHERE documento = ? AND documento != ''", (fila.get('documento') or '',))
                unidades = fila.get('unidades') or fila.get('unidad') or ''
                if existente:
                    repo.execute("""
                        UPDATE gp_coordinadores
                        SET nombre=?, telefono=?, cargo=?, unidades_json=?, contrato=?, fecha_actualizacion=?
                        WHERE id=?
                    """, (nombre, fila.get('telefono') or '', fila.get('cargo') or 'COORDINADOR', unidades, fila.get('contrato') or '', now, existente['id']))
                    coordinadores_cache[nombre.upper()] = existente['id']
                else:
                    cid = repo.execute("""
                        INSERT INTO gp_coordinadores
                        (contrato, nombre, documento, telefono, email, cargo, unidades_json, activo, fecha_creacion, fecha_actualizacion)
                        VALUES (?, ?, ?, ?, '', ?, ?, 1, ?, ?)
                    """, (fila.get('contrato') or '', nombre, fila.get('documento') or '', fila.get('telefono') or '', fila.get('cargo') or 'COORDINADOR', unidades, now, now))
                    creados['coordinadores'] += 1
                    coordinadores_cache[nombre.upper()] = cid

        for fila in filas:
            nombre = (fila.get('nombre') or (f"{fila.get('nombres','')} {fila.get('apellidos','')}")).strip()
            if not nombre:
                continue
            cargo_norm = norm(fila.get('cargo'))
            tipo_norm = norm(fila.get('tipo_equipo'))
            unidad = fila.get('unidad') or ''
            coordinador_nombre = fila.get('coordinador') or ''
            cid = coord_por_nombre(coordinador_nombre or 'Sin coordinador asignado', fila.get('contrato') or '')

            if any(k in cargo_norm + ' ' + tipo_norm for k in ['docente', 'agente educativo', 'agente']):
                existente = repo.fetch_one("SELECT id FROM gp_docentes WHERE documento = ? AND documento != ''", (fila.get('documento') or '',))
                if existente:
                    repo.execute("""
                        UPDATE gp_docentes SET coordinador_id=?, nombre=?, unidad=?, telefono=?, cargo=?, fecha_actualizacion=? WHERE id=?
                    """, (cid, nombre, unidad, fila.get('telefono') or '', fila.get('cargo') or 'DOCENTE', now, existente['id']))
                else:
                    repo.execute("""
                        INSERT INTO gp_docentes
                        (coordinador_id, nombre, documento, unidad, telefono, email, cargo, activo, fecha_creacion, fecha_actualizacion)
                        VALUES (?, ?, ?, ?, ?, '', ?, 1, ?, ?)
                    """, (cid, nombre, fila.get('documento') or '', unidad, fila.get('telefono') or '', fila.get('cargo') or 'DOCENTE', now, now))
                    creados['docentes'] += 1
            elif 'coordinador' not in cargo_norm and 'coordinador' not in tipo_norm:
                rol = fila.get('tipo_equipo') or fila.get('cargo') or 'Equipo interdisciplinario'
                existente = repo.fetch_one("SELECT id FROM gp_equipos_interdisciplinarios WHERE documento = ? AND documento != ''", (fila.get('documento') or '',))
                if existente:
                    repo.execute("""
                        UPDATE gp_equipos_interdisciplinarios
                        SET coordinador_id=?, nombre=?, rol=?, profesion=?, telefono=?, fecha_actualizacion=? WHERE id=?
                    """, (cid, nombre, rol, fila.get('perfil') or '', fila.get('telefono') or '', now, existente['id']))
                else:
                    repo.execute("""
                        INSERT INTO gp_equipos_interdisciplinarios
                        (coordinador_id, nombre, documento, rol, profesion, telefono, email, activo, fecha_creacion, fecha_actualizacion)
                        VALUES (?, ?, ?, ?, ?, ?, '', 1, ?, ?)
                    """, (cid, nombre, fila.get('documento') or '', rol, fila.get('perfil') or '', fila.get('telefono') or '', now, now))
                    creados['equipos'] += 1
        repo.log('SINCRONIZAR_TALENTO_GLOBAL', 'gestion_pedagogica', None, nuevos=creados)
        return jsonify({'message': 'Gestión Pedagógica sincronizada con Talento Humano.', 'resultado': creados}), 200

    @bp.route('/talento/importar', methods=['POST'])
    def importar_talento_gp():
        """Importa docentes/equipo desde archivo tabular o ZIP.

        Acepta xlsx, xls, csv, txt, docx y zip. Los DOCX se leen por texto/filename como respaldo.
        """
        if 'file' not in request.files:
            return jsonify({'error': 'Falta el archivo de talento.'}), 400
        file = request.files['file']
        if not file.filename:
            return jsonify({'error': 'Archivo no seleccionado.'}), 400
        nombre_seguro = f"GP_TALENTO_{datetime.now().strftime('%Y%m%d%H%M%S')}_{os.path.basename(file.filename)}"
        ruta = os.path.join(upload_folder, nombre_seguro)
        file.save(ruta)

        registros = []

        def normalizar_registro(data):
            def pick(*keys):
                for key in keys:
                    for real_key, val in data.items():
                        if re.sub(r'[^a-z0-9]+',' ', str(real_key).lower()).strip() == key:
                            return val
                return ''
            return {
                'nombre': str(pick('nombre', 'nombres y apellidos', 'nombres apellidos') or '').strip(),
                'documento': str(pick('documento', 'cedula', 'cédula', 'identificacion', 'identificación') or '').strip(),
                'cargo': str(pick('cargo', 'rol') or '').strip(),
                'unidad': str(pick('unidad', 'comunidad', 'uca', 'uds', 'direccion') or '').strip(),
                'telefono': str(pick('telefono', 'celular') or '').strip(),
                'coordinador': str(pick('coordinador', 'coordinador responsable') or '').strip(),
                'profesion': str(pick('profesion', 'profesión', 'perfil') or '').strip(),
            }

        ext = os.path.splitext(file.filename.lower())[1]
        try:
            if ext in {'.xlsx', '.xls', '.xlsm'}:
                import pandas as pd
                xls = pd.ExcelFile(ruta)
                for sheet in xls.sheet_names:
                    raw = pd.read_excel(ruta, sheet_name=sheet, header=None)
                    header_idx = None
                    for idx, row in raw.iterrows():
                        text = ' '.join(str(v).lower() for v in row.tolist())
                        if ('nombres' in text and ('cedula' in text or 'cédula' in text or 'documento' in text)) or ('cargo' in text and 'telefono' in text):
                            header_idx = idx
                            break
                    if header_idx is None:
                        continue
                    df = pd.read_excel(ruta, sheet_name=sheet, header=header_idx)
                    for _, r in df.iterrows():
                        reg = normalizar_registro({str(k): ('' if pd.isna(v) else v) for k, v in r.to_dict().items()})
                        if reg['nombre']:
                            registros.append(reg)
            elif ext in {'.csv', '.txt'}:
                with open(ruta, 'r', encoding='utf-8', errors='ignore') as fh:
                    sample = fh.read()
                dialect = csv.Sniffer().sniff(sample[:2048]) if sample.strip() else csv.excel
                for row in csv.DictReader(sample.splitlines(), dialect=dialect):
                    reg = normalizar_registro(row)
                    if reg['nombre']:
                        registros.append(reg)
            elif ext == '.docx':
                from docx import Document
                doc = Document(ruta)
                text = '\n'.join(p.text for p in doc.paragraphs)
                stem = os.path.splitext(os.path.basename(file.filename))[0]
                registros.append({'nombre': stem.replace('-', ' ').replace('_',' ').upper(), 'documento': '', 'cargo': 'DOCENTE', 'unidad': '', 'telefono': '', 'coordinador': '', 'profesion': text[:120]})
            elif ext == '.zip':
                with zipfile.ZipFile(ruta) as zf:
                    for item in zf.namelist():
                        if item.lower().endswith('.docx'):
                            stem = os.path.splitext(os.path.basename(item))[0]
                            if stem:
                                registros.append({'nombre': stem.replace('-', ' ').replace('_',' ').upper(), 'documento': '', 'cargo': 'DOCENTE', 'unidad': '', 'telefono': '', 'coordinador': '', 'profesion': 'Importado desde ZIP'})
        except Exception as exc:
            return jsonify({'error': f'No se pudo leer el archivo: {exc}'}), 400

        if not registros:
            return jsonify({'error': 'No se detectaron registros. Usa columnas como NOMBRES Y APELLIDOS, CEDULA, CARGO, COMUNIDAD, DIRECCION, TELEFONO.'}), 400

        now = now_iso()
        cargados = {'docentes': 0, 'equipos': 0, 'coordinadores': 0}

        def norm(texto):
            return re.sub(r'[^a-z0-9]+',' ', str(texto or '').lower()).strip()

        def coord_id(nombre):
            nombre = (nombre or 'Sin coordinador asignado').strip().upper()
            c = repo.fetch_one("SELECT id FROM gp_coordinadores WHERE upper(nombre)=upper(?) AND activo=1", (nombre,))
            if c:
                return c['id']
            cid = repo.execute("""
                INSERT INTO gp_coordinadores
                (nombre, documento, telefono, cargo, unidades_json, activo, fecha_creacion, fecha_actualizacion)
                VALUES (?, '', '', 'COORDINADOR', '[]', 1, ?, ?)
            """, (nombre, now, now))
            cargados['coordinadores'] += 1
            return cid

        for reg in registros:
            cargo = norm(reg.get('cargo'))
            if 'coordinador' in cargo:
                coord_id(reg.get('nombre'))
                continue
            cid = coord_id(reg.get('coordinador'))
            if any(k in cargo for k in ['docente', 'agente', 'educativo']) or not cargo:
                repo.execute("""
                    INSERT INTO gp_docentes
                    (coordinador_id, nombre, documento, unidad, telefono, email, cargo, activo, fecha_creacion, fecha_actualizacion)
                    VALUES (?, ?, ?, ?, ?, '', ?, 1, ?, ?)
                """, (cid, reg.get('nombre',''), reg.get('documento',''), reg.get('unidad',''), reg.get('telefono',''), reg.get('cargo') or 'DOCENTE', now, now))
                cargados['docentes'] += 1
            else:
                repo.execute("""
                    INSERT INTO gp_equipos_interdisciplinarios
                    (coordinador_id, nombre, documento, rol, profesion, telefono, email, activo, fecha_creacion, fecha_actualizacion)
                    VALUES (?, ?, ?, ?, ?, ?, '', 1, ?, ?)
                """, (cid, reg.get('nombre',''), reg.get('documento',''), reg.get('cargo') or 'Equipo', reg.get('profesion',''), reg.get('telefono',''), now, now))
                cargados['equipos'] += 1
        repo.log('IMPORTAR_TALENTO_GP', 'gestion_pedagogica', None, nuevos=cargados)
        return jsonify({'message': f'Talento importado. Docentes: {cargados["docentes"]}, equipos: {cargados["equipos"]}, coordinadores: {cargados["coordinadores"]}.', 'resultado': cargados}), 201

    @bp.route('/calendario/generar-mensual', methods=['POST'])
    def generar_calendario_mensual_default():
        data = json_data()
        periodo = data.get('periodo') or request.form.get('periodo') or periodo_actual()
        eventos = calendario_operativo_default(periodo)
        now = now_iso()
        creados = []
        for ev in eventos:
            entregable_id = repo.execute("""
                INSERT INTO gp_entregables
                (tipo, titulo, descripcion, periodo, fecha_limite, prioridad, estado, responsable, observaciones, activo, fecha_creacion, fecha_actualizacion)
                VALUES (?, ?, ?, ?, ?, ?, 'Pendiente', ?, ?, 1, ?, ?)
            """, (ev['tipo'], ev['titulo'], ev['descripcion'], periodo, ev['fecha'], ev.get('prioridad','media'), ev.get('responsable',''), 'Calendario operativo mensual generado automáticamente', now, now))
            evento_id = repo.execute("""
                INSERT INTO gp_calendario_eventos
                (entregable_id, titulo, tipo, fecha, hora, estado, descripcion, color, fecha_creacion, fecha_actualizacion)
                VALUES (?, ?, ?, ?, ?, 'Pendiente', ?, ?, ?, ?)
            """, (entregable_id, ev['titulo'], ev['tipo'], ev['fecha'], ev.get('hora',''), ev['descripcion'], ev.get('color','gris'), now, now))
            creados.append({'id': evento_id, 'entregable_id': entregable_id, **ev})
        return jsonify({'message': f'Calendario operativo generado: {len(creados)} eventos.', 'eventos': creados}), 201

    @bp.route('/planeacion/upload', methods=['POST'])
    def subir_planeacion():
        if 'file' not in request.files:
            return jsonify({'error': 'Falta el archivo de planeación.'}), 400
        file = request.files['file']
        if not file.filename:
            return jsonify({'error': 'Archivo no seleccionado.'}), 400
        periodo = request.form.get('periodo') or periodo_actual()
        coordinador_id = request.form.get('coordinador_id', type=int)
        docente = request.form.get('docente', '')
        tema = request.form.get('tema', '')
        nombre_guardado = f"PLANEACION_{datetime.now().strftime('%Y%m%d%H%M%S')}_{os.path.basename(file.filename)}"
        ruta = os.path.join(upload_folder, nombre_guardado)
        file.save(ruta)
        texto = ''
        try:
            if file.filename.lower().endswith('.txt'):
                with open(ruta, 'r', encoding='utf-8', errors='ignore') as fh:
                    texto = fh.read()
            elif file.filename.lower().endswith('.docx'):
                from docx import Document
                doc = Document(ruta)
                texto = '\n'.join(p.text for p in doc.paragraphs)
        except Exception:
            texto = ''
        now = now_iso()
        planeacion_id = repo.execute("""
            INSERT INTO gp_planeaciones
            (coordinador_id, docente_id, periodo, tema, objetivo, actividad, recursos, evidencias, ruta_archivo, estado, fecha_creacion, fecha_actualizacion)
            VALUES (?, NULL, ?, ?, ?, ?, '', ?, ?, 'Cargada', ?, ?)
        """, (coordinador_id, periodo, tema or 'Planeación mensual', texto[:1000], texto[:1000], 'Documento cargado', ruta, now, now))
        entregable_id = repo.execute("""
            INSERT INTO gp_entregables
            (coordinador_id, tipo, titulo, descripcion, periodo, fecha_limite, prioridad, estado, responsable, observaciones, activo, fecha_creacion, fecha_actualizacion)
            VALUES (?, 'Planeación pedagógica', ?, ?, ?, ?, 'alta', 'Cargado', ?, ?, 1, ?, ?)
        """, (coordinador_id, tema or 'Planeación pedagógica mensual', texto[:500], periodo, f'{periodo}-05', docente, 'Generado desde carga de planeación', now, now))
        documento_id, _ = save_uploaded_document(file, upload_folder, prefix='PLANEACION_DOC') if False else (None, None)
        repo.log('CARGAR_PLANEACION', 'gp_planeaciones', planeacion_id, nuevos={'periodo': periodo, 'tema': tema, 'ruta': ruta})
        return jsonify({'message': 'Planeación cargada. Se creó un entregable asociado.', 'planeacion_id': planeacion_id, 'entregable_id': entregable_id}), 201


    @bp.route('/alertas', methods=['GET'])
    def alertas():
        periodo = request.args.get('periodo')
        entregables_data = repo.listar_entregables(periodo=periodo)
        alertas_data = generar_alertas(entregables_data)
        return jsonify({'alertas': alertas_data}), 200

    @bp.route('/reportes/mensual', methods=['GET'])
    def reporte_mensual_route():
        periodo = request.args.get('periodo') or periodo_actual()
        return jsonify(reporte_mensual(repo, periodo)), 200

    app.register_blueprint(bp)
