from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "referencia_archivos" / "image001.png"
HEADER = ROOT / "encabezado_oficial.png"
OUTPUT = Path.home() / "Downloads" / "Formato listado encuentros EDITABLE.docx"


def borders(table):
    props = table._tbl.tblPr
    element = props.first_child_found_in("w:tblBorders")
    if element is None:
        element = OxmlElement("w:tblBorders")
        props.append(element)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:color"), "000000")
        element.append(node)


def set_cell(cell, text, bold=False, center=False, size=7):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold


with Image.open(SOURCE) as image:
    image.crop((0, 0, image.width, 116)).save(HEADER)

doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = Inches(0.18)
section.bottom_margin = Inches(0.18)
section.left_margin = Inches(0.18)
section.right_margin = Inches(0.18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
p.add_run().add_picture(str(HEADER), width=Inches(10.55))

title = doc.add_table(rows=1, cols=1)
title.alignment = WD_TABLE_ALIGNMENT.CENTER
borders(title)
set_cell(title.cell(0, 0), "LISTADO DE ASISTENCIA", bold=True, center=True, size=9)

info = doc.add_table(rows=4, cols=1)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
borders(info)
set_cell(info.cell(0, 0), "TEMA: {{tema}}", bold=True, size=7)
set_cell(info.cell(1, 0), "FECHA: {{fecha}}     HORA INICIO: {{hora_inicio}}     HORA FINAL: {{hora_final}}     UCA: {{unidad}}", bold=True, size=7)
set_cell(info.cell(2, 0), "PROFESIONAL: {{profesional}}                                      CARGO: {{cargo}}", bold=True, size=7)
set_cell(info.cell(3, 0), "MODALIDAD DE ATENCIÓN: PROPIA E INTERCULTURAL     SERVICIO DE ATENCIÓN: EDUCACIÓN INICIAL DIARIA", bold=True, size=7)

headers = ["No.", "NOMBRE DEL BENEFICIARIO", "No. DOCUMENTO", "TELÉFONO", "ACUDIENTE", "No. DOCUMENTO", "FIRMA"]
table = doc.add_table(rows=21, cols=7)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
borders(table)
widths = [0.36, 2.42, 1.08, 1.18, 2.38, 1.08, 1.55]
for index, (header, width) in enumerate(zip(headers, widths)):
    table.columns[index].width = Inches(width)
    set_cell(table.cell(0, index), header, bold=True, center=True, size=7)
for row_number in range(1, 21):
    set_cell(table.cell(row_number, 0), str(row_number), center=True, size=7)
    table.rows[row_number].height = Inches(0.27)

doc.save(OUTPUT)
print(OUTPUT)
